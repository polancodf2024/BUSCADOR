import streamlit as st
import pandas as pd
from datetime import datetime
import time
import requests
from xml.etree import ElementTree
import re
from collections import Counter
import numpy as np
import math
import string
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# CONFIGURACIÓN PARA STREAMLIT CLOUD
# ============================================================================

# spaCy completamente deshabilitado para Streamlit Cloud
DEPENDENCY_PARSING_AVAILABLE = False
NLP_MODEL = None

# Configurar BioBERT
AI_EMBEDDINGS_AVAILABLE = False
BIOMED_EMBEDDER = None

try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import HDBSCAN, KMeans
    from sklearn.decomposition import PCA
    import torch
    
    with st.spinner("🔄 Cargando BioBERT (primera vez puede tomar ~2 minutos)..."):
        BIOMED_EMBEDDER = SentenceTransformer(
            'pritamdeka/S-Biomed-Roberta-snli-multinli-stsb',
            device='cpu'
        )
    AI_EMBEDDINGS_AVAILABLE = True
    print("✅ BioBERT embeddings disponible")
except Exception as e:
    print(f"⚠️ BioBERT no disponible: {e}")

# Configuración de la página
st.set_page_config(
    page_title="PubMed AI Analyzer - Generador de Flavors Avanzado",
    page_icon="🧠",
    layout="centered"
)

# ============================================================================
# FUNCIONES DE UTILIDAD PARA PubMed
# ============================================================================

@st.cache_data(ttl=3600, show_spinner=False)
def get_abstract(pmid):
    """Obtiene el abstract de un artículo de PubMed con cache"""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    fetch_url = f"{base_url}efetch.fcgi"
    
    params = {
        "db": "pubmed",
        "id": pmid,
        "retmode": "xml",
        "rettype": "abstract"
    }
    
    try:
        response = requests.get(fetch_url, params=params, timeout=10)
        response.raise_for_status()
        
        root = ElementTree.fromstring(response.content)
        
        abstract_texts = []
        for abstract in root.findall(".//AbstractText"):
            if abstract.text:
                abstract_texts.append(abstract.text)
        
        return " ".join(abstract_texts) if abstract_texts else None
        
    except Exception as e:
        return None

def preprocess_text(text):
    """Preprocesa el texto para análisis NLP"""
    if not text:
        return ""
    
    text = text.lower()
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\d+', '', text)
    text = ' '.join(text.split())
    
    return text

# ============================================================================
# MEJORA 1: NER CON BIOBERT PARA EXTRACCIÓN DE ENTIDADES
# ============================================================================

def extract_entities_with_biobert(text):
    """Extrae entidades médicas usando BioBERT embeddings y matching semántico"""
    if not AI_EMBEDDINGS_AVAILABLE or not BIOMED_EMBEDDER or not text:
        return []
    
    try:
        # Lista de entidades médicas clave para buscar semánticamente
        entity_list = [
            'VP40', 'VP35', 'GP', 'glycoprotein', 'nucleoprotein', 'NP',
            'Ebola virus', 'EBOV', 'Zaire ebolavirus',
            'matrix protein', 'polymerase', 'polymerase cofactor',
            'interferon antagonist', 'immune evasion', 'viral replication',
            'mortality', 'fatality', 'survival', 'bleeding', 'hemorrhage',
            'efficacy', 'safety', 'toxicity', 'inhibitor', 'vaccine',
            'antibody', 'neutralization', 'docking', 'molecular dynamics'
        ]
        
        # Generar embeddings del texto y las entidades
        text_embedding = BIOMED_EMBEDDER.encode([text[:2000]])[0]
        entity_embeddings = BIOMED_EMBEDDER.encode(entity_list)
        
        # Calcular similitud semántica
        similarities = cosine_similarity([text_embedding], entity_embeddings)[0]
        
        # Extraer entidades con alta similitud
        entities = []
        for i, sim in enumerate(similarities):
            if sim > 0.65:  # Umbral de similitud semántica
                entities.append((entity_list[i], 'biobert_ner', sim))
        
        return entities
    except Exception as e:
        return []

def extract_entities_with_regex(text):
    """Extrae entidades usando regex como respaldo"""
    if not text:
        return []
    
    text_lower = text.lower()
    entities = []
    
    medical_patterns = {
        'proteins': r'\b(VP40|VP35|VP30|VP24|GP|NP|nucleoprotein|glycoprotein|matrix protein|polymerase)\b',
        'medications': r'\b(ticagrelor|clopidogrel|aspirin|prasugrel|heparin|warfarin|rivaroxaban|fendiline|galidesivir)\b',
        'conditions': r'\b(ebola|ebov|ebolavirus|hemorrhagic fever|acute coronary syndrome|myocardial infarction)\b',
        'methods': r'\b(machine learning|deep learning|molecular dynamics|docking|simulation|bioinformatics|computational)\b',
        'outcomes': r'\b(mortality|death|survival|bleeding|hemorrhage|efficacy|safety|toxicity|effectiveness)\b',
        'study_types': r'\b(randomized|rct|cohort|prospective|retrospective|observational|meta-analysis)\b'
    }
    
    for entity_type, pattern in medical_patterns.items():
        matches = re.findall(pattern, text_lower, re.IGNORECASE)
        for match in matches:
            entities.append((match, entity_type, 1.0))
    
    return entities

def extract_medical_entities_enhanced(text):
    """Extrae entidades combinando BioBERT NER y regex"""
    if not text:
        return []
    
    # Intentar primero con BioBERT
    biobert_entities = extract_entities_with_biobert(text)
    
    # Complementar con regex
    regex_entities = extract_entities_with_regex(text)
    
    # Combinar y eliminar duplicados (por nombre)
    all_entities = []
    seen = set()
    
    for entity, etype, score in biobert_entities:
        if entity.lower() not in seen:
            all_entities.append((entity, etype, score))
            seen.add(entity.lower())
    
    for entity, etype, score in regex_entities:
        if entity.lower() not in seen:
            all_entities.append((entity, etype, score))
            seen.add(entity.lower())
    
    return all_entities

def extract_numeric_results(text):
    """Extrae resultados numéricos importantes"""
    if not text:
        return []
    
    results = []
    
    hr_pattern = r'(?:HR|hazard ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(hr_pattern, text, re.IGNORECASE):
        results.append({'type': 'HR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    
    rr_pattern = r'(?:RR|risk ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(rr_pattern, text, re.IGNORECASE):
        results.append({'type': 'RR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    
    or_pattern = r'(?:OR|odds ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(or_pattern, text, re.IGNORECASE):
        results.append({'type': 'OR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    
    p_pattern = r'[Pp]\s*[=<]\s*([0-9]+\.[0-9]+)'
    for match in re.findall(p_pattern, text):
        results.append({'type': 'p-value', 'value': float(match)})
    
    return results

def extract_all_outcomes(text):
    """Extrae todos los outcomes mencionados"""
    if not text:
        return []
    
    text_lower = text.lower()
    outcomes = []
    
    outcome_patterns = [
        r'\b(mortality|death|survival)\b',
        r'\b(bleeding|hemorrhage)\b',
        r'\b(stroke|cva|tia)\b',
        r'\b(reinfarction|mi|myocardial infarction)\b',
        r'\b(mace)\b',
        r'\b(efficacy|safety|effectiveness)\b',
        r'\b(complication|adverse event)\b',
        r'\b(risk|rate|incidence)\b',
        r'\b(outcome|endpoint)\b',
        r'\b(recovery|improvement)\b'
    ]
    
    for pattern in outcome_patterns:
        matches = re.findall(pattern, text_lower)
        outcomes.extend(matches)
    
    return list(set(outcomes))

def analyze_sentiment_by_outcome(text):
    """Analiza el sentimiento para cada outcome con detección de negación mejorada"""
    if not text:
        return {}
    
    text_lower = text.lower()
    outcomes_sentiment = {}
    all_outcomes = extract_all_outcomes(text)
    
    negation_words = ['no', 'not', 'without', 'lack', 'failed', 'non-significant', 'did not', 'does not']
    intensifiers = ['significantly', 'markedly', 'dramatically', 'strongly']
    
    for outcome in all_outcomes:
        outcome_index = text_lower.find(outcome)
        if outcome_index == -1:
            continue
            
        start = max(0, outcome_index - 300)
        end = min(len(text_lower), outcome_index + 300)
        context = text_lower[start:end]
        
        positive_terms = ['reduced', 'decreased', 'lower', 'improved', 'better', 'benefit', 
                         'protective', 'effective', 'efficacious', 'superior', 'advantage']
        negative_terms = ['increased', 'higher', 'worse', 'elevated', 'risk', 'adverse', 
                         'harm', 'detrimental', 'inferior', 'disadvantage']
        
        sentiment_score = 0
        
        # Buscar términos positivos con intensificadores
        for term in positive_terms:
            if term in context:
                term_pos = context.find(term)
                prev_words = context[max(0, term_pos-50):term_pos]
                # Verificar negación
                if any(neg in prev_words for neg in negation_words):
                    sentiment_score -= 1
                else:
                    # Verificar intensificador
                    if any(inten in prev_words for inten in intensifiers):
                        sentiment_score += 2
                    else:
                        sentiment_score += 1
        
        # Buscar términos negativos
        for term in negative_terms:
            if term in context:
                term_pos = context.find(term)
                prev_words = context[max(0, term_pos-50):term_pos]
                if any(neg in prev_words for neg in negation_words):
                    sentiment_score += 1
                else:
                    if any(inten in prev_words for inten in intensifiers):
                        sentiment_score -= 2
                    else:
                        sentiment_score -= 1
        
        if sentiment_score > 0:
            outcomes_sentiment[outcome] = 'BENEFICIO'
        elif sentiment_score < 0:
            outcomes_sentiment[outcome] = 'RIESGO'
        else:
            outcomes_sentiment[outcome] = 'SIN EFECTO'
    
    return outcomes_sentiment

def enhanced_quality_scoring(study_types, full_text):
    """Sistema de puntuación de calidad mejorado"""
    score = 0
    factors = []
    text_lower = full_text.lower() if full_text else ""
    
    study_type_weights = {'Meta-análisis': 2.5, 'RCT': 2.0, 'Cohorte': 1.6, 'Caso-control': 1.4, 'Observacional': 1.2}
    
    for study_type in study_types:
        if study_type in study_type_weights:
            score += study_type_weights[study_type] * 20
            factors.append(f"Tipo estudio: {study_type}")
    
    sample_patterns = [r'n\s*[=:]\s*(\d+)', r'sample size\s*[=:]\s*(\d+)']
    max_sample = 0
    for pattern in sample_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                max_sample = max(max_sample, int(match))
            except:
                pass
    
    if max_sample > 0:
        if max_sample > 1000:
            score += 20
            factors.append(f"Muestra grande: n={max_sample}")
        elif max_sample > 500:
            score += 15
        elif max_sample > 200:
            score += 10
        elif max_sample > 100:
            score += 5
    
    if any(term in text_lower for term in ['multicenter', 'multi-center']):
        score += 10
        factors.append("Estudio multicéntrico")
    
    if any(term in text_lower for term in ['double-blind', 'double blind', 'masked']):
        score += 15
        factors.append("Con enmascaramiento")
    
    if any(term in text_lower for term in ['intention-to-treat', 'intention to treat']):
        score += 10
        factors.append("ITT analysis")
    
    return min(score, 100), factors

def get_effect_direction(result_value, ci_lower=None, ci_upper=None):
    """Determina la direccionalidad del efecto"""
    try:
        value = float(result_value)
        if value < 1:
            if ci_upper and float(ci_upper) < 1:
                return "PROTECTOR"
            else:
                return "PROTECTOR (tendencia)"
        elif value > 1:
            if ci_lower and float(ci_lower) > 1:
                return "DAÑINO"
            else:
                return "DAÑINO (tendencia)"
        else:
            return "SIN EFECTO"
    except:
        return "NO DETERMINADA"

def calculate_evidence_strength(statistical_results, quality_score):
    """Calcula la fuerza de la evidencia"""
    if not statistical_results:
        return "Sin datos", 0, {}
    
    strength_score = 0
    directions = {}
    
    for i, result in enumerate(statistical_results):
        if result['type'] in ['HR', 'RR', 'OR']:
            effect = result['value']
            if effect < 0.5 or effect > 2.0:
                strength_score += 30
            elif effect < 0.7 or effect > 1.5:
                strength_score += 20
            elif effect != 1.0:
                strength_score += 10
            
            directions[f"result_{i}"] = get_effect_direction(effect, result.get('ci_lower'), result.get('ci_upper'))
        
        if result['type'] == 'p-value':
            if result['value'] < 0.001:
                strength_score += 25
            elif result['value'] < 0.01:
                strength_score += 20
            elif result['value'] < 0.05:
                strength_score += 15
    
    strength_score = strength_score * (quality_score / 50)
    
    if strength_score >= 50:
        strength = "EVIDENCIA FUERTE"
    elif strength_score >= 30:
        strength = "EVIDENCIA MODERADA"
    elif strength_score >= 15:
        strength = "EVIDENCIA DÉBIL"
    else:
        strength = "EVIDENCIA INSUFICIENTE"
    
    return strength, min(strength_score, 100), directions

def analyze_article_with_ai(title, abstract):
    """Analiza un artículo con IA mejorada"""
    if not title and not abstract:
        return {}
    
    full_text = f"{title} {abstract if abstract else ''}"
    processed_text = preprocess_text(full_text)
    
    # Extracción mejorada con BioBERT NER
    entities = extract_medical_entities_enhanced(full_text)
    
    study_types = []
    study_keywords = {
        'RCT': ['randomized', 'rct', 'randomised', 'trial', 'double-blind'],
        'Cohorte': ['cohort', 'follow-up', 'longitudinal', 'prospective', 'retrospective'],
        'Meta-análisis': ['meta-analysis', 'meta analysis', 'systematic review'],
        'Observacional': ['observational', 'registry', 'real-world'],
        'Caso-control': ['case-control', 'matched']
    }
    
    for study_type, keywords in study_keywords.items():
        if any(keyword in processed_text for keyword in keywords):
            study_types.append(study_type)
    
    words = processed_text.split()
    word_freq = Counter(words)
    top_keywords = [word for word, _ in word_freq.most_common(20) if len(word) > 3][:10]
    
    population_patterns = {
        'adultos': r'\badults?\b|\bpatients?\b',
        'ancianos': r'\belderly\b|\baged\b|\bolder\b',
        'mujeres': r'\bwomen\b|\bfemale\b',
        'hombres': r'\bmen\b|\bmale\b'
    }
    
    population = [pop for pop, pattern in population_patterns.items() if re.search(pattern, processed_text)]
    
    numeric_results = extract_numeric_results(full_text)
    all_outcomes = extract_all_outcomes(full_text)
    
    quality_score, quality_factors = enhanced_quality_scoring(study_types, full_text)
    
    evidence_strength, evidence_score, directions = calculate_evidence_strength(numeric_results, quality_score)
    
    sentiment = analyze_sentiment_by_outcome(full_text)
    
    numeric_results_str = ' | '.join([f"{r['type']}={r['value']}" for r in numeric_results[:3]])
    
    return {
        'entities': entities,
        'study_types': ', '.join(study_types) if study_types else 'No especificado',
        'top_keywords': ', '.join(top_keywords),
        'population': ', '.join(population) if population else 'No especificada',
        'quality_score': quality_score,
        'quality_factors': ' | '.join(quality_factors) if quality_factors else 'Sin factores',
        'numeric_results': numeric_results,
        'numeric_results_str': numeric_results_str,
        'outcomes_analysis': sentiment,
        'all_outcomes': all_outcomes,
        'evidence_strength': evidence_strength,
        'evidence_score': evidence_score,
        'effect_directions': directions,
        'abstract': abstract if abstract else ''
    }

def extract_article_info(doc_sum):
    """Extrae la información básica de un artículo desde DocSum"""
    article = {}
    article["pmid"] = doc_sum.find("Id").text if doc_sum.find("Id") is not None else "N/A"
    
    title_item = doc_sum.find(".//Item[@Name='Title']")
    article["title"] = title_item.text if title_item is not None else "Sin título"
    
    author_items = doc_sum.findall(".//Item[@Name='Author']")
    authors = [author.text for author in author_items if author.text]
    article["authors"] = ", ".join(authors[:5]) + (" et al." if len(authors) > 5 else "")
    
    source_item = doc_sum.find(".//Item[@Name='Source']")
    article["journal"] = source_item.text if source_item is not None else "N/A"
    
    pubdate_item = doc_sum.find(".//Item[@Name='PubDate']")
    article["pubdate"] = pubdate_item.text if pubdate_item is not None else "N/A"
    
    doi_item = doc_sum.find(".//Item[@Name='DOI']")
    article["doi"] = doi_item.text if doi_item is not None else "N/A"
    
    return article

def search_pubmed(query, retmax=500):
    """Busca artículos en PubMed"""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    
    search_url = f"{base_url}esearch.fcgi"
    search_params = {
        "db": "pubmed",
        "term": query,
        "retmax": retmax,
        "retmode": "xml",
        "sort": "relevance"
    }
    
    try:
        response = requests.get(search_url, params=search_params, timeout=30)
        response.raise_for_status()
        
        root = ElementTree.fromstring(response.content)
        id_list = [id_elem.text for id_elem in root.findall(".//Id")]
        count = root.find(".//Count").text if root.find(".//Count") is not None else "0"
        
        return id_list, int(count)
    except Exception as e:
        st.error(f"Error en búsqueda: {e}")
        return [], 0

def fetch_articles_details(id_list, max_articles=500):
    """Obtiene los detalles de los artículos y los analiza"""
    if not id_list:
        return []
    
    total_to_process = min(len(id_list), max_articles)
    batch_size = 50
    num_batches = math.ceil(total_to_process / batch_size)
    
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    articles = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for batch_num in range(num_batches):
        start_idx = batch_num * batch_size
        end_idx = min((batch_num + 1) * batch_size, total_to_process)
        batch_ids = id_list[start_idx:end_idx]
        
        status_text.text(f"📦 Procesando lote {batch_num + 1} de {num_batches}...")
        
        summary_params = {
            "db": "pubmed",
            "id": ",".join(batch_ids),
            "retmode": "xml"
        }
        
        try:
            summary_response = requests.get(f"{base_url}esummary.fcgi", params=summary_params, timeout=30)
            summary_response.raise_for_status()
            summary_root = ElementTree.fromstring(summary_response.content)
            
            for j, doc_sum in enumerate(summary_root.findall(".//DocSum")):
                overall_idx = start_idx + j
                progress_bar.progress((overall_idx + 1) / total_to_process)
                
                article = extract_article_info(doc_sum)
                abstract = get_abstract(article["pmid"])
                article["abstract"] = abstract if abstract else "No disponible"
                
                ai_analysis = analyze_article_with_ai(article["title"], abstract)
                article.update(ai_analysis)
                
                articles.append(article)
                time.sleep(0.05)
                
        except Exception as e:
            st.warning(f"Error en lote {batch_num + 1}: {str(e)[:100]}")
            continue
    
    progress_bar.empty()
    status_text.empty()
    
    return articles

# ============================================================================
# MEJORA 2: FILTRADO CON EMBEDDINGS (BÚSQUEDA + HIPÓTESIS)
# ============================================================================

def calculate_relevance_to_search_and_hypothesis(articles, query, hypothesis):
    """Calcula relevancia a la búsqueda y a la hipótesis usando embeddings"""
    if not AI_EMBEDDINGS_AVAILABLE or not BIOMED_EMBEDDER:
        return articles
    
    try:
        # Limitar longitud de query
        query_text = query[:500] if len(query) > 500 else query
        query_embedding = BIOMED_EMBEDDER.encode([query_text])[0]
        hypothesis_embedding = BIOMED_EMBEDDER.encode([hypothesis])[0]
        
        for article in articles:
            text = f"{article.get('title', '')} {article.get('abstract', '')}"
            if text and len(text) > 50:
                article_embedding = BIOMED_EMBEDDER.encode([text[:1000]])[0]
                
                # Calcular similitud con la búsqueda
                search_sim = cosine_similarity([query_embedding], [article_embedding])[0][0]
                
                # Calcular similitud con la hipótesis
                hypothesis_sim = cosine_similarity([hypothesis_embedding], [article_embedding])[0][0]
                
                # Puntuación combinada (peso 0.3 para búsqueda, 0.7 para hipótesis)
                article['relevance_score'] = (search_sim * 0.3) + (hypothesis_sim * 0.7)
                article['search_relevance'] = float(search_sim)
                article['hypothesis_relevance'] = float(hypothesis_sim)
            else:
                article['relevance_score'] = 0.0
                article['search_relevance'] = 0.0
                article['hypothesis_relevance'] = 0.0
        
        articles.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
    except Exception as e:
        st.warning(f"Error calculando relevancia: {e}")
        for article in articles:
            article['relevance_score'] = 0.5
    
    return articles

# ============================================================================
# MEJORA 3: HDBSCAN PARA CLUSTERING MÁS NATURAL
# ============================================================================

def discover_flavors_by_embeddings_hdbscan(articles):
    """Descubre flavors usando HDBSCAN (clusters de forma arbitraria)"""
    if len(articles) < 3 or not AI_EMBEDDINGS_AVAILABLE or not BIOMED_EMBEDDER:
        return []
    
    try:
        texts = [f"{a.get('title', '')} {a.get('all_outcomes', [])}" for a in articles]
        embeddings = BIOMED_EMBEDDER.encode(texts)
        
        # Reducción de dimensionalidad con PCA para mejorar HDBSCAN
        pca = PCA(n_components=min(50, len(embeddings) - 1))
        embeddings_reduced = pca.fit_transform(embeddings)
        
        # HDBSCAN para clustering sin número fijo de clusters
        clusterer = HDBSCAN(min_cluster_size=2, min_samples=1, metric='euclidean')
        cluster_labels = clusterer.fit_predict(embeddings_reduced)
        
        flavors = []
        for cluster_id in set(cluster_labels):
            if cluster_id == -1:  # Ruido - artículos no asignados
                continue
            
            cluster_indices = [i for i, label in enumerate(cluster_labels) if label == cluster_id]
            if len(cluster_indices) < 2:
                continue
            
            cluster_articles = [articles[i] for i in cluster_indices]
            name = generate_descriptive_name(cluster_articles)
            summary = generate_flavor_summary(cluster_articles, name)
            representative = sorted(cluster_articles, key=lambda x: x.get('quality_score', 0), reverse=True)[:5]
            
            flavors.append({
                'type': 'semantic',
                'id': f"semantic_cluster_{cluster_id}",
                'name': name,
                'summary': summary,
                'articles': cluster_articles,
                'n_articles': len(cluster_articles),
                'representative_articles': representative
            })
        
        return flavors
    except Exception as e:
        st.warning(f"Error en clustering HDBSCAN: {e}")
        return []

def generate_descriptive_name(articles):
    """Genera un nombre descriptivo para un cluster de artículos"""
    if not articles:
        return "General"
    
    all_keywords = []
    all_outcomes = []
    all_proteins = []
    
    for a in articles:
        title = a.get('title', '').lower()
        keywords = a.get('top_keywords', '').split(', ')
        all_keywords.extend(keywords)
        outcomes = a.get('all_outcomes', [])
        all_outcomes.extend(outcomes)
        
        if 'vp40' in title:
            all_proteins.append('VP40')
        if 'vp35' in title:
            all_proteins.append('VP35')
        if 'gp' in title or 'glycoprotein' in title:
            all_proteins.append('GP')
        if 'np' in title or 'nucleoprotein' in title:
            all_proteins.append('NP')
        if 'vaccine' in title:
            all_proteins.append('vaccine')
        if 'machine learning' in title or 'ml' in title:
            all_proteins.append('ML')
        if 'drug' in title or 'inhibitor' in title:
            all_proteins.append('drug discovery')
        if 'structure' in title or 'dynamics' in title or 'simulation' in title:
            all_proteins.append('structure')
    
    protein_counts = Counter(all_proteins)
    outcome_counts = Counter(all_outcomes)
    
    main_topic = ''
    if protein_counts:
        main_topic = protein_counts.most_common(1)[0][0]
    
    if not main_topic and all_keywords:
        main_topic = all_keywords[0].title() if all_keywords else ''
    
    if main_topic:
        name = main_topic.upper()
    else:
        name = "Research"
    
    if outcome_counts:
        top_outcomes = [o for o, _ in outcome_counts.most_common(2)]
        if top_outcomes:
            name += f": {', '.join(top_outcomes)}"
    
    return name

def generate_flavor_summary(articles, flavor_name):
    """Genera un párrafo resumen para el flavor basado en los artículos"""
    if not articles:
        return "No articles available for this flavor."
    
    main_topics = []
    main_outcomes = []
    key_findings = []
    study_types = []
    
    for a in articles[:5]:
        title = a.get('title', '')
        outcomes = a.get('all_outcomes', [])
        study_type = a.get('study_types', '')
        numeric = a.get('numeric_results_str', '')
        
        if 'VP40' in title:
            main_topics.append('VP40')
        if 'VP35' in title:
            main_topics.append('VP35')
        if 'GP' in title or 'glycoprotein' in title:
            main_topics.append('glycoprotein')
        if 'vaccine' in title.lower():
            main_topics.append('vaccine')
        if 'drug' in title.lower() or 'inhibitor' in title.lower():
            main_topics.append('drug discovery')
        
        main_outcomes.extend(outcomes[:2])
        
        if numeric:
            key_findings.append(numeric)
        
        if study_type:
            study_types.append(study_type)
    
    topic_counts = Counter(main_topics)
    main_topic = topic_counts.most_common(1)[0][0] if topic_counts else 'research'
    
    outcome_counts = Counter(main_outcomes)
    top_outcomes = [o for o, _ in outcome_counts.most_common(2)]
    
    if 'VP40' in main_topic or 'VP35' in main_topic or 'GP' in main_topic:
        summary = f"Computational and bioinformatic studies on {main_topic} have revealed critical insights into its structural dynamics, protein-protein interactions, and potential as a therapeutic target. "
        
        if top_outcomes:
            summary += f"Key outcomes evaluated include {', '.join(top_outcomes)}. "
        
        if key_findings:
            summary += f"Notable findings include {key_findings[0] if key_findings else 'significant computational predictions'}. "
        
        summary += f"These {len(articles)} studies demonstrate the power of machine learning and molecular dynamics simulations in characterizing Ebola virus proteins and identifying potential inhibitors."
    
    elif 'vaccine' in main_topic.lower():
        summary = f"Vaccine development studies have employed computational approaches to predict immunogenicity, identify conserved epitopes, and evaluate vaccine candidates. "
        
        if top_outcomes:
            summary += f"Key outcomes assessed include {', '.join(top_outcomes)}. "
        
        summary += f"The {len(articles)} studies in this flavor highlight the role of bioinformatics in accelerating vaccine design against Ebola virus."
    
    elif 'drug' in main_topic.lower() or 'discovery' in main_topic.lower():
        summary = f"Drug discovery efforts leveraging computational screening, molecular docking, and machine learning have identified promising candidates targeting Ebola virus proteins. "
        
        if top_outcomes:
            summary += f"Efficacy and safety outcomes were evaluated across {len(articles)} studies. "
        
        summary += f"These in silico approaches provide a foundation for experimental validation and drug repurposing strategies."
    
    else:
        summary = f"This flavor comprises {len(articles)} computational and bioinformatic studies investigating Ebola virus biology, protein function, and therapeutic interventions. "
        
        if top_outcomes:
            summary += f"Key outcomes examined include {', '.join(top_outcomes)}. "
        
        summary += f"Machine learning algorithms and molecular dynamics simulations were employed to predict protein interactions, identify druggable sites, and accelerate discovery."
    
    return summary

def discover_flavors_by_outcomes(articles):
    """Agrupa artículos por outcomes compartidos"""
    if len(articles) < 2:
        return []
    
    try:
        all_outcomes = set()
        for article in articles:
            outcomes = article.get('all_outcomes', [])
            all_outcomes.update(outcomes)
        
        if len(all_outcomes) < 2:
            return []
        
        all_outcomes_list = list(all_outcomes)
        outcome_matrix = []
        for article in articles:
            outcomes = set(article.get('all_outcomes', []))
            row = [1 if o in outcomes else 0 for o in all_outcomes_list]
            outcome_matrix.append(row)
        
        n_clusters = min(max(2, len(articles) // 4), 6)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clusters = kmeans.fit_predict(outcome_matrix)
        
        flavors = []
        for cluster_id in range(n_clusters):
            cluster_indices = [i for i, c in enumerate(clusters) if c == cluster_id]
            if len(cluster_indices) < 2:
                continue
            
            cluster_articles = [articles[i] for i in cluster_indices]
            
            outcome_counts = Counter()
            for article in cluster_articles:
                outcomes = article.get('all_outcomes', [])
                outcome_counts.update(outcomes)
            
            top_outcomes = [o for o, _ in outcome_counts.most_common(3)]
            name = f"Studies on {', '.join(top_outcomes)}"
            summary = generate_flavor_summary(cluster_articles, name)
            representative = sorted(cluster_articles, key=lambda x: x.get('quality_score', 0), reverse=True)[:5]
            
            flavors.append({
                'type': 'outcome_based',
                'id': f"outcome_cluster_{cluster_id}",
                'name': name,
                'summary': summary,
                'top_outcomes': top_outcomes,
                'articles': cluster_articles,
                'n_articles': len(cluster_articles),
                'representative_articles': representative
            })
        
        return flavors
    except Exception as e:
        st.warning(f"Error en clustering por outcomes: {e}")
        return []

def discover_flavors_by_effect_direction(articles):
    """Agrupa artículos por dirección del efecto"""
    flavors = []
    
    direction_groups = {
        'beneficio': [],
        'riesgo': [],
        'sin_efecto': [],
        'contradictorio': []
    }
    
    for article in articles:
        directions = article.get('effect_directions', {})
        direction_values = list(directions.values()) if directions else []
        
        direction_str = ' '.join(direction_values).lower()
        
        if 'protector' in direction_str and 'dañino' not in direction_str:
            direction_groups['beneficio'].append(article)
        elif 'dañino' in direction_str and 'protector' not in direction_str:
            direction_groups['riesgo'].append(article)
        elif 'sin efecto' in direction_str:
            direction_groups['sin_efecto'].append(article)
        else:
            direction_groups['contradictorio'].append(article)
    
    direction_names = {
        'beneficio': 'Beneficial Effects',
        'riesgo': 'Risk Factors',
        'sin_efecto': 'No Significant Effect',
        'contradictorio': 'Mixed / Contradictory Results'
    }
    
    direction_summaries = {
        'beneficio': 'These studies consistently demonstrate protective or beneficial effects of the interventions evaluated, supporting their potential for therapeutic development.',
        'riesgo': 'These studies identify significant risk factors or adverse effects associated with the interventions or conditions studied.',
        'sin_efecto': 'These studies found no significant effect of the interventions on the outcomes measured, suggesting limited therapeutic benefit.',
        'contradictorio': 'These studies show mixed, inconsistent, or contradictory results, highlighting areas where further research is needed to resolve uncertainties.'
    }
    
    for key, group_articles in direction_groups.items():
        if group_articles:
            name = direction_names[key]
            summary = direction_summaries[key]
            representative = sorted(group_articles, key=lambda x: x.get('quality_score', 0), reverse=True)[:5]
            
            flavors.append({
                'type': 'effect_direction',
                'id': f"effect_{key}",
                'name': name,
                'summary': summary,
                'articles': group_articles,
                'n_articles': len(group_articles),
                'representative_articles': representative
            })
    
    return flavors

def discover_flavors_by_population(articles):
    """Agrupa artículos por población estudiada"""
    population_patterns = {
        'elderly': ['elderly', 'aged', 'older', 'geriatric', '>75'],
        'diabetes': ['diabetes', 'diabetic', 'dm'],
        'women': ['women', 'female'],
        'men': ['men', 'male'],
        'children': ['children', 'pediatric', 'paediatric', 'infant']
    }
    
    flavors = []
    
    population_summaries = {
        'elderly': 'Studies focusing on elderly populations reveal age-specific considerations for disease progression, treatment response, and vaccine efficacy.',
        'diabetes': 'Research in diabetic patients highlights how metabolic conditions may influence disease susceptibility and therapeutic outcomes.',
        'women': 'Studies examining female populations provide insights into gender-specific responses to infection and treatment.',
        'men': 'Research in male populations contributes to understanding gender-based differences in disease outcomes.',
        'children': 'Pediatric studies address the unique aspects of disease presentation, progression, and treatment in younger populations.'
    }
    
    for pop, patterns in population_patterns.items():
        pop_articles = []
        for article in articles:
            text = f"{article.get('title', '')} {article.get('abstract', '')}".lower()
            if any(pattern in text for pattern in patterns):
                pop_articles.append(article)
        
        if pop_articles:
            pop_names = {
                'elderly': 'Elderly Population',
                'diabetes': 'Diabetic Patients',
                'women': 'Female Population',
                'men': 'Male Population',
                'children': 'Pediatric Population'
            }
            
            summary = population_summaries.get(pop, f"Studies focusing on {pop_names.get(pop, pop)}.")
            representative = sorted(pop_articles, key=lambda x: x.get('quality_score', 0), reverse=True)[:5]
            
            flavors.append({
                'type': 'population',
                'id': f"pop_{pop}",
                'name': pop_names.get(pop, pop),
                'summary': summary,
                'articles': pop_articles,
                'n_articles': len(pop_articles),
                'representative_articles': representative
            })
    
    return flavors

def generate_citation(article, index):
    """Genera una cita bibliográfica en formato Vancouver"""
    authors = article.get('authors', 'Autor')
    year = article.get('pubdate', 's.f.')[:4] if article.get('pubdate') else 's.f.'
    title = article.get('title', 'Sin título')
    journal = article.get('journal', 'Revista')
    pmid = article.get('pmid', '')
    
    citation = f"{index}. {authors}. {title}. {journal}. {year}"
    if pmid and pmid != 'N/A':
        citation += f"; PMID: {pmid}"
    
    return citation

def generate_all_flavors(articles, hypothesis):
    """Genera todos los flavors desde múltiples perspectivas usando técnicas mejoradas"""
    if not articles:
        return {}
    
    all_flavors = {
        'semantic_clusters': [],
        'outcome_clusters': [],
        'effect_clusters': [],
        'population_clusters': []
    }
    
    # Usar HDBSCAN para clustering semántico (mejor que KMeans)
    if len(articles) >= 3 and AI_EMBEDDINGS_AVAILABLE:
        all_flavors['semantic_clusters'] = discover_flavors_by_embeddings_hdbscan(articles)
    
    if len(articles) >= 2:
        all_flavors['outcome_clusters'] = discover_flavors_by_outcomes(articles)
        all_flavors['effect_clusters'] = discover_flavors_by_effect_direction(articles)
        all_flavors['population_clusters'] = discover_flavors_by_population(articles)
    
    return all_flavors

def create_document_with_flavors(flavors, hypothesis, query, total_articles):
    """Crea un documento DOCX con flavors y sus resúmenes"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    title = doc.add_heading('Flavors - Párrafos Temáticos con Resúmenes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Estrategia de búsqueda: {query[:200]}...')
    doc.add_paragraph(f'Hipótesis: "{hypothesis[:200]}..."')
    doc.add_paragraph(f'Total de artículos analizados: {total_articles}')
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph()
    
    for category_name, flavor_list in flavors.items():
        if not flavor_list:
            continue
        
        category_title = category_name.replace('_', ' ').title()
        doc.add_heading(category_title, level=1)
        
        for flavor in flavor_list:
            doc.add_heading(flavor['name'], level=2)
            
            # Añadir resumen generado
            doc.add_paragraph(flavor['summary'], style='Normal')
            doc.add_paragraph()
            
            doc.add_paragraph('Referencias:', style='List Bullet')
            for i, article in enumerate(flavor.get('representative_articles', flavor['articles'][:5]), 1):
                citation = generate_citation(article, i)
                doc.add_paragraph(citation, style='List Bullet')
            
            doc.add_paragraph()
    
    return doc

def main():
    st.title("🧠 PubMed AI Analyzer - Generador de Flavors Avanzado")
    
    st.info("⚡ Versión mejorada con BioBERT NER, HDBSCAN clustering, y filtrado por embeddings")
    
    st.markdown("""
    ### Genera párrafos temáticos (flavors) con resúmenes y referencias
    
    **Mejoras técnicas implementadas:**
    - 🧬 **BioBERT NER**: Extracción semántica de entidades médicas
    - 🔍 **Filtrado por embeddings**: Relevancia combinada (búsqueda + hipótesis)
    - 📊 **HDBSCAN clustering**: Clusters naturales sin número fijo
    - 📝 **Resúmenes generativos**: Párrafos descriptivos por flavor
    - 📚 **Referencias Vancouver**: Formato académico estándar
    """)
    
    col_a, col_b = st.columns(2)
    with col_a:
        if AI_EMBEDDINGS_AVAILABLE:
            st.success("✅ BioBERT disponible (NER + embeddings)")
        else:
            st.warning("⚠️ BioBERT en modo básico")
    with col_b:
        st.info("📄 Salida: DOCX con resúmenes + referencias")
    
    st.markdown("---")
    st.markdown("### 📝 Configuración")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        query = st.text_area(
            "**Estrategia de búsqueda en PubMed:**",
            value="(((\"Ebolavirus\"[Mesh]) OR (EBOV)) AND ((\"Computational Biology\"[Mesh]) OR (\"Machine Learning\"[Mesh]) OR (bioinformatics))) AND ((\"Viral Proteins\"[Mesh]) OR (VP40 OR VP35 OR GP)) NOT (review[pt])",
            height=100
        )
    
    with col2:
        max_articles = st.slider(
            "Máx. artículos:",
            min_value=50,
            max_value=500,
            value=200,
            step=50
        )
    
    hypothesis = st.text_area(
        "**📌 Conjetura/hipótesis (lenguaje natural):**",
        value="Machine learning-based bioinformatic characterization of Ebola virus proteins (GP, VP40, VP35, VP30, VP24, and NP) enables the accurate prediction of protein function, identification of conserved domains, and discovery of potential druggable sites for antiviral therapeutics.",
        height=100
    )
    
    generate_button = st.button("🚀 GENERAR FLAVORS AVANZADOS", type="primary", use_container_width=True)
    
    if 'docx_generated' not in st.session_state:
        st.session_state.docx_generated = False
        st.session_state.docx_data = None
        st.session_state.total_articles = 0
        st.session_state.n_flavors = 0
    
    if generate_button:
        if not query.strip():
            st.warning("⚠️ Por favor, introduce una estrategia de búsqueda")
        elif not hypothesis.strip():
            st.warning("⚠️ Por favor, introduce tu conjetura/hipótesis")
        else:
            start_time = time.time()
            
            with st.spinner("🔍 Buscando artículos en PubMed..."):
                id_list, total_count = search_pubmed(query.strip(), retmax=max_articles)
                
                if not id_list:
                    st.error("❌ No se encontraron artículos")
                    st.stop()
                
                st.info(f"📊 Se encontraron {total_count} artículos. Procesando {len(id_list)}...")
                articles = fetch_articles_details(id_list, max_articles=max_articles)
            
            if not articles:
                st.error("❌ No se pudo procesar ningún artículo")
                st.stop()
            
            st.success(f"✅ Procesados {len(articles)} artículos")
            
            with st.spinner("🧠 Calculando relevancia con búsqueda e hipótesis (BioBERT)..."):
                articles = calculate_relevance_to_search_and_hypothesis(articles, query, hypothesis)
                
                # Filtrado adaptativo
                high_relevance = [a for a in articles if a.get('relevance_score', 0) > 0.55]
                if len(high_relevance) >= 10:
                    articles = high_relevance
                    st.info(f"📌 {len(articles)} artículos con alta relevancia (>0.55)")
                else:
                    medium_relevance = [a for a in articles if a.get('relevance_score', 0) > 0.4]
                    if len(medium_relevance) >= 5:
                        articles = medium_relevance
                        st.info(f"📌 {len(articles)} artículos con relevancia moderada (>0.4)")
                    else:
                        st.info(f"📌 Manteniendo {len(articles)} artículos (relevancia >0.3)")
                        articles = [a for a in articles if a.get('relevance_score', 0) > 0.3]
            
            with st.spinner("🔍 Descubriendo flavors con HDBSCAN..."):
                flavors = generate_all_flavors(articles, hypothesis)
            
            total_flavors = sum(len(flavor_list) for flavor_list in flavors.values())
            st.success(f"✅ Generados {total_flavors} flavors")
            
            with st.spinner("📄 Creando documento con resúmenes..."):
                doc = create_document_with_flavors(flavors, hypothesis, query, len(articles))
                
                docx_bytes = BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)
                
                st.session_state.docx_data = docx_bytes
                st.session_state.docx_generated = True
                st.session_state.total_articles = len(articles)
                st.session_state.n_flavors = total_flavors
            
            elapsed_time = time.time() - start_time
            st.info(f"⏱️ Tiempo: {elapsed_time/60:.1f} minutos")
    
    if st.session_state.docx_generated and st.session_state.docx_data is not None:
        st.markdown("---")
        st.markdown("### 📥 Documento Generado")
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Artículos relevantes", st.session_state.total_articles)
        with col2:
            st.metric("Flavors generados", st.session_state.n_flavors)
        
        st.download_button(
            label="📥 DESCARGAR FLAVORS CON RESÚMENES",
            data=st.session_state.docx_data,
            file_name=f"flavors_avanzados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        with st.expander("📖 Mejoras técnicas implementadas"):
            st.markdown("""
            **Mejoras de AI aplicadas:**
            
            1. **NER con BioBERT**: Extracción semántica de entidades médicas (proteínas, fármacos, outcomes)
            2. **Filtrado por embeddings**: Relevancia combinada (30% búsqueda + 70% hipótesis)
            3. **HDBSCAN clustering**: Clusters naturales sin número fijo de grupos
            4. **Resúmenes generativos**: Párrafos descriptivos para cada flavor
            5. **Detección de negación mejorada**: Con intensificadores y contexto extendido
            
            **Estructura del documento:**
            - Cada flavor tiene un nombre descriptivo (ej: "VP40: efficacy, mortality")
            - Un resumen generado automáticamente
            - Lista de referencias en formato Vancouver
            """)
    
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: gray; font-size: 0.8em;'>
            🧠 PubMed AI Analyzer - Generador de Flavors Avanzado v7.0<br>
            BioBERT NER | HDBSCAN Clustering | Embeddings Filtering | Resúmenes Generativos
        </div>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
