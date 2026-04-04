import streamlit as st
import pandas as pd
from datetime import datetime
import time
import requests
from xml.etree import ElementTree
import re
from collections import Counter
import numpy as np
import math
import string
import os
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import hashlib
import paramiko
import json
from typing import List, Dict, Optional
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
warnings.filterwarnings('ignore')

# ============================================================================
# CONFIGURACIÓN DE EMBEDDINGS CON FALLBACK ROBUSTO
# ============================================================================

AI_EMBEDDINGS_AVAILABLE = False
BIOMED_EMBEDDER = None
FALLBACK_EMBEDDER = None
USE_FALLBACK = False

try:
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import HDBSCAN, KMeans, AgglomerativeClustering
    from sklearn.decomposition import PCA, TruncatedSVD
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.preprocessing import normalize
    
    try:
        from sentence_transformers import SentenceTransformer
        import torch
        
        with st.spinner("🔄 Loading embeddings model..."):
            try:
                BIOMED_EMBEDDER = SentenceTransformer(
                    'pritamdeka/S-Biomed-Roberta-snli-multinli-stsb',
                    device='cpu'
                )
                AI_EMBEDDINGS_AVAILABLE = True
                USE_FALLBACK = False
                st.success("✅ BioBERT embeddings available")
            except:
                try:
                    FALLBACK_EMBEDDER = SentenceTransformer(
                        'all-MiniLM-L6-v2',
                        device='cpu'
                    )
                    AI_EMBEDDINGS_AVAILABLE = True
                    USE_FALLBACK = True
                    st.warning("⚠️ BioBERT unavailable, using SBERT")
                except:
                    AI_EMBEDDINGS_AVAILABLE = False
    except ImportError:
        AI_EMBEDDINGS_AVAILABLE = False
        
except Exception as e:
    print(f"⚠️ Some ML libraries not available: {e}")
    AI_EMBEDDINGS_AVAILABLE = False

TFIDF_AVAILABLE = True

st.set_page_config(
    page_title="PubMed AI Analyzer - Advanced Flavor Generator",
    page_icon="🧠",
    layout="centered"
)

# ============================================================================
# FUNCIONES DE CORREO ELECTRÓNICO
# ============================================================================

def send_status_email(to_email, session_id, block_num, total_blocks, articles_in_block, total_articles, status):
    """Enviar correo de estado del procesamiento de un bloque"""
    try:
        smtp_server = st.secrets["smtp_server"]
        smtp_port = st.secrets["smtp_port"]
        email_user = st.secrets["email_user"]
        email_password = st.secrets["email_password"]
        
        if not email_user or not email_password:
            st.warning("⚠️ Credenciales de email no configuradas")
            return False
        
        msg = MIMEMultipart()
        msg['From'] = email_user
        msg['To'] = to_email
        msg['Subject'] = f"[PubMed Analyzer] Progreso - Bloque {block_num}/{total_blocks}"
        
        if status == "started_block":
            body = f"""
Estimado usuario,

El bloque {block_num} de {total_blocks} ha comenzado su procesamiento.

Detalles:
- ID de sesión: {session_id[:30]}...
- Bloque: {block_num}/{total_blocks}
- Artículos en este bloque: {articles_in_block}

Recibirá una notificación cuando este bloque se complete.

Saludos cordiales,
PubMed AI Analyzer
"""
        elif status == "completed" and block_num > 0:
            body = f"""
Estimado usuario,

El bloque {block_num} de {total_blocks} ha sido procesado exitosamente.

Detalles:
- ID de sesión: {session_id[:30]}...
- Bloque: {block_num}/{total_blocks}
- Artículos en este bloque: {articles_in_block}
- Total procesado hasta ahora: {total_articles}

Saludos cordiales,
PubMed AI Analyzer
"""
        elif status == "started":
            body = f"""
Estimado usuario,

El procesamiento de su búsqueda en PubMed ha comenzado.

Detalles:
- ID de sesión: {session_id[:30]}...
- Total de bloques: {total_blocks}
- Artículos totales: {total_articles}

Recibirá notificaciones al iniciar y completar cada bloque.

Saludos cordiales,
PubMed AI Analyzer
"""
        else:
            body = f"""
Estimado usuario,

El bloque {block_num} de {total_blocks} ha sido procesado.

Saludos cordiales,
PubMed AI Analyzer
"""
        
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(email_user, email_password)
        server.send_message(msg)
        server.quit()
        
        st.success(f"📧 Correo enviado: {'Inicio' if status == 'started_block' else 'Completado'} Bloque {block_num}")
        return True
    except Exception as e:
        st.error(f"❌ Error en correo: {e}")
        return False


def send_final_email(to_email, session_id, total_articles, relevance_threshold, docx_bytes, session_id_short):
    """Enviar correo final con el archivo DOCX adjunto"""
    try:
        smtp_server = st.secrets["smtp_server"]
        smtp_port = st.secrets["smtp_port"]
        email_user = st.secrets["email_user"]
        email_password = st.secrets["email_password"]
        
        if not email_user or not email_password:
            st.warning("⚠️ Credenciales de email no configuradas")
            return False
        
        msg = MIMEMultipart()
        msg['From'] = email_user
        msg['To'] = to_email
        msg['Subject'] = f"[PubMed Analyzer] ¡Proceso completado! - {session_id_short}..."
        
        body = f"""
Estimado usuario,

¡El procesamiento de su búsqueda en PubMed ha finalizado exitosamente!

Detalles finales:
- ID de sesión: {session_id}
- Artículos procesados: {total_articles}
- Threshold de relevancia: {relevance_threshold}
- Fecha de finalización: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Se adjunta el archivo flavors_{session_id_short}.docx con los flavors generados.

Saludos cordiales,
PubMed AI Analyzer
"""
        
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        if hasattr(docx_bytes, 'seek'):
            docx_bytes.seek(0)
            docx_data = docx_bytes.getvalue()
        else:
            docx_data = docx_bytes
        
        docx_part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        docx_part.set_payload(docx_data)
        encoders.encode_base64(docx_part)
        docx_part.add_header('Content-Disposition', f'attachment; filename="flavors_{session_id_short}.docx"')
        msg.attach(docx_part)
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(email_user, email_password)
        server.send_message(msg)
        server.quit()
        
        st.success(f"📧 Correo final enviado a {to_email} con DOCX adjunto")
        return True
    except Exception as e:
        st.error(f"❌ Error al enviar correo final: {e}")
        return False


# ============================================================================
# CLASES PARA ALMACENAMIENTO REMOTO CON CSV
# ============================================================================

class RemoteCSVStorage:
    def __init__(self, host, port, username, password, remote_dir):
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.remote_dir = remote_dir
        self.client = None
        self.sftp = None
        self._connect()
    
    def _connect(self):
        try:
            self.client = paramiko.SSHClient()
            self.client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            self.client.connect(
                hostname=self.host,
                port=self.port,
                username=self.username,
                password=self.password,
                timeout=30
            )
            self.sftp = self.client.open_sftp()
            return True
        except Exception as e:
            st.error(f"❌ Error de conexión remota: {e}")
            return False
    
    def _ensure_dir(self):
        try:
            self.sftp.stat(self.remote_dir)
        except FileNotFoundError:
            try:
                self.sftp.mkdir(self.remote_dir)
            except:
                pass
    
    def _get_file_path(self, filename):
        return f"{self.remote_dir}/{filename}"
    
    def read_csv(self, filename):
        try:
            if not self.sftp:
                return None
            self._ensure_dir()
            file_path = self._get_file_path(filename)
            with self.sftp.open(file_path, 'r') as f:
                content = f.read().decode('utf-8')
                if content.strip():
                    return pd.read_csv(StringIO(content))
                return None
        except FileNotFoundError:
            return None
        except Exception as e:
            st.warning(f"⚠️ Error leyendo {filename}: {e}")
            return None
    
    def write_csv(self, filename, df):
        try:
            if not self.sftp:
                return False
            self._ensure_dir()
            file_path = self._get_file_path(filename)
            csv_content = df.to_csv(index=False, encoding='utf-8-sig')
            with self.sftp.open(file_path, 'w') as f:
                f.write(csv_content.encode('utf-8-sig'))
            return True
        except Exception as e:
            st.error(f"❌ Error escribiendo {filename}: {e}")
            return False
    
    def append_to_csv(self, filename, new_data):
        try:
            existing_df = self.read_csv(filename)
            if existing_df is not None and not existing_df.empty:
                df = pd.concat([existing_df, new_data], ignore_index=True)
                if 'pmid' in df.columns and 'session_id' in df.columns:
                    df = df.drop_duplicates(subset=['pmid', 'session_id'], keep='last')
            else:
                df = new_data
            return self.write_csv(filename, df)
        except Exception as e:
            st.error(f"❌ Error append a {filename}: {e}")
            return False
    
    def close(self):
        if hasattr(self, 'sftp') and self.sftp:
            try:
                self.sftp.close()
            except:
                pass
        if hasattr(self, 'client') and self.client:
            try:
                self.client.close()
            except:
                pass


class UserSessionManager:
    def __init__(self, remote_storage, login):
        self.remote = remote_storage
        self.login = login
        self.articles_file = f"{login}_articles.csv"
        self.sessions_file = f"{login}_sessions.csv"
        self.checkpoints_file = f"{login}_checkpoints.csv"
    
    def create_session(self, query: str, hypothesis: str, relevance_threshold: float, email: str) -> str:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        query_hash = hashlib.md5(query.encode()).hexdigest()[:8]
        session_id = f"{timestamp}_{query_hash}"
        
        session_data = pd.DataFrame([{
            'session_id': session_id,
            'login': self.login,
            'query': query[:500],
            'hypothesis': hypothesis[:500],
            'query_hash': query_hash,
            'relevance_threshold': relevance_threshold,
            'total_found': 0,
            'total_processed': 0,
            'start_time': datetime.now().isoformat(),
            'end_time': '',
            'status': 'running',
            'flavors_generated': False,
            'user_email': email
        }])
        self.remote.append_to_csv(self.sessions_file, session_data)
        return session_id
    
    def mark_flavors_generated(self, session_id: str):
        df = self.remote.read_csv(self.sessions_file)
        if df is None or df.empty:
            return
        df = df.copy()
        mask = df['session_id'] == session_id
        df.loc[mask, 'flavors_generated'] = True
        self.remote.write_csv(self.sessions_file, df)
    
    def get_user_email(self, session_id: str) -> Optional[str]:
        df = self.remote.read_csv(self.sessions_file)
        if df is None or df.empty:
            return None
        session_data = df[df['session_id'] == session_id]
        if session_data.empty:
            return None
        return session_data.iloc[0].get('user_email', None)
    
    def save_articles_batch(self, articles: List[Dict], session_id: str):
        if not articles:
            return
        
        records = []
        for i, article in enumerate(articles):
            if i % 100 == 0 and i > 0:
                st.write(f"   Procesando artículo {i+1}/{len(articles)}...")
            
            pmid = article.get('pmid', '')
            if not pmid:
                continue
            
            record = {
                'pmid': str(pmid),
                'session_id': str(session_id),
                'login': str(self.login),
                'title': str(article.get('title', ''))[:1000],
                'authors': str(article.get('authors', ''))[:500],
                'journal': str(article.get('journal', ''))[:200],
                'pubdate': str(article.get('pubdate', '')),
                'doi': str(article.get('doi', ''))[:100],
                'abstract': str(article.get('abstract', ''))[:5000],
                'study_types': str(article.get('study_types', 'Not specified')),
                'quality_score': float(article.get('quality_score', 0) or 0),
                'evidence_strength': str(article.get('evidence_strength', 'No data')),
                'top_keywords': str(article.get('top_keywords', '')),
                'population': str(article.get('population', 'Not specified')),
                'outcomes': ','.join(str(o) for o in (article.get('all_outcomes', []) or []) if o),
                'numeric_results': str(article.get('numeric_results_str', '')),
                'relevance_score': float(article.get('relevance_score', 0.5) or 0.5),
                'search_relevance': float(article.get('search_relevance', 0.5) or 0.5),
                'hypothesis_relevance': float(article.get('hypothesis_relevance', 0.5) or 0.5),
                'embedding_used': 'TF-IDF',
                'processed_date': datetime.now().isoformat(),
                'block_number': int(article.get('block_number', 0) or 0)
            }
            records.append(record)
        
        if not records:
            return
        
        df_new = pd.DataFrame(records)
        existing_df = self.remote.read_csv(self.articles_file)
        
        if existing_df is not None and not existing_df.empty:
            df = pd.concat([existing_df, df_new], ignore_index=True)
            df = df.drop_duplicates(subset=['pmid', 'session_id'], keep='last')
        else:
            df = df_new
        
        self.remote.write_csv(self.articles_file, df)
        st.success(f"✅ Guardados {len(df_new)} artículos")
    
    def save_checkpoint(self, session_id: str, block_num: int, batch_size: int,
                        start_idx: int, end_idx: int, articles_processed: int):
        checkpoint_data = pd.DataFrame([{
            'session_id': str(session_id),
            'login': str(self.login),
            'batch_number': int(block_num),
            'batch_size': int(batch_size),
            'start_idx': int(start_idx),
            'end_idx': int(end_idx),
            'articles_processed': int(articles_processed),
            'checkpoint_time': str(datetime.now().isoformat()),
            'status': 'completed'
        }])
        self.remote.append_to_csv(self.checkpoints_file, checkpoint_data)
    
    def get_completed_blocks(self, session_id: str) -> set:
        df = self.remote.read_csv(self.checkpoints_file)
        if df is None or df.empty:
            return set()
        session_checkpoints = df[df['session_id'] == session_id]
        return set(int(cp['batch_number']) for _, cp in session_checkpoints.iterrows())
    
    def get_processed_pmids(self, session_id: str) -> set:
        df = self.remote.read_csv(self.articles_file)
        if df is None or df.empty:
            return set()
        session_articles = df[df['session_id'] == session_id]
        return set(session_articles['pmid'].tolist())
    
    def update_session(self, session_id: str, status: str, total_found: int = None, total_processed: int = None):
        df = self.remote.read_csv(self.sessions_file)
        if df is None or df.empty:
            return
        df = df.copy()
        mask = df['session_id'] == session_id
        if total_found is not None:
            df.loc[mask, 'total_found'] = int(total_found)
        if total_processed is not None:
            df.loc[mask, 'total_processed'] = int(total_processed)
        df.loc[mask, 'status'] = str(status)
        df.loc[mask, 'end_time'] = str(datetime.now().isoformat())
        self.remote.write_csv(self.sessions_file, df)
    
    def get_user_sessions(self) -> pd.DataFrame:
        df = self.remote.read_csv(self.sessions_file)
        if df is None or df.empty:
            return pd.DataFrame()
        return df[df['login'] == self.login].sort_values('start_time', ascending=False)
    
    def get_session_articles(self, session_id: str) -> pd.DataFrame:
        df = self.remote.read_csv(self.articles_file)
        if df is None or df.empty:
            return pd.DataFrame()
        return df[df['session_id'] == session_id]
    
    def get_session_stats(self, session_id: str) -> Dict:
        articles_df = self.get_session_articles(session_id)
        if articles_df.empty:
            return {'total_articles': 0, 'avg_quality': 0, 'strong_evidence': 0, 'avg_relevance': 0}
        
        avg_quality = articles_df['quality_score'].mean() if 'quality_score' in articles_df.columns else 0
        avg_relevance = articles_df['relevance_score'].mean() if 'relevance_score' in articles_df.columns else 0
        
        return {
            'total_articles': len(articles_df),
            'avg_quality': avg_quality,
            'strong_evidence': 0,
            'avg_relevance': avg_relevance
        }
    
    def export_session_to_csv(self, session_id: str) -> bytes:
        articles_df = self.get_session_articles(session_id)
        if articles_df.empty:
            return None
        export_columns = ['pmid', 'title', 'authors', 'journal', 'pubdate', 'doi',
                          'study_types', 'quality_score', 'evidence_strength',
                          'relevance_score', 'outcomes', 'processed_date', 'block_number']
        available_cols = [col for col in export_columns if col in articles_df.columns]
        export_df = articles_df[available_cols]
        csv_buffer = StringIO()
        export_df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
        csv_buffer.seek(0)
        return csv_buffer.getvalue().encode('utf-8-sig')
    
    def get_session_info(self, session_id: str) -> Optional[Dict]:
        df = self.remote.read_csv(self.sessions_file)
        if df is None or df.empty:
            return None
        session_data = df[df['session_id'] == session_id]
        if session_data.empty:
            return None
        return session_data.iloc[0].to_dict()
    
    def delete_session(self, session_id: str) -> bool:
        try:
            articles_df = self.remote.read_csv(self.articles_file)
            if articles_df is not None and not articles_df.empty:
                articles_df = articles_df[articles_df['session_id'] != session_id]
                self.remote.write_csv(self.articles_file, articles_df)
            
            sessions_df = self.remote.read_csv(self.sessions_file)
            if sessions_df is not None and not sessions_df.empty:
                sessions_df = sessions_df[sessions_df['session_id'] != session_id]
                self.remote.write_csv(self.sessions_file, sessions_df)
            
            checkpoints_df = self.remote.read_csv(self.checkpoints_file)
            if checkpoints_df is not None and not checkpoints_df.empty:
                checkpoints_df = checkpoints_df[checkpoints_df['session_id'] != session_id]
                self.remote.write_csv(self.checkpoints_file, checkpoints_df)
            
            return True
        except Exception as e:
            st.error(f"Error eliminando sesión: {e}")
            return False
    
    def get_articles_by_block(self, session_id: str, block_number: int) -> pd.DataFrame:
        articles_df = self.get_session_articles(session_id)
        if articles_df.empty:
            return pd.DataFrame()
        if 'block_number' in articles_df.columns:
            return articles_df[articles_df['block_number'] == block_number]
        else:
            BLOCK_SIZE = 500  # MODIFICADO: 500 artículos por bloque
            start_idx = (block_number - 1) * BLOCK_SIZE
            end_idx = block_number * BLOCK_SIZE
            return articles_df.iloc[start_idx:end_idx] if len(articles_df) > start_idx else pd.DataFrame()
    
    def has_flavors_generated(self, session_id: str) -> bool:
        df = self.remote.read_csv(self.sessions_file)
        if df is None or df.empty:
            return False
        session_data = df[df['session_id'] == session_id]
        if session_data.empty:
            return False
        return session_data.iloc[0].get('flavors_generated', False)


# ============================================================================
# FUNCIONES DE UTILIDAD PARA PubMed
# ============================================================================

def make_request_with_retry(url, params, max_retries=5, initial_delay=5):
    delay = initial_delay
    for attempt in range(max_retries):
        try:
            response = requests.get(url, params=params, timeout=30)
            if response.status_code == 429:
                wait_time = delay * (2 ** attempt)
                st.warning(f"⚠️ Rate limit. Waiting {wait_time}s...")
                time.sleep(wait_time)
                continue
            response.raise_for_status()
            return response
        except:
            if attempt == max_retries - 1:
                raise
            time.sleep(delay)
    return None


def search_pubmed_complete(query, max_articles=1000):
    """Search articles - limitado a max_articles (1000 = 2 bloques de 500)"""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    search_url = f"{base_url}esearch.fcgi"
    
    try:
        st.info("🔍 Buscando artículos en PubMed...")
        
        # Primero obtener el total
        params = {"db": "pubmed", "term": query, "retmax": 1, "retmode": "xml"}
        response = make_request_with_retry(search_url, params)
        if response is None:
            return [], 0
        
        root = ElementTree.fromstring(response.content)
        total_count = int(root.find(".//Count").text) if root.find(".//Count") is not None else 0
        st.info(f"📊 PubMed encontró {total_count} artículos")
        
        if total_count == 0:
            return [], 0
        
        # Limitar a max_articles
        articles_to_fetch = min(total_count, max_articles)
        st.info(f"📊 Se procesarán {articles_to_fetch} artículos ({articles_to_fetch//500} bloques de 500)")
        
        # Obtener todos los IDs
        all_ids = []
        retstart = 0
        batch_size = 10000
        
        progress_bar = st.progress(0)
        
        while retstart < articles_to_fetch:
            remaining = articles_to_fetch - retstart
            current_batch = min(batch_size, remaining)
            
            fetch_params = {
                "db": "pubmed",
                "term": query,
                "retstart": retstart,
                "retmax": current_batch,
                "retmode": "xml"
            }
            
            try:
                fetch_response = make_request_with_retry(search_url, fetch_params)
                if fetch_response is None:
                    break
                
                fetch_root = ElementTree.fromstring(fetch_response.content)
                batch_ids = [id_elem.text for id_elem in fetch_root.findall(".//Id")]
                
                if not batch_ids:
                    break
                
                all_ids.extend(batch_ids)
                st.info(f"   ✅ Descargados {len(batch_ids)} IDs (total: {len(all_ids)})")
                
                progress_bar.progress(min(len(all_ids) / articles_to_fetch, 1.0))
                retstart += current_batch
                time.sleep(0.5)
                
            except Exception as e:
                st.error(f"Error: {e}")
                break
        
        progress_bar.empty()
        st.success(f"✅ Recuperados {len(all_ids)} IDs")
        
        if len(all_ids) < 500:
            st.error(f"❌ No hay suficientes artículos (mínimo 500, hay {len(all_ids)})")
            return [], 0
        
        return all_ids, len(all_ids)
        
    except Exception as e:
        st.error(f"Error: {e}")
        return [], 0


@st.cache_data(ttl=3600, show_spinner=False)
def get_abstract(pmid):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    fetch_url = f"{base_url}efetch.fcgi"
    params = {"db": "pubmed", "id": pmid, "retmode": "xml", "rettype": "abstract"}
    try:
        response = make_request_with_retry(fetch_url, params)
        if response is None:
            return None
        root = ElementTree.fromstring(response.content)
        abstract_texts = []
        for abstract in root.findall(".//AbstractText"):
            if abstract.text:
                abstract_texts.append(abstract.text)
        return " ".join(abstract_texts) if abstract_texts else None
    except:
        return None


def preprocess_text(text):
    if not text:
        return ""
    text = text.lower()
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\d+', '', text)
    return ' '.join(text.split())


# ============================================================================
# EXTRACCIÓN DE TÉRMINOS
# ============================================================================

def extract_key_terms_from_query(query):
    terms = set()
    mesh_pattern = r'"([^"]+)"\[Mesh\]'
    for term in re.findall(mesh_pattern, query, re.IGNORECASE):
        terms.add(term.lower())
    tiab_pattern = r'"([^"]+)"\[tiab\]'
    for term in re.findall(tiab_pattern, query, re.IGNORECASE):
        terms.add(term.lower())
    quoted_pattern = r'"([^"]+)"'
    for term in re.findall(quoted_pattern, query):
        if len(term) > 3 and not term.lower().endswith('mesh') and not term.lower().endswith('tiab'):
            terms.add(term.lower())
    return list(terms)


def extract_key_terms_from_hypothesis(hypothesis):
    if not hypothesis:
        return []
    terms = set()
    hypothesis_lower = hypothesis.lower()
    words = hypothesis_lower.split()
    for i in range(len(words)-1):
        if len(words[i]) > 4 and len(words[i+1]) > 3:
            terms.add(f"{words[i]} {words[i+1]}")
        if len(words[i]) > 3:
            terms.add(words[i])
    pattern = r'\b([a-z]+(?:\s+[a-z]+){1,3})\b'
    for match in re.findall(pattern, hypothesis_lower):
        if len(match.split()) >= 2:
            terms.add(match)
    return list(terms)


# ============================================================================
# ENTITY EXTRACTION
# ============================================================================

def get_embedder():
    if BIOMED_EMBEDDER is not None:
        return BIOMED_EMBEDDER
    elif FALLBACK_EMBEDDER is not None:
        return FALLBACK_EMBEDDER
    return None


def extract_entities_with_embeddings(text, dynamic_terms):
    embedder = get_embedder()
    if not embedder or not text or not dynamic_terms or not AI_EMBEDDINGS_AVAILABLE:
        return []
    try:
        text_embedding = embedder.encode([text[:2000]])[0]
        entity_embeddings = embedder.encode(dynamic_terms)
        similarities = cosine_similarity([text_embedding], entity_embeddings)[0]
        threshold = 0.60 if USE_FALLBACK else 0.65
        entities = []
        for i, sim in enumerate(similarities):
            if sim > threshold:
                entities.append((dynamic_terms[i], 'semantic', sim))
        return entities
    except:
        return []


def extract_entities_with_regex(text, dynamic_terms):
    if not text or not dynamic_terms:
        return []
    text_lower = text.lower()
    entities = []
    for term in dynamic_terms:
        if term.lower() in text_lower:
            entities.append((term, 'regex_match', 1.0))
    return entities


def extract_medical_entities_enhanced(text, query_terms, hypothesis_terms):
    if not text:
        return []
    all_terms = list(set(query_terms + hypothesis_terms))
    embedding_entities = extract_entities_with_embeddings(text, all_terms)
    regex_entities = extract_entities_with_regex(text, all_terms)
    all_entities = []
    seen = set()
    for entity, etype, score in embedding_entities:
        if entity.lower() not in seen:
            all_entities.append((entity, etype, score))
            seen.add(entity.lower())
    for entity, etype, score in regex_entities:
        if entity.lower() not in seen:
            all_entities.append((entity, etype, score))
            seen.add(entity.lower())
    return all_entities


def extract_numeric_results(text):
    if not text:
        return []
    results = []
    hr_pattern = r'(?:HR|hazard ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(hr_pattern, text, re.IGNORECASE):
        results.append({'type': 'HR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    rr_pattern = r'(?:RR|risk ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(rr_pattern, text, re.IGNORECASE):
        results.append({'type': 'RR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    or_pattern = r'(?:OR|odds ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(or_pattern, text, re.IGNORECASE):
        results.append({'type': 'OR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    p_pattern = r'[Pp]\s*[=<]\s*([0-9]+\.[0-9]+)'
    for match in re.findall(p_pattern, text):
        results.append({'type': 'p-value', 'value': float(match)})
    return results


def extract_all_outcomes(text):
    if not text:
        return []
    text_lower = text.lower()
    outcomes = set()
    indicators = ['mortality', 'death', 'survival', 'rupture', 'bleeding', 'hemorrhage', 
                  'stroke', 'reinfarction', 'complication', 'risk', 'rate', 'incidence', 
                  'outcome', 'endpoint', 'recovery', 'improvement', 'efficacy', 'safety']
    for indicator in indicators:
        if indicator in text_lower:
            outcomes.add(indicator)
    return list(outcomes)


def analyze_sentiment_by_outcome(text):
    if not text:
        return {}
    text_lower = text.lower()
    outcomes_sentiment = {}
    all_outcomes = extract_all_outcomes(text)
    negation_words = ['no', 'not', 'without', 'lack', 'failed', 'non-significant']
    intensifiers = ['significantly', 'markedly', 'dramatically', 'strongly']
    for outcome in all_outcomes:
        outcome_index = text_lower.find(outcome)
        if outcome_index == -1:
            continue
        start = max(0, outcome_index - 300)
        end = min(len(text_lower), outcome_index + 300)
        context = text_lower[start:end]
        positive_terms = ['reduced', 'decreased', 'lower', 'improved', 'better', 'benefit', 
                         'protective', 'effective', 'efficacious', 'superior', 'advantage']
        negative_terms = ['increased', 'higher', 'worse', 'elevated', 'risk', 'adverse', 
                         'harm', 'detrimental', 'inferior', 'disadvantage']
        sentiment_score = 0
        for term in positive_terms:
            if term in context:
                term_pos = context.find(term)
                prev_words = context[max(0, term_pos-50):term_pos]
                if any(neg in prev_words for neg in negation_words):
                    sentiment_score -= 1
                else:
                    if any(inten in prev_words for inten in intensifiers):
                        sentiment_score += 2
                    else:
                        sentiment_score += 1
        for term in negative_terms:
            if term in context:
                term_pos = context.find(term)
                prev_words = context[max(0, term_pos-50):term_pos]
                if any(neg in prev_words for neg in negation_words):
                    sentiment_score += 1
                else:
                    if any(inten in prev_words for inten in intensifiers):
                        sentiment_score -= 2
                    else:
                        sentiment_score -= 1
        if sentiment_score > 0:
            outcomes_sentiment[outcome] = 'BENEFIT'
        elif sentiment_score < 0:
            outcomes_sentiment[outcome] = 'RISK'
        else:
            outcomes_sentiment[outcome] = 'NO EFFECT'
    return outcomes_sentiment


def enhanced_quality_scoring(study_types, full_text):
    score = 0
    factors = []
    text_lower = full_text.lower() if full_text else ""
    study_type_weights = {'Meta-analysis': 2.5, 'RCT': 2.0, 'Cohort': 1.6, 'Case-control': 1.4, 'Observational': 1.2}
    for study_type in study_types:
        if study_type in study_type_weights:
            score += study_type_weights[study_type] * 20
            factors.append(f"Study type: {study_type}")
    sample_patterns = [r'n\s*[=:]\s*(\d+)', r'sample size\s*[=:]\s*(\d+)']
    max_sample = 0
    for pattern in sample_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                max_sample = max(max_sample, int(match))
            except:
                pass
    if max_sample > 0:
        if max_sample > 1000:
            score += 20
            factors.append(f"Large sample: n={max_sample}")
        elif max_sample > 500:
            score += 15
        elif max_sample > 200:
            score += 10
        elif max_sample > 100:
            score += 5
    if any(term in text_lower for term in ['multicenter', 'multi-center']):
        score += 10
        factors.append("Multicenter study")
    if any(term in text_lower for term in ['double-blind', 'double blind', 'masked']):
        score += 15
        factors.append("Blinded")
    return min(score, 100), factors


def get_effect_direction(result_value, ci_lower=None, ci_upper=None):
    try:
        value = float(result_value)
        if value < 1:
            if ci_upper and float(ci_upper) < 1:
                return "PROTECTIVE"
            return "PROTECTIVE (trend)"
        elif value > 1:
            if ci_lower and float(ci_lower) > 1:
                return "HARMFUL"
            return "HARMFUL (trend)"
        return "NO EFFECT"
    except:
        return "NOT DETERMINED"


def calculate_evidence_strength(statistical_results, quality_score):
    if not statistical_results:
        return "No data", 0, {}
    strength_score = 0
    directions = {}
    for i, result in enumerate(statistical_results):
        if result['type'] in ['HR', 'RR', 'OR']:
            effect = result['value']
            if effect < 0.5 or effect > 2.0:
                strength_score += 30
            elif effect < 0.7 or effect > 1.5:
                strength_score += 20
            elif effect != 1.0:
                strength_score += 10
            directions[f"result_{i}"] = get_effect_direction(effect, result.get('ci_lower'), result.get('ci_upper'))
        if result['type'] == 'p-value':
            if result['value'] < 0.001:
                strength_score += 25
            elif result['value'] < 0.01:
                strength_score += 20
            elif result['value'] < 0.05:
                strength_score += 15
    strength_score = strength_score * (quality_score / 50)
    if strength_score >= 50:
        strength = "STRONG EVIDENCE"
    elif strength_score >= 30:
        strength = "MODERATE EVIDENCE"
    elif strength_score >= 15:
        strength = "WEAK EVIDENCE"
    else:
        strength = "INSUFFICIENT EVIDENCE"
    return strength, min(strength_score, 100), directions


def analyze_article_with_ai(title, abstract, query_terms, hypothesis_terms):
    if not title and not abstract:
        return {}
    full_text = f"{title} {abstract if abstract else ''}"
    processed_text = preprocess_text(full_text)
    entities = extract_medical_entities_enhanced(full_text, query_terms, hypothesis_terms)
    
    study_types = []
    study_keywords = {
        'RCT': ['randomized', 'rct', 'randomised', 'trial', 'double-blind'],
        'Cohort': ['cohort', 'follow-up', 'longitudinal', 'prospective', 'retrospective'],
        'Meta-analysis': ['meta-analysis', 'meta analysis', 'systematic review'],
        'Observational': ['observational', 'registry', 'real-world'],
        'Case-control': ['case-control', 'matched']
    }
    for study_type, keywords in study_keywords.items():
        if any(keyword in processed_text for keyword in keywords):
            study_types.append(study_type)
    
    words = processed_text.split()
    word_freq = Counter(words)
    top_keywords = [word for word, _ in word_freq.most_common(20) if len(word) > 3][:10]
    
    population_patterns = {
        'adults': r'\badults?\b|\bpatients?\b',
        'elderly': r'\belderly\b|\baged\b|\bolder\b',
        'women': r'\bwomen\b|\bfemale\b',
        'men': r'\bmen\b|\bmale\b',
        'diabetes': r'\bdiabetes\b|\bdiabetic\b',
        'hypertension': r'\bhypertension\b|\bhypertensive\b'
    }
    population = [pop for pop, pattern in population_patterns.items() if re.search(pattern, processed_text)]
    
    numeric_results = extract_numeric_results(full_text)
    all_outcomes = extract_all_outcomes(full_text)
    quality_score, quality_factors = enhanced_quality_scoring(study_types, full_text)
    evidence_strength, evidence_score, directions = calculate_evidence_strength(numeric_results, quality_score)
    sentiment = analyze_sentiment_by_outcome(full_text)
    numeric_results_str = ' | '.join([f"{r['type']}={r['value']}" for r in numeric_results[:3]])
    
    return {
        'entities': entities,
        'study_types': ', '.join(study_types) if study_types else 'Not specified',
        'top_keywords': ', '.join(top_keywords),
        'population': ', '.join(population) if population else 'Not specified',
        'quality_score': quality_score,
        'quality_factors': ' | '.join(quality_factors) if quality_factors else 'No factors',
        'numeric_results': numeric_results,
        'numeric_results_str': numeric_results_str,
        'outcomes_analysis': sentiment,
        'all_outcomes': all_outcomes,
        'evidence_strength': evidence_strength,
        'evidence_score': evidence_score,
        'effect_directions': directions,
        'abstract': abstract if abstract else ''
    }


def extract_article_info(doc_sum):
    article = {}
    article["pmid"] = doc_sum.find("Id").text if doc_sum.find("Id") is not None else "N/A"
    title_item = doc_sum.find(".//Item[@Name='Title']")
    article["title"] = title_item.text if title_item is not None else "No title"
    author_items = doc_sum.findall(".//Item[@Name='Author']")
    authors = [author.text for author in author_items if author.text]
    article["authors"] = ", ".join(authors[:5]) + (" et al." if len(authors) > 5 else "")
    source_item = doc_sum.find(".//Item[@Name='Source']")
    article["journal"] = source_item.text if source_item is not None else "N/A"
    pubdate_item = doc_sum.find(".//Item[@Name='PubDate']")
    article["pubdate"] = pubdate_item.text if pubdate_item is not None else "N/A"
    doi_item = doc_sum.find(".//Item[@Name='DOI']")
    article["doi"] = doi_item.text if doi_item is not None else "N/A"
    return article


def fetch_articles_details(id_list, query_terms, hypothesis_terms, block_number=0):
    if not id_list:
        return []
    total_to_process = len(id_list)
    batch_size = 30
    num_batches = math.ceil(total_to_process / batch_size)
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    articles = []
    progress_bar = st.progress(0)
    
    for batch_num in range(num_batches):
        start_idx = batch_num * batch_size
        end_idx = min((batch_num + 1) * batch_size, total_to_process)
        batch_ids = id_list[start_idx:end_idx]
        
        st.write(f"   📦 Lote {batch_num + 1}/{num_batches} ({len(batch_ids)} artículos)...")
        
        summary_params = {"db": "pubmed", "id": ",".join(batch_ids), "retmode": "xml"}
        
        for retry in range(3):
            try:
                summary_response = make_request_with_retry(f"{base_url}esummary.fcgi", summary_params)
                if summary_response is None:
                    if retry < 2:
                        time.sleep(5)
                        continue
                    else:
                        raise Exception("Max retries exceeded")
                
                summary_root = ElementTree.fromstring(summary_response.content)
                
                for j, doc_sum in enumerate(summary_root.findall(".//DocSum")):
                    overall_idx = start_idx + j
                    progress_bar.progress((overall_idx + 1) / total_to_process)
                    
                    article = extract_article_info(doc_sum)
                    time.sleep(0.5)
                    abstract = get_abstract(article["pmid"])
                    article["abstract"] = abstract if abstract else "Not available"
                    ai_analysis = analyze_article_with_ai(article["title"], abstract, query_terms, hypothesis_terms)
                    article.update(ai_analysis)
                    article["block_number"] = block_number
                    articles.append(article)
                
                break
                
            except Exception as e:
                if retry == 2:
                    st.warning(f"Error en lote {batch_num + 1}: {str(e)[:50]}")
                else:
                    time.sleep(5)
        
        time.sleep(1)
    
    progress_bar.empty()
    return articles


def process_articles_in_independent_blocks(id_list, query_terms, hypothesis_terms, 
                                           session_manager, session_id, query, hypothesis, user_email):
    BLOCK_SIZE = 500  # MODIFICADO: 500 artículos por bloque
    total_to_process = len(id_list)
    total_blocks = (total_to_process + BLOCK_SIZE - 1) // BLOCK_SIZE
    
    st.info(f"📦 Total de artículos a procesar: {total_to_process}")
    st.info(f"📦 Se procesarán en {total_blocks} bloques de {BLOCK_SIZE} artículos cada uno")
    
    # Email de inicio del proceso general
    if user_email:
        st.info(f"📧 Enviando correo de inicio a {user_email}...")
        send_status_email(user_email, session_id, 0, total_blocks, 0, total_to_process, "started")
    
    all_articles = []
    
    if session_manager:
        completed_blocks = session_manager.get_completed_blocks(session_id)
        if completed_blocks:
            st.info(f"🔄 Bloques ya completados: {sorted(completed_blocks)}")
        existing_df = session_manager.get_session_articles(session_id)
        if not existing_df.empty:
            all_articles = existing_df.to_dict('records')
            processed_pmids = set(existing_df['pmid'].tolist())
            st.info(f"📊 {len(processed_pmids)} artículos ya procesados")
        else:
            processed_pmids = set()
    else:
        completed_blocks = set()
        processed_pmids = set()
    
    for block_num in range(1, total_blocks + 1):
        if block_num in completed_blocks:
            st.info(f"⏭️ Bloque {block_num}/{total_blocks} ya completado. Saltando...")
            continue
        
        start_idx = (block_num - 1) * BLOCK_SIZE
        end_idx = min(block_num * BLOCK_SIZE, total_to_process)
        block_ids = id_list[start_idx:end_idx]
        new_ids = [pid for pid in block_ids if pid not in processed_pmids]
        
        if not new_ids:
            st.info(f"⏭️ Bloque {block_num} sin nuevos IDs. Marcando como completado...")
            if session_manager:
                session_manager.save_checkpoint(session_id, block_num, BLOCK_SIZE, start_idx, end_idx, len(all_articles))
            continue
        
        # EMAIL: INICIO DEL BLOQUE
        if user_email:
            st.info(f"📧 Enviando correo de INICIO del bloque {block_num}...")
            send_status_email(user_email, session_id, block_num, total_blocks, len(new_ids), 0, "started_block")
        
        st.markdown(f"## 📦 PROCESANDO BLOQUE {block_num}/{total_blocks}")
        st.info(f"Artículos {start_idx+1}-{end_idx} ({len(new_ids)} nuevos)")
        
        block_articles = fetch_articles_details(new_ids, query_terms, hypothesis_terms, block_num)
        
        if block_articles:
            st.info(f"🧠 Calculando relevancia para {len(block_articles)} artículos...")
            block_articles = calculate_relevance_to_search_and_hypothesis(block_articles, query, hypothesis)
        
        if session_manager and block_articles:
            session_manager.save_articles_batch(block_articles, session_id)
            for article in block_articles:
                processed_pmids.add(article.get('pmid', ''))
            session_manager.save_checkpoint(session_id, block_num, BLOCK_SIZE, start_idx, end_idx, len(all_articles) + len(block_articles))
            
            st.success(f"✅ BLOQUE {block_num} COMPLETADO. {len(block_articles)} artículos guardados.")
            
            # EMAIL: COMPLETACIÓN DEL BLOQUE
            if user_email:
                st.info(f"📧 Enviando correo de COMPLETACIÓN del bloque {block_num}...")
                send_status_email(user_email, session_id, block_num, total_blocks, len(block_articles), len(all_articles) + len(block_articles), "completed")
        
        all_articles.extend(block_articles)
        
        if block_num < total_blocks:
            st.info("⏳ Esperando 5 segundos antes del siguiente bloque...")
            time.sleep(5)
    
    if session_manager:
        session_manager.update_session(session_id, 'completed', total_found=total_to_process, total_processed=len(all_articles))
    
    st.success(f"🎉 PROCESO COMPLETADO. Total: {len(all_articles)} artículos en {total_blocks} bloques")
    return all_articles, total_blocks


def calculate_relevance_to_search_and_hypothesis(articles, query, hypothesis):
    embedder = get_embedder()
    if not embedder or not AI_EMBEDDINGS_AVAILABLE:
        for article in articles:
            article['relevance_score'] = 0.5
            article['search_relevance'] = 0.5
            article['hypothesis_relevance'] = 0.5
        return articles
    try:
        query_text = query[:500] if len(query) > 500 else query
        query_embedding = embedder.encode([query_text])[0]
        hypothesis_embedding = embedder.encode([hypothesis])[0]
        for article in articles:
            text = f"{article.get('title', '')} {article.get('abstract', '')}"
            if text and len(text) > 50:
                article_embedding = embedder.encode([text[:1000]])[0]
                search_sim = cosine_similarity([query_embedding], [article_embedding])[0][0]
                hypothesis_sim = cosine_similarity([hypothesis_embedding], [article_embedding])[0][0]
                article['relevance_score'] = (search_sim * 0.3) + (hypothesis_sim * 0.7)
                article['search_relevance'] = float(search_sim)
                article['hypothesis_relevance'] = float(hypothesis_sim)
            else:
                article['relevance_score'] = 0.0
                article['search_relevance'] = 0.0
                article['hypothesis_relevance'] = 0.0
        articles.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
    except Exception as e:
        st.warning(f"Error en relevancia: {e}")
        for article in articles:
            article['relevance_score'] = 0.5
    return articles


def filter_articles_by_relevance(articles, relevance_threshold):
    filtered = [a for a in articles if a.get('relevance_score', 0) >= relevance_threshold]
    if articles:
        st.write(f"**📊 Filtro:**")
        st.write(f"   - Threshold: {relevance_threshold}")
        st.write(f"   - Antes: {len(articles)}")
        st.write(f"   - Después: {len(filtered)} ({len(filtered)/len(articles)*100:.1f}%)")
    return filtered


# ============================================================================
# FUNCIONES DE CLUSTERING
# ============================================================================

def extract_topic_keywords_tfidf(articles, n_keywords=5):
    if not articles or len(articles) < 2:
        return []
    texts = []
    for a in articles[:20]:
        title = a.get('title', '')
        outcomes = ' '.join(a.get('all_outcomes', []))
        text = f"{title} {outcomes}"
        if text:
            texts.append(text)
    if len(texts) < 2:
        return []
    try:
        vectorizer = TfidfVectorizer(max_features=100, stop_words='english')
        tfidf_matrix = vectorizer.fit_transform(texts)
        feature_names = vectorizer.get_feature_names_out()
        scores = np.array(tfidf_matrix.sum(axis=0)).flatten()
        top_indices = scores.argsort()[-n_keywords:][::-1]
        return [feature_names[i] for i in top_indices if scores[i] > 0][:n_keywords]
    except:
        return []


def determine_flavor_aspect_and_difference(articles, flavor_name, query_terms, hypothesis_terms):
    if not articles:
        return "Clinical analysis", "Integrative approach"
    all_outcomes = []
    all_methods = []
    all_populations = []
    all_entities = []
    for a in articles:
        if not a:
            continue
        outcomes = a.get('all_outcomes', [])
        if outcomes:
            all_outcomes.extend(outcomes)
        text = f"{a.get('title', '')} {a.get('abstract', '')}".lower()
        if 'cohort' in text:
            all_methods.append('cohort study')
        if 'registry' in text:
            all_methods.append('registry analysis')
        if 'meta-analysis' in text:
            all_methods.append('meta-analysis')
        pop = a.get('population', '')
        if pop and pop != 'Not specified':
            all_populations.extend(pop.split(', '))
        entities = a.get('entities', [])
        if entities:
            for entity, etype, score in entities:
                all_entities.append(entity)
    outcome_counts = Counter(all_outcomes)
    method_counts = Counter(all_methods)
    entity_counts = Counter(all_entities)
    
    aspect = ""
    if entity_counts:
        aspect = f"Clinical Analysis of {entity_counts.most_common(1)[0][0].title()}"
    elif outcome_counts:
        aspect = f"Evaluation of {outcome_counts.most_common(1)[0][0].capitalize()} Outcomes"
    else:
        aspect = "Clinical Outcomes Analysis"
    
    difference = ""
    if method_counts.get('meta-analysis', 0) > len(articles) * 0.3:
        difference = "Synthesizes evidence through meta-analytic methods"
    elif method_counts.get('registry analysis', 0) > len(articles) * 0.4:
        difference = "Leverages large-scale registry data"
    elif method_counts.get('cohort study', 0) > len(articles) * 0.5:
        difference = "Longitudinal cohort design"
    else:
        difference = "Integrates diverse clinical studies"
    
    return aspect, difference


def get_text_embeddings(texts):
    embedder = get_embedder()
    if embedder and AI_EMBEDDINGS_AVAILABLE:
        try:
            return embedder.encode(texts)
        except:
            pass
    try:
        vectorizer = TfidfVectorizer(max_features=500, stop_words='english')
        tfidf_matrix = vectorizer.fit_transform(texts)
        svd = TruncatedSVD(n_components=min(50, tfidf_matrix.shape[1] - 1), random_state=42)
        return svd.fit_transform(tfidf_matrix)
    except:
        pass
    return None


def discover_flavors_by_embeddings_hdbscan(articles, query_terms, hypothesis_terms):
    if len(articles) < 3:
        return []
    try:
        texts = [f"{a.get('title', '')} {' '.join(a.get('all_outcomes', []))}" for a in articles if a]
        embeddings = get_text_embeddings(texts)
        if embeddings is None:
            return []
        n_components = min(50, embeddings.shape[1], len(embeddings) - 1)
        if n_components > 1:
            pca = PCA(n_components=n_components)
            embeddings_reduced = pca.fit_transform(embeddings)
        else:
            embeddings_reduced = embeddings
        clusterer = HDBSCAN(min_cluster_size=3, min_samples=2, metric='euclidean')
        cluster_labels = clusterer.fit_predict(embeddings_reduced)
        flavors = []
        for cluster_id in set(cluster_labels):
            if cluster_id == -1:
                continue
            cluster_indices = [i for i, label in enumerate(cluster_labels) if label == cluster_id]
            if len(cluster_indices) < 3:
                continue
            cluster_articles = [articles[i] for i in cluster_indices]
            name = generate_descriptive_name(cluster_articles, query_terms, hypothesis_terms)
            representative = sorted(cluster_articles, key=lambda x: x.get('quality_score', 0), reverse=True)[:5]
            flavors.append({
                'type': 'semantic',
                'id': f"semantic_cluster_{cluster_id}",
                'name': name,
                'articles': cluster_articles,
                'n_articles': len(cluster_articles),
                'representative_articles': representative
            })
        return flavors
    except:
        return []


def discover_flavors_by_outcomes(articles, query_terms, hypothesis_terms):
    if len(articles) < 3:
        return []
    try:
        all_outcomes = set()
        for article in articles:
            if article:
                outcomes = article.get('all_outcomes', [])
                if outcomes:
                    all_outcomes.update(outcomes)
        if len(all_outcomes) < 2:
            return []
        all_outcomes_list = list(all_outcomes)
        outcome_matrix = []
        for article in articles:
            if article:
                outcomes = set(article.get('all_outcomes', []))
                row = [1 if o in outcomes else 0 for o in all_outcomes_list]
                outcome_matrix.append(row)
        n_clusters = min(max(3, len(articles) // 15), 4)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clusters = kmeans.fit_predict(outcome_matrix)
        flavors = []
        for cluster_id in range(n_clusters):
            cluster_indices = [i for i, c in enumerate(clusters) if c == cluster_id]
            if len(cluster_indices) < 3:
                continue
            cluster_articles = [articles[i] for i in cluster_indices]
            outcome_counts = Counter()
            for article in cluster_articles:
                if article:
                    outcomes = article.get('all_outcomes', [])
                    if outcomes:
                        outcome_counts.update(outcomes)
            top_outcomes = [o for o, _ in outcome_counts.most_common(3)]
            name = f"Studies on {', '.join(top_outcomes)}"
            representative = sorted(cluster_articles, key=lambda x: x.get('quality_score', 0) if x else 0, reverse=True)[:5]
            flavors.append({
                'type': 'outcome_based',
                'id': f"outcome_cluster_{cluster_id}",
                'name': name,
                'articles': cluster_articles,
                'n_articles': len(cluster_articles),
                'representative_articles': representative
            })
        return flavors
    except:
        return []


def merge_small_clusters(flavors, target_count=4):
    if not flavors:
        return []
    if len(flavors) <= target_count:
        return flavors
    flavors_sorted = sorted(flavors, key=lambda x: x['n_articles'], reverse=True)
    merged_flavors = flavors_sorted[:target_count-1]
    remaining_articles = []
    for flavor in flavors_sorted[target_count-1:]:
        remaining_articles.extend(flavor['articles'])
    if remaining_articles:
        name = "Additional Clinical Studies"
        all_outcomes = []
        for article in remaining_articles:
            if article:
                outcomes = article.get('all_outcomes', [])
                if outcomes:
                    all_outcomes.extend(outcomes)
        outcome_counts = Counter(all_outcomes)
        if outcome_counts:
            top_outcomes = [o for o, _ in outcome_counts.most_common(2)]
            if top_outcomes:
                name += f": {', '.join(top_outcomes)}"
        representative = sorted(remaining_articles, key=lambda x: x.get('quality_score', 0) if x else 0, reverse=True)[:5]
        merged_flavors.append({
            'type': 'merged',
            'id': 'merged_cluster',
            'name': name,
            'articles': remaining_articles,
            'n_articles': len(remaining_articles),
            'representative_articles': representative
        })
    return merged_flavors


def assign_articles_to_best_flavor(flavors):
    if not flavors:
        return []
    article_to_flavor = {}
    article_scores = {}
    for flavor in flavors:
        flavor_id = flavor['id']
        flavor_articles = flavor['articles']
        for article in flavor_articles:
            if not article:
                continue
            article_id = article.get('pmid', id(article))
            score = article.get('quality_score', 50) / 100.0
            flavor_name_lower = flavor['name'].lower()
            title_lower = article.get('title', '').lower()
            if any(keyword in title_lower for keyword in flavor_name_lower.split()[:3]):
                score += 0.2
            if article_id not in article_scores or score > article_scores[article_id]:
                article_scores[article_id] = score
                article_to_flavor[article_id] = flavor_id
    flavor_article_map = {flavor['id']: [] for flavor in flavors}
    for article in [a for flavor in flavors for a in flavor['articles'] if a]:
        article_id = article.get('pmid', id(article))
        if article_id in article_to_flavor:
            assigned_flavor = article_to_flavor[article_id]
            if article not in flavor_article_map[assigned_flavor]:
                flavor_article_map[assigned_flavor].append(article)
    for flavor in flavors:
        flavor['articles'] = flavor_article_map.get(flavor['id'], [])
        flavor['n_articles'] = len(flavor['articles'])
        flavor['representative_articles'] = sorted(
            flavor['articles'], 
            key=lambda x: x.get('quality_score', 0) if x else 0, 
            reverse=True
        )[:5]
    return [f for f in flavors if f['n_articles'] >= 2]


def generate_descriptive_name(articles, query_terms, hypothesis_terms):
    if not articles:
        return "Clinical Studies"
    all_keywords = []
    all_outcomes = []
    all_entities = []
    for a in articles:
        if not a:
            continue
        keywords = a.get('top_keywords', '')
        if keywords:
            all_keywords.extend(keywords.split(', '))
        outcomes = a.get('all_outcomes', [])
        if outcomes:
            all_outcomes.extend(outcomes)
        entities = a.get('entities', [])
        if entities:
            for entity, etype, score in entities:
                all_entities.append(entity)
    outcome_counts = Counter(all_outcomes)
    entity_counts = Counter(all_entities)
    if entity_counts:
        top_entity = entity_counts.most_common(1)[0][0]
        if len(top_entity) > 3:
            return top_entity.title()
    if outcome_counts:
        top_outcomes = [o for o, _ in outcome_counts.most_common(2)]
        return f"Studies on {', '.join(top_outcomes)}"
    if all_keywords:
        return all_keywords[0].title()
    return "Clinical Studies"


def generate_all_flavors(articles, query_terms, hypothesis_terms):
    if not articles:
        return {}
    articles = [a for a in articles if a is not None]
    if not articles:
        return {}
    all_flavors = []
    if len(articles) >= 5:
        semantic_flavors = discover_flavors_by_embeddings_hdbscan(articles, query_terms, hypothesis_terms)
        all_flavors.extend(semantic_flavors)
    if len(articles) >= 5:
        outcome_flavors = discover_flavors_by_outcomes(articles, query_terms, hypothesis_terms)
        all_flavors.extend(outcome_flavors)
    if not all_flavors:
        name = generate_descriptive_name(articles, query_terms, hypothesis_terms)
        all_flavors = [{
            'type': 'default',
            'id': 'default_cluster',
            'name': name,
            'articles': articles,
            'n_articles': len(articles),
            'representative_articles': articles[:5]
        }]
    all_flavors = assign_articles_to_best_flavor(all_flavors)
    merged_flavors = merge_small_clusters(all_flavors, target_count=4)
    flavors_by_category = {
        'semantic_clusters': [],
        'outcome_clusters': [],
        'merged_clusters': []
    }
    for flavor in merged_flavors:
        category = flavor.get('type', 'semantic')
        if category == 'semantic':
            flavors_by_category['semantic_clusters'].append(flavor)
        elif category == 'outcome_based':
            flavors_by_category['outcome_clusters'].append(flavor)
        else:
            flavors_by_category['merged_clusters'].append(flavor)
    return flavors_by_category


def generate_citation_text(article, index):
    if not article:
        return f"{index}. [No article data]"
    authors = str(article.get('authors', 'Author'))
    year = str(article.get('pubdate', 'n.d.'))[:4] if article.get('pubdate') else 'n.d.'
    title = str(article.get('title', 'No title'))
    journal = str(article.get('journal', 'Journal'))
    pmid = str(article.get('pmid', ''))
    citation = f"{index}. {authors}. {title}. {journal}. {year}"
    if pmid and pmid != 'N/A':
        citation += f"; PMID: {pmid}"
    return citation


def generate_flavor_summary_with_citations(articles, flavor_name, section, query_terms, hypothesis_terms):
    if not articles:
        return "No articles available for this flavor.", []
    articles = [a for a in articles if a is not None]
    if not articles:
        return "No articles available for this flavor.", []
    aspect, difference = determine_flavor_aspect_and_difference(articles, flavor_name, query_terms, hypothesis_terms)
    main_outcomes = []
    key_findings = []
    study_types = []
    key_articles = []
    for idx, a in enumerate(articles[:10], 1):
        if not a:
            continue
        outcomes = a.get('all_outcomes', []) or []
        study_type = a.get('study_types', '') or ''
        numeric = a.get('numeric_results_str', '') or ''
        authors = a.get('authors', 'Author') or 'Author'
        authors = authors.split(',')[0] if authors else 'Author'
        main_outcomes.extend(outcomes[:2])
        if numeric:
            key_findings.append((idx, authors, numeric))
        if study_type and study_type != 'Not specified':
            study_types.append(study_type)
        key_articles.append((idx, authors))
    outcome_counts = Counter(main_outcomes)
    study_type_counts = Counter(study_types)
    top_outcomes = [o for o, _ in outcome_counts.most_common(3)]
    top_study_types = [st for st, _ in study_type_counts.most_common(2)]
    header = f"### {flavor_name}\n\n**🎯 Aspect:** {aspect}\n\n**🔍 Difference:** {difference}\n\n---\n\n"
    if section == 'introduction':
        summary = header
        summary += f"This flavor groups {len(articles)} clinical studies "
        if top_outcomes:
            summary += f"investigating {', '.join(top_outcomes[:2])} "
        summary += f"in the context of the research topic. "
        if top_study_types:
            summary += f"Methodological approaches include {', '.join(top_study_types[:2])}, "
        summary += f"providing comprehensive insights into clinical outcomes. "
        if top_outcomes:
            summary += f"Primary outcomes examined include {', '.join(top_outcomes[:3])}. "
        if key_articles:
            refs = [f"{auth} ({idx})" for idx, auth in key_articles[:3] if auth]
            if refs:
                summary += f"Key contributions from {', '.join(refs)} have advanced understanding. "
        if key_findings:
            find_refs = [f"{auth} ({idx})" for idx, auth, _ in key_findings[:3] if auth]
            if find_refs:
                summary += f"Notable findings from {', '.join(find_refs)} include {key_findings[0][2] if key_findings else 'significant associations'}. "
        summary += f"The convergence of these {len(articles)} studies demonstrates clinical importance."
    else:
        summary = header
        summary += f"Our analysis aligns with findings from {len(articles)} clinical studies. "
        if top_study_types:
            summary += f"The methodological convergence across {', '.join(top_study_types[:2])} has enabled identification of consistent risk factors. "
        if key_articles:
            refs = [f"{auth} ({idx})" for idx, auth in key_articles[:3] if auth]
            if refs:
                summary += f"Specifically, {', '.join(refs)} reported consistent results. "
        if key_findings:
            find_refs = [f"{auth} ({idx})" for idx, auth, _ in key_findings[:3] if auth]
            if find_refs:
                summary += f"Quantitative findings from {', '.join(find_refs)} support our conclusions. "
        summary += f"Heterogeneity in approaches across these studies highlights the need for standardized protocols."
    citations = []
    for i, article in enumerate(articles[:15], 1):
        if article:
            citations.append(generate_citation_text(article, i))
    return summary, citations


def create_document_with_flavors(flavors, hypothesis, query, total_articles, relevance_threshold, query_terms, hypothesis_terms):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    title = doc.add_heading('Thematic Flavors for Scientific Article', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Search strategy: {str(query)[:200]}...')
    doc.add_paragraph(f'Hypothesis: "{str(hypothesis)[:200]}..."')
    doc.add_paragraph(f'Total articles analyzed: {total_articles}')
    doc.add_paragraph(f'Relevance threshold applied: {relevance_threshold}')
    doc.add_paragraph(f'Generation date: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'Embedding model: TF-IDF + LSA (fallback)')
    doc.add_paragraph()
    doc.add_heading('FLAVORS FOR INTRODUCTION', level=1)
    doc.add_paragraph('The following paragraphs are designed for the Introduction section.')
    doc.add_paragraph()
    for category_name, flavor_list in flavors.items():
        if not flavor_list:
            continue
        category_title = category_name.replace('_', ' ').title()
        doc.add_heading(category_title, level=2)
        for flavor in flavor_list:
            intro_summary, intro_citations = generate_flavor_summary_with_citations(
                flavor.get('representative_articles', flavor['articles'][:10]), 
                flavor['name'], 
                'introduction',
                query_terms,
                hypothesis_terms
            )
            doc.add_paragraph(intro_summary)
            doc.add_paragraph()
            doc.add_paragraph('References:', style='List Bullet')
            for citation in intro_citations:
                doc.add_paragraph(citation, style='List Bullet')
            doc.add_paragraph()
    doc.add_page_break()
    doc.add_heading('FLAVORS FOR DISCUSSION', level=1)
    doc.add_paragraph('The following paragraphs are designed for the Discussion section.')
    doc.add_paragraph()
    for category_name, flavor_list in flavors.items():
        if not flavor_list:
            continue
        category_title = category_name.replace('_', ' ').title()
        doc.add_heading(category_title, level=2)
        for flavor in flavor_list:
            disc_summary, disc_citations = generate_flavor_summary_with_citations(
                flavor.get('representative_articles', flavor['articles'][:10]), 
                flavor['name'], 
                'discussion',
                query_terms,
                hypothesis_terms
            )
            doc.add_paragraph(disc_summary)
            doc.add_paragraph()
            doc.add_paragraph('References:', style='List Bullet')
            for citation in disc_citations:
                doc.add_paragraph(citation, style='List Bullet')
            doc.add_paragraph()
    return doc


def export_articles_to_csv(articles):
    if not articles:
        return None
    data = []
    for article in articles:
        if not article:
            continue
        abstract = article.get('abstract', '') or ''
        outcomes = article.get('all_outcomes', []) or []
        row = {
            'PMID': str(article.get('pmid', '')),
            'Title': str(article.get('title', '')),
            'Authors': str(article.get('authors', '')),
            'Journal': str(article.get('journal', '')),
            'Publication Date': str(article.get('pubdate', '')),
            'DOI': str(article.get('doi', '')),
            'Study Types': str(article.get('study_types', '')),
            'Quality Score': float(article.get('quality_score', 0)),
            'Evidence Strength': str(article.get('evidence_strength', '')),
            'Top Keywords': str(article.get('top_keywords', '')),
            'Population': str(article.get('population', '')),
            'Numeric Results': str(article.get('numeric_results_str', '')),
            'Relevance Score': float(article.get('relevance_score', 0)),
            'Outcomes': ', '.join(str(o) for o in outcomes if o) if outcomes else '',
            'Block Number': int(article.get('block_number', 0)),
            'Abstract (first 500 chars)': str(abstract)[:500] if abstract else ''
        }
        data.append(row)
    if not data:
        return None
    df = pd.DataFrame(data)
    csv_buffer = StringIO()
    df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
    csv_buffer.seek(0)
    return csv_buffer.getvalue().encode('utf-8-sig')


def generate_flavors_from_saved_session_only(session_manager, session_id, relevance_threshold, block_number=None):
    if block_number is not None:
        articles_df = session_manager.get_articles_by_block(session_id, block_number)
        block_label = f"bloque {block_number}"
    else:
        articles_df = session_manager.get_session_articles(session_id)
        block_label = "sesión completa"
    if articles_df.empty:
        st.error(f"❌ No hay artículos en {block_label}")
        return None, None, None, None
    session_info = session_manager.get_session_info(session_id)
    if session_info is None:
        st.error("❌ No se pudo obtener información de la sesión")
        return None, None, None, None
    query = session_info.get('query', '')
    hypothesis = session_info.get('hypothesis', '')
    articles = articles_df.to_dict('records')
    articles = [a for a in articles if a is not None]
    st.info(f"✅ Cargados {len(articles)} artículos de {block_label}")
    threshold = relevance_threshold if relevance_threshold is not None else float(session_info.get('relevance_threshold', 0.35))
    filtered_articles = [a for a in articles if a.get('relevance_score', 0) is not None and a.get('relevance_score', 0) >= threshold]
    st.info(f"📊 Filtro ({threshold}): {len(filtered_articles)} de {len(articles)} artículos")
    if len(filtered_articles) < 5:
        st.error(f"❌ No hay suficientes artículos (mínimo 5, hay {len(filtered_articles)})")
        return None, None, None, None
    query_terms = extract_key_terms_from_query(query)
    hypothesis_terms = extract_key_terms_from_hypothesis(hypothesis)
    with st.spinner("🔍 Generando flavors..."):
        flavors = generate_all_flavors(filtered_articles, query_terms, hypothesis_terms)
    return flavors, filtered_articles, query, hypothesis


def generate_automatic_flavors(session_manager, session_id, relevance_threshold, user_email):
    st.markdown("---")
    st.markdown("## 🎨 GENERANDO FLAVORS AUTOMÁTICAMENTE...")
    st.info("Generando flavors desde los artículos procesados...")
    
    flavors, filtered_articles, query, hypothesis = generate_flavors_from_saved_session_only(
        session_manager, session_id, relevance_threshold, block_number=None
    )
    
    if flavors:
        session_manager.mark_flavors_generated(session_id)
        display_flavors_preview(flavors)
        
        query_terms = extract_key_terms_from_query(query)
        hypothesis_terms = extract_key_terms_from_hypothesis(hypothesis)
        
        doc = create_document_with_flavors(flavors, hypothesis, query, len(filtered_articles),
                                          relevance_threshold, query_terms, hypothesis_terms)
        
        docx_bytes_for_download = BytesIO()
        doc.save(docx_bytes_for_download)
        docx_bytes_for_download.seek(0)
        
        docx_bytes_for_email = BytesIO()
        doc.save(docx_bytes_for_email)
        docx_bytes_for_email.seek(0)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        st.markdown("---")
        st.markdown("## 📥 RESULTADO FINAL")
        st.success("✅ ¡Flavors generados exitosamente!")
        
        col_left, col_right = st.columns(2)
        with col_left:
            st.download_button("💾 DESCARGAR FLAVORS (DOCX)", data=docx_bytes_for_download,
                              file_name=f"flavors_{session_id[:20]}_{timestamp}.docx",
                              mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                              type="primary")
        with col_right:
            csv_data = export_articles_to_csv(filtered_articles)
            if csv_data:
                st.download_button("📊 DESCARGAR ARTÍCULOS (CSV)", data=csv_data,
                                  file_name=f"articles_{session_id[:20]}_{timestamp}.csv", mime="text/csv")
        
        if user_email:
            st.info(f"📧 Enviando archivo DOCX a {user_email}...")
            email_sent = send_final_email(user_email, session_id, len(filtered_articles), 
                                         relevance_threshold, docx_bytes_for_email, session_id[:20])
            if email_sent:
                st.success(f"✅ DOCX enviado a {user_email}")
            else:
                st.warning(f"⚠️ No se pudo enviar el correo")
        return True
    return False


def display_flavors_preview(flavors):
    st.markdown("---")
    st.markdown("## 🎨 Generated Flavors")
    total_flavors = 0
    for category_name, flavor_list in flavors.items():
        if not flavor_list:
            continue
        total_flavors += len(flavor_list)
        category_title = category_name.replace('_', ' ').title()
        st.markdown(f"### 📂 {category_title}")
        for i, flavor in enumerate(flavor_list, 1):
            with st.expander(f"🔹 Flavor {i}: {flavor.get('name', 'Unnamed')} ({flavor.get('n_articles', 0)} articles)", expanded=(i==1)):
                col1, col2 = st.columns(2)
                with col1:
                    rep_articles = flavor.get('representative_articles', []) or flavor.get('articles', [])[:5]
                    aspect, difference = determine_flavor_aspect_and_difference(
                        rep_articles[:5], flavor.get('name', 'Clinical Studies'), [], []
                    )
                    st.markdown(f"**🎯 Aspect:** {aspect}")
                    st.markdown(f"**🔍 Difference:** {difference}")
                with col2:
                    st.markdown(f"**📊 Articles:** {flavor.get('n_articles', 0)}")
                    st.markdown(f"**📚 Representative articles:**")
                    for j, article in enumerate(flavor.get('representative_articles', [])[:3], 1):
                        if article:
                            title = str(article.get('title', 'No title'))[:80]
                            st.markdown(f"   {j}. {title}...")
    st.info(f"✅ Total flavors: {total_flavors}")
    return total_flavors


def display_session_exporter(session_manager):
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 📤 Exportar sesión")
    user_sessions = session_manager.get_user_sessions()
    if user_sessions.empty:
        st.sidebar.info("No hay sesiones")
        return
    session_options = {}
    for _, session in user_sessions.iterrows():
        session_id = session['session_id']
        flavors_status = "✅" if session.get('flavors_generated', False) else "⏳"
        session_name = f"{flavors_status} {session_id[:25]}... - {session['start_time'][:16]} ({session.get('total_processed', 0)} artículos)"
        session_options[session_name] = session_id
    selected = st.sidebar.selectbox("Seleccionar sesión:", list(session_options.keys()))
    if selected:
        sid = session_options[selected]
        stats = session_manager.get_session_stats(sid)
        st.sidebar.markdown(f"📄 Artículos: {stats['total_articles']} | ⭐ Calidad: {stats['avg_quality']:.1f}")
        if st.sidebar.button("📥 Exportar CSV"):
            csv_data = session_manager.export_session_to_csv(sid)
            if csv_data:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                st.sidebar.download_button("💾 DESCARGAR CSV", data=csv_data,
                                          file_name=f"session_{sid[:8]}_{timestamp}.csv", mime="text/csv")


# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    st.title("🧠 PubMed AI Analyzer - Advanced Flavor Generator")
    
    if 'user_login' not in st.session_state:
        st.markdown("### 🔐 Identificación de usuario")
        col1, col2 = st.columns(2)
        with col1:
            login = st.text_input("Login:", placeholder="ejemplo: juan_perez")
        with col2:
            user_email = st.text_input("📧 Correo electrónico:", placeholder="usuario@dominio.com")
        
        if st.button("✅ Continuar", type="primary"):
            if login.strip() and user_email.strip() and '@' in user_email:
                st.session_state.user_login = login.strip().lower()
                st.session_state.user_email = user_email.strip()
                st.session_state.new_search_mode = True
                st.rerun()
            else:
                st.warning("⚠️ Complete ambos campos correctamente")
        return
    
    if 'new_search_mode' not in st.session_state:
        st.session_state.new_search_mode = True
    
    session_manager = None
    selected_session_id = None
    
    try:
        remote_storage = RemoteCSVStorage(
            host=st.secrets["remote_host"],
            port=st.secrets["remote_port"],
            username=st.secrets["remote_user"],
            password=st.secrets["remote_password"],
            remote_dir=st.secrets["remote_dir"]
        )
        session_manager = UserSessionManager(remote_storage, st.session_state.user_login)
        
        st.sidebar.success(f"👤 Usuario: {st.session_state.user_login}")
        st.sidebar.info(f"📧 Email: {st.session_state.user_email}")
        
        if st.sidebar.button("🔄 Cambiar usuario"):
            del st.session_state.user_login
            del st.session_state.user_email
            st.rerun()
        
        st.sidebar.markdown("---")
        user_sessions = session_manager.get_user_sessions()
        if not user_sessions.empty:
            session_options = {}
            for _, session in user_sessions.iterrows():
                sid = session['session_id']
                status = "✅" if session.get('flavors_generated', False) else "⏳"
                name = f"{status} {sid[:25]}... ({session.get('total_processed', 0)} artículos)"
                session_options[name] = sid
            selected = st.sidebar.selectbox("Sesiones guardadas:", ["[NUEVA BÚSQUEDA]"] + list(session_options.keys()))
            if selected == "[NUEVA BÚSQUEDA]":
                st.session_state.new_search_mode = True
                selected_session_id = None
            else:
                st.session_state.new_search_mode = False
                selected_session_id = session_options[selected]
        else:
            st.session_state.new_search_mode = True
        
        display_session_exporter(session_manager)
        
    except Exception as e:
        st.warning(f"⚠️ Error de conexión remota: {e}")
        session_manager = None
        st.session_state.new_search_mode = True
    
    st.info("⚡ 2 BLOQUES DE 500 ARTÍCULOS | Emails de inicio y fin por bloque | DOCX final por email")
    st.markdown("---")
    
    if selected_session_id and session_manager and not st.session_state.new_search_mode:
        session_info = session_manager.get_session_info(selected_session_id)
        if session_info:
            st.info(f"📌 Usando sesión: {selected_session_id[:25]}... | Artículos: {session_info.get('total_processed', 0)} | Flavors: {'✅' if session_manager.has_flavors_generated(selected_session_id) else '⏳'}")
            
            articles_df = session_manager.get_session_articles(selected_session_id)
            total_articles = len(articles_df)
            BLOCK_SIZE = 500
            num_blocks = (total_articles + BLOCK_SIZE - 1) // BLOCK_SIZE if total_articles > 0 else 0
            
            if num_blocks > 0:
                block_options = ["TODA LA SESIÓN"] + [f"Bloque {i+1}" for i in range(num_blocks)]
                selected_block = st.selectbox("Seleccionar:", block_options)
                threshold = st.slider("Relevance threshold:", 0.0, 0.9, 0.35, 0.05)
                
                if st.button("🎨 GENERAR FLAVORS", type="primary"):
                    if selected_block == "TODA LA SESIÓN":
                        flavors, articles, query, hypothesis = generate_flavors_from_saved_session_only(
                            session_manager, selected_session_id, threshold, None)
                        label = "sesion_completa"
                    else:
                        block_num = int(selected_block.split()[1])
                        flavors, articles, query, hypothesis = generate_flavors_from_saved_session_only(
                            session_manager, selected_session_id, threshold, block_num)
                        label = f"bloque_{block_num}"
                    
                    if flavors:
                        session_manager.mark_flavors_generated(selected_session_id)
                        display_flavors_preview(flavors)
                        qt = extract_key_terms_from_query(query)
                        ht = extract_key_terms_from_hypothesis(hypothesis)
                        doc = create_document_with_flavors(flavors, hypothesis, query, len(articles), threshold, qt, ht)
                        docx_bytes = BytesIO()
                        doc.save(docx_bytes)
                        docx_bytes.seek(0)
                        st.download_button("💾 DESCARGAR DOCX", data=docx_bytes, file_name=f"flavors_{label}.docx",
                                          mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        csv_data = export_articles_to_csv(articles)
                        if csv_data:
                            st.download_button("📊 DESCARGAR CSV", data=csv_data, file_name=f"articles_{label}.csv", mime="text/csv")
    else:
        st.info("🆕 **Nueva Búsqueda**")
        
        query = st.text_area("**PubMed search strategy:**",
            value="(\"myocardial infarction\"[Mesh] OR \"myocardial infarction\"[tiab]) AND (\"heart rupture\"[Mesh] OR \"cardiac rupture\"[tiab] OR \"ventricular septal rupture\"[tiab] OR \"free wall rupture\"[tiab] OR \"intramyocardial dissecting hematoma\"[tiab])",
            height=100)
        
        threshold = st.slider("Relevance threshold:", 0.0, 0.9, 0.35, 0.05)
        
        hypothesis = st.text_area("**📌 Hypothesis:**",
            value="Intramyocardial dissections occurring as a complication of myocardial infarction follow predictable anatomical pathways along established tissue planes, with distinct patterns based on timing of presentation and location within the ventricular wall.",
            height=100)
        
        auto_flavors = st.checkbox("🤖 Generar flavors automáticamente", value=True)
        
        if st.button("🚀 GENERATE (2 BLOQUES DE 500)", type="primary"):
            if not query.strip() or not hypothesis.strip():
                st.warning("⚠️ Complete los campos")
            else:
                start = time.time()
                qt = extract_key_terms_from_query(query)
                ht = extract_key_terms_from_hypothesis(hypothesis)
                st.info(f"📝 Términos: {len(qt)} de búsqueda, {len(ht)} de hipótesis")
                
                with st.spinner("🔍 Buscando artículos (máximo 1000 = 2 bloques de 500)..."):
                    id_list, total = search_pubmed_complete(query.strip(), max_articles=1000)
                    if not id_list:
                        st.error("❌ No se encontraron artículos")
                        st.stop()
                    st.info(f"📊 Se procesarán {len(id_list)} artículos (2 bloques de 500)")
                    
                    if session_manager:
                        sid = session_manager.create_session(query, hypothesis, threshold, st.session_state.user_email)
                        articles, blocks = process_articles_in_independent_blocks(
                            id_list, qt, ht, session_manager, sid, query, hypothesis, st.session_state.user_email)
                    else:
                        articles = fetch_articles_details(id_list, qt, ht)
                        blocks = 2
                
                if not articles:
                    st.error("❌ No se procesaron artículos")
                    st.stop()
                
                st.success(f"✅ Procesados {len(articles)} artículos de {blocks} bloques")
                
                with st.spinner("🧠 Calculando relevancia..."):
                    articles = calculate_relevance_to_search_and_hypothesis(articles, query, hypothesis)
                    filtered = filter_articles_by_relevance(articles, threshold)
                    articles = filtered
                
                if len(articles) < 5:
                    st.error(f"❌ Pocos artículos después del filtro ({len(articles)})")
                    st.stop()
                
                st.info(f"⏱️ Tiempo: {(time.time()-start)/60:.1f} minutos")
                
                if auto_flavors and session_manager and 'sid' in locals():
                    success = generate_automatic_flavors(session_manager, sid, threshold, st.session_state.user_email)
                    if success:
                        st.success("🎉 ¡Proceso completado! DOCX enviado a tu correo.")
    
    st.markdown("---")
    st.markdown("<div style='text-align: center; color: gray;'>PubMed AI Analyzer v31.0 | ✅ 2 BLOQUES DE 500 | Emails de inicio y fin por bloque</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
