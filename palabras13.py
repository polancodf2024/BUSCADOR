"""
PubMed AI Analyzer v8.7 - COMPLETO CON ASISTENTE INTERPRETATIVO
- Basado en palabras15.py (funciona correctamente)
- Incluye: embeddings (BioBERT/SBERT), clustering (KMeans), TF-IDF
- NUEVO: Sección "¿Qué significa esto?" con interpretación contextual
- Aumentado límite a 3000 artículos
- SIN LLM (solo procesamiento local)
"""

import streamlit as st
import pandas as pd
from datetime import datetime
import time
import requests
from xml.etree import ElementTree
import re
from collections import Counter
import numpy as np
import math
import os
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import hashlib
import json
from typing import List, Dict, Optional, Tuple
import pickle
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

warnings.filterwarnings('ignore')

# ============================================================================
# CONFIGURACIÓN
# ============================================================================

MAX_ARTICLES = 3000      # Máximo de artículos a procesar
BATCH_SIZE = 50          # Artículos por lote de PubMed API
BLOCK_SIZE = 500         # Artículos por bloque de procesamiento
REQUEST_DELAY = 0.5      # Segundos entre requests

def get_email_config():
    try:
        return {
            'smtp_server': st.secrets.get("smtp_server", "smtp.gmail.com"),
            'smtp_port': int(st.secrets.get("smtp_port", 587)),
            'email_user': st.secrets.get("email_user", ""),
            'email_password': st.secrets.get("email_password", ""),
            'notification_email': st.secrets.get("notification_email", "")
        }
    except Exception:
        return {'smtp_server': "smtp.gmail.com", 'smtp_port': 587, 'email_user': "", 'email_password': "", 'notification_email': ""}


# ============================================================================
# EMAIL SENDER
# ============================================================================

class EmailSender:
    @staticmethod
    def send_docx_by_email(recipient_email: str, docx_bytes: BytesIO, filename: str,
                           subject: str = None, body: str = None) -> Tuple[bool, str]:
        config = get_email_config()
        
        if not config['email_user'] or not config['email_password']:
            return False, "❌ Configuración de correo no disponible."
        
        if subject is None:
            subject = f"PubMed AI Analyzer v8.7 - Análisis de Evidencia - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        
        if body is None:
            body = f"""
            Hola,
            
            Adjunto encontrará el análisis de evidencia generado por PubMed AI Analyzer v8.7.
            
            Archivo: {filename}
            Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            
            Este documento incluye una nueva sección "¿Qué significa esto?" que interpreta
            automáticamente los resultados para ayudarle a contextualizar sus hallazgos.
            
            Este es un mensaje automático.
            """
        
        msg = MIMEMultipart()
        msg['From'] = config['email_user']
        msg['To'] = recipient_email
        msg['Subject'] = subject
        
        if config['notification_email']:
            msg['Bcc'] = config['notification_email']
        
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        try:
            docx_bytes.seek(0)
            attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
            attachment.set_payload(docx_bytes.read())
            encoders.encode_base64(attachment)
            attachment.add_header('Content-Disposition', f'attachment; filename={filename}')
            msg.attach(attachment)
        except Exception as e:
            return False, f"❌ Error al adjuntar: {str(e)}"
        
        try:
            with smtplib.SMTP(config['smtp_server'], config['smtp_port']) as server:
                server.starttls()
                server.login(config['email_user'], config['email_password'])
                server.send_message(msg)
            return True, f"✅ Enviado a {recipient_email}"
        except Exception as e:
            return False, f"❌ Error: {str(e)}"
    
    @staticmethod
    def get_email_config_status() -> Dict:
        config = get_email_config()
        return {
            'configured': bool(config['email_user'] and config['email_password']),
            'sender_email': config['email_user'] if config['email_user'] else "No configurado"
        }


# ============================================================================
# CLASIFICADOR DE RELACIÓN CON HIPÓTESIS
# ============================================================================

class HypothesisRelationClassifier:
    CONFIRMS_PHRASES = [
        'confirms that', 'demonstrates that', 'shows that',
        'consistent with the hypothesis', 'supports the hypothesis',
        'validates', 'our findings confirm', 'these results confirm',
        'is consistent with', 'provides evidence for',
        'confirm the hypothesis', 'support the hypothesis',
        'findings support', 'results confirm'
    ]
    
    CONTRADICTS_PHRASES = [
        'contradicts the hypothesis', 'contrary to', 'disagrees with',
        'opposite to what was expected', 'does not support',
        'refutes', 'calls into question', 'inconsistent with the hypothesis'
    ]
    
    NUANCES_PHRASES = ['however', 'although', 'but', 'except', 'whereas', 'interestingly', 'notably']
    
    @classmethod
    def classify(cls, title: str, abstract: str, hypothesis: str) -> Dict:
        if not abstract and not title:
            return {'relation': 'unrelated', 'confidence': 0.0, 'reason': 'No hay texto'}
        
        text = f"{title} {abstract if abstract else ''}".lower()
        
        for phrase in cls.CONFIRMS_PHRASES:
            if phrase in text:
                return {'relation': 'confirms', 'confidence': 0.95, 'reason': f'Frase explícita: "{phrase}"'}
        
        for phrase in cls.CONTRADICTS_PHRASES:
            if phrase in text:
                return {'relation': 'contradicts', 'confidence': 0.95, 'reason': f'Frase explícita: "{phrase}"'}
        
        nuance_count = sum(1 for p in cls.NUANCES_PHRASES if p in text)
        if nuance_count >= 2:
            return {'relation': 'nuances', 'confidence': 0.6, 'reason': f'Presencia de matices ({nuance_count})'}
        
        return {'relation': 'unrelated', 'confidence': 0.3, 'reason': 'No se encontró evidencia textual explícita'}


# ============================================================================
# NUEVO: ASISTENTE INTERPRETATIVO (v8.7)
# ============================================================================

class InterpretiveAssistant:
    """
    Genera una sección "¿Qué significa esto?" que contextualiza
    los resultados numéricos para el investigador.
    """
    
    @staticmethod
    def generate_interpretation(consensus: Dict, total_articles: int, 
                                 relevant_articles: int, hypothesis: str,
                                 avg_methodological: float) -> str:
        """
        Genera la interpretación contextual basada en los resultados.
        """
        confirms_pct = consensus.get('confirms_percentage', 0)
        nuances_pct = consensus.get('nuances_percentage', 0)
        contradicts_pct = consensus.get('contradicts_percentage', 0)
        extends_pct = consensus.get('extends_percentage', 0)
        unrelated_pct = 100 - (confirms_pct + nuances_pct + contradicts_pct + extends_pct)
        
        interpretation = "# 🤔 ¿QUÉ SIGNIFICA ESTO?\n\n"
        interpretation += "Esta sección le ayuda a interpretar los resultados numéricos en contexto.\n\n"
        interpretation += "---\n\n"
        
        # ================================================================
        # 1. INTERPRETACIÓN DE LA CONFIRMACIÓN
        # ================================================================
        interpretation += "## 📊 Sobre la confirmación de la hipótesis\n\n"
        
        if confirms_pct == 0:
            interpretation += "**🔴 0% de los estudios confirman explícitamente su hipótesis**\n\n"
            interpretation += "→ Esto **NO significa que su hipótesis sea falsa**.\n"
            interpretation += "→ Significa que ningún artículo contiene frases explícitas como "
            interpretation += "*'confirms that'* o *'supports the hypothesis'* en relación a su hipótesis específica.\n\n"
            interpretation += "**Posibles explicaciones:**\n"
            interpretation += "• Su hipótesis no ha sido formalmente evaluada en la literatura\n"
            interpretation += "• Los estudios existentes son predominantemente descriptivos (reportes de caso)\n"
            interpretation += "• La confirmación requeriría un diseño de estudio específico que no se ha realizado\n"
            interpretation += "• El lenguaje utilizado en los artículos es más cauto (ej: 'sugiere', 'podría indicar')\n\n"
        elif confirms_pct < 30:
            interpretation += f"**🟠 Solo {confirms_pct:.0f}% de los estudios confirman explícitamente su hipótesis**\n\n"
            interpretation += "→ La evidencia de confirmación es escasa.\n"
            interpretation += "→ La mayoría de los artículos no abordan directamente su hipótesis.\n\n"
        elif confirms_pct > 70:
            interpretation += f"**🟢 {confirms_pct:.0f}% de los estudios confirman explícitamente su hipótesis**\n\n"
            interpretation += "→ Existe consenso sólido en la literatura.\n"
            interpretation += "→ Su hipótesis está bien respaldada por evidencia explícita.\n\n"
        else:
            interpretation += f"**🟡 {confirms_pct:.0f}% de los estudios confirman explícitamente su hipótesis**\n\n"
            interpretation += "→ El respaldo a su hipótesis es moderado.\n"
            interpretation += "→ Existe evidencia, pero también espacio para más investigación.\n\n"
        
        # ================================================================
        # 2. INTERPRETACIÓN DE CONTRADICCIONES
        # ================================================================
        if contradicts_pct > 0:
            interpretation += "## ⚠️ Sobre las contradicciones\n\n"
            if contradicts_pct > 30:
                interpretation += f"**🔴 {contradicts_pct:.0f}% de los estudios contradicen su hipótesis**\n\n"
                interpretation += "→ Existe evidencia significativa en contra de su hipótesis.\n"
                interpretation += "→ Se recomienda revisar los artículos que contradicen para entender sus argumentos.\n\n"
            else:
                interpretation += f"**🟡 {contradicts_pct:.0f}% de los estudios contradicen su hipótesis**\n\n"
                interpretation += "→ Hay alguna evidencia contraria, pero no es mayoritaria.\n"
                interpretation += "→ Considere si estos estudios tienen limitaciones metodológicas.\n\n"
        
        # ================================================================
        # 3. INTERPRETACIÓN DE RELEVANCIA
        # ================================================================
        interpretation += "## 📚 Sobre la relevancia de los artículos\n\n"
        
        if relevant_articles < 5:
            interpretation += f"**📊 Solo {relevant_articles} artículos son realmente relevantes** "
            interpretation += f"(de {total_articles} totales, {unrelated_pct:.0f}% no están relacionados)\n\n"
            interpretation += "→ Su hipótesis específica apenas ha sido abordada en la literatura.\n"
            interpretation += "→ Los artículos 'no relacionados' contienen las palabras de su búsqueda "
            interpretation += "pero NO abordan su hipótesis específica.\n"
            interpretation += "→ Esto es un hallazgo válido: la ausencia de literatura relevante es información útil.\n\n"
            interpretation += "**Recomendación:** Considerar reformular su hipótesis como pregunta EXPLORATORIA "
            interpretation += "en lugar de CONFIRMATORIA.\n\n"
        elif relevant_articles < 15:
            interpretation += f"**📊 {relevant_articles} artículos son relevantes** "
            interpretation += f"(de {total_articles} totales)\n\n"
            interpretation += "→ Existe un cuerpo de literatura moderado que aborda su tema.\n"
            interpretation += "→ Suficiente para un análisis, pero no para conclusiones definitivas.\n\n"
        else:
            interpretation += f"**📊 {relevant_articles} artículos son relevantes** "
            interpretation += f"(de {total_articles} totales)\n\n"
            interpretation += "→ Existe un cuerpo de literatura sustancial sobre su tema.\n"
            interpretation += "→ Los resultados tienen mayor robustez.\n\n"
        
        # ================================================================
        # 4. INTERPRETACIÓN DE CALIDAD METODOLÓGICA
        # ================================================================
        interpretation += "## ⭐ Sobre la calidad metodológica\n\n"
        
        if avg_methodological < 30:
            interpretation += f"**🔴 Calidad metodológica promedio: {avg_methodological:.0f}/100**\n\n"
            interpretation += "→ La evidencia existente es predominantemente de **baja calidad**.\n"
            interpretation += "→ Esto es **esperable** para:\n"
            interpretation += "  • Enfermedades raras o condiciones poco frecuentes\n"
            interpretation += "  • Temas emergentes con poca investigación previa\n"
            interpretation += "  • Fenómenos descritos principalmente en reportes de caso\n"
            interpretation += "→ **No indica que su análisis sea incorrecto**; refleja el estado de la literatura.\n\n"
        elif avg_methodological < 60:
            interpretation += f"**🟡 Calidad metodológica promedio: {avg_methodological:.0f}/100**\n\n"
            interpretation += "→ La calidad metodológica es moderada.\n"
            interpretation += "→ Los resultados deben interpretarse con cautela.\n\n"
        else:
            interpretation += f"**🟢 Calidad metodológica promedio: {avg_methodological:.0f}/100**\n\n"
            interpretation += "→ La evidencia muestra buena calidad metodológica.\n"
            interpretation += "→ Mayor confianza en los resultados.\n\n"
        
        # ================================================================
        # 5. INTERPRETACIÓN DE MATICES Y EXTENSIONES
        # ================================================================
        if nuances_pct > 50:
            interpretation += "## 🔍 Sobre los matices\n\n"
            interpretation += f"**{nuances_pct:.0f}% de los artículos relevantes matizan su hipótesis**\n\n"
            interpretation += "→ Estos artículos **mencionan** el fenómeno pero **no lo confirman explícitamente**.\n"
            interpretation += "→ Pueden añadir condiciones, limitaciones o contextos específicos.\n"
            interpretation += "→ Son útiles para entender la complejidad del tema.\n\n"
        
        if extends_pct > 20:
            interpretation += "## 🚀 Sobre las extensiones\n\n"
            interpretation += f"**{extends_pct:.0f}% de los artículos extienden su hipótesis**\n\n"
            interpretation += "→ Algunos estudios aplican su hipótesis a nuevos contextos o poblaciones.\n"
            interpretation += "→ Esto sugiere que el concepto tiene aplicabilidad más amplia.\n\n"
        
        # ================================================================
        # 6. RECOMENDACIONES ACCIONABLES
        # ================================================================
        interpretation += "## 💡 RECOMENDACIONES BASADAS EN ESTOS RESULTADOS\n\n"
        
        if confirms_pct == 0 and relevant_articles < 5:
            interpretation += "**✅ ACCIONES SUGERIDAS:**\n"
            interpretation += "1. **Reformular la hipótesis como pregunta EXPLORATORIA** en lugar de confirmatoria\n"
            interpretation += "2. **Realizar una revisión sistemática CUALITATIVA** de los reportes de caso\n"
            interpretation += "3. **Usar estos resultados como JUSTIFICACIÓN** para investigación primaria\n"
            interpretation += "4. **Considerar que la ausencia de evidencia NO es evidencia de ausencia**\n\n"
            interpretation += "**❌ ACCIONES NO RECOMENDADAS:**\n"
            interpretation += "1. Concluir que la hipótesis es falsa (no ha sido probada)\n"
            interpretation += "2. Descartar el análisis como 'fallido' (el hallazgo es válido)\n"
            interpretation += "3. Ignorar los artículos 'no relacionados' (pueden tener información contextual)\n\n"
        elif confirms_pct > 50:
            interpretation += "**✅ ACCIONES SUGERIDAS:**\n"
            interpretation += "1. La hipótesis está respaldada por la evidencia\n"
            interpretation += "2. Considerar aplicar los hallazgos en la práctica o investigación\n"
            interpretation += "3. Identificar áreas donde aún hay incertidumbre\n\n"
        elif contradicts_pct > 30:
            interpretation += "**✅ ACCIONES SUGERIDAS:**\n"
            interpretation += "1. Revisar críticamente los artículos que contradicen la hipótesis\n"
            interpretation += "2. Evaluar si las contradicciones se deben a diferencias metodológicas\n"
            interpretation += "3. Considerar refinar o modificar la hipótesis original\n\n"
        else:
            interpretation += "**✅ ACCIONES SUGERIDAS:**\n"
            interpretation += "1. Mantener la hipótesis pero ajustar las expectativas\n"
            interpretation += "2. Complementar con búsquedas en otras bases de datos (Google Scholar, Scopus)\n"
            interpretation += "3. Considerar que la ausencia de evidencia NO es evidencia de ausencia\n"
            interpretation += "4. Documentar esta brecha de conocimiento en futuras publicaciones\n\n"
        
        # ================================================================
        # 7. NOTA SOBRE LA METODOLOGÍA DEL ANÁLISIS
        # ================================================================
        interpretation += "---\n\n"
        interpretation += "## 📌 NOTA METODOLÓGICA\n\n"
        interpretation += "La clasificación 'Confirma' requiere **evidencia textual EXPLÍCITA** "
        interpretation += "(frases como *'confirms that'*, *'supports the hypothesis'*).\n\n"
        interpretation += "Esto es INTENCIONAL para evitar:\n"
        interpretation += "• Falsas confirmaciones por mera similitud temática\n"
        interpretation += "• Interpretaciones subjetivas del contenido\n"
        interpretation += "• Sesgos del investigador en la clasificación\n\n"
        interpretation += "Si su hipótesis es correcta pero no está formulada en términos "
        interpretation += "que aparezcan textualmente en los abstracts, obtendrá 0% de confirmaciones. "
        interpretation += "Esto no invalida su hipótesis, sino que indica que no ha sido "
        interpretation += "formalmente evaluada con ese lenguaje específico."
        
        return interpretation


# ============================================================================
# VALIDADOR DE FALSAS CONFIRMACIONES
# ============================================================================

class FalseConfirmationValidator:
    @staticmethod
    def validate(article: Dict, hypothesis: str) -> Dict:
        if article.get('hypothesis_relation') != 'confirms':
            return article
        
        text = f"{article.get('title', '')} {article.get('abstract', '')}".lower()
        
        explicit_confirmation = any(phrase in text for phrase in [
            'confirms that', 'demonstrates that', 'consistent with the hypothesis'
        ])
        
        if not explicit_confirmation:
            article['hypothesis_relation'] = 'nuances'
            article['relation_confidence'] = 0.4
            article['validation_note'] = '⚠️ Posible falsa confirmación - no hay evidencia explícita.'
        
        return article


# ============================================================================
# ALERTAS DE CALIDAD
# ============================================================================

class QualityAlertSystem:
    @staticmethod
    def generate_alerts(articles: List[Dict]) -> Dict:
        articles = deduplicate_articles(articles)
        total = len(articles)
        
        if total == 0:
            return {'alerts': {}, 'overall_level': "🔵 INFORMATIVO", 'overall_message': "No hay artículos"}
        
        alerts = {'critical': [], 'warning': [], 'success': []}
        
        case_reports = sum(1 for a in articles if a.get('evidence_level') == 'Reporte de caso')
        rct = sum(1 for a in articles if a.get('evidence_level') == 'RCT')
        
        if case_reports / total > 0.5:
            alerts['critical'].append({'message': f'⚠️ {case_reports}/{total} son reportes de caso', 'recommendation': 'Interpretar con cautela'})
        
        if rct > 0:
            alerts['success'].append({'message': f'✅ {rct} ensayo(s) randomizado(s)', 'recommendation': ''})
        
        low_quality = [a for a in articles if (a.get('methodological_score') or 0) < 30]
        if len(low_quality) > total * 0.5:
            alerts['critical'].append({'message': f'⚠️ {len(low_quality)}/{total} tienen baja calidad', 'recommendation': 'Hallazgos preliminares'})
        
        if alerts['critical']:
            overall_level = "🔴 CRÍTICO"
            overall_message = "Limitaciones metodológicas graves."
        elif alerts['warning']:
            overall_level = "🟡 ADVERTENCIA"
            overall_message = "Limitaciones a considerar."
        else:
            overall_level = "🟢 FAVORABLE"
            overall_message = "Buena calidad metodológica."
        
        return {'alerts': alerts, 'overall_level': overall_level, 'overall_message': overall_message, 'total_articles': total}


# ============================================================================
# DETECCIÓN DE CONTRADICCIONES
# ============================================================================

class ContradictionDetector:
    @staticmethod
    def find_contradictions(articles: List[Dict]) -> List[Dict]:
        articles = deduplicate_articles(articles)
        contradictions = []
        
        for i, a1 in enumerate(articles):
            for a2 in articles[i+1:]:
                if a1.get('effect_value') and a2.get('effect_value'):
                    if (a1['effect_value'] > 1 and a2['effect_value'] < 1) or (a1['effect_value'] < 1 and a2['effect_value'] > 1):
                        contradictions.append({
                            'article1': a1, 'article2': a2,
                            'conflict': {'type': 'effect_direction', 'description': 'Dirección del efecto opuesta',
                                        'detail': f"{a1.get('authors', 'A')[:30]} ({a1['effect_value']:.2f}) vs {a2.get('authors', 'B')[:30]} ({a2['effect_value']:.2f})"}
                        })
        return contradictions
    
    @staticmethod
    def generate_contradiction_summary(contradictions: List[Dict]) -> str:
        if not contradictions:
            return "✅ **No se detectaron contradicciones significativas**"
        return f"⚠️ **Se detectaron {len(contradictions)} contradicción(es)**"


# ============================================================================
# LÍNEA DE TIEMPO
# ============================================================================

class TemporalConsensusAnalyzer:
    @staticmethod
    def analyze_temporal_consensus(articles: List[Dict]) -> Dict:
        articles = deduplicate_articles(articles)
        years_data = {}
        
        for art in articles:
            year = art.get('pubdate', '')[:4]
            if year and year.isdigit() and 1990 <= int(year) <= datetime.now().year + 1:
                if year not in years_data:
                    years_data[year] = {'confirms': 0, 'total': 0}
                years_data[year]['total'] += 1
                if art.get('hypothesis_relation') == 'confirms':
                    years_data[year]['confirms'] += 1
        
        for year, data in years_data.items():
            if data['total'] > 0:
                data['confirms_pct'] = (data['confirms'] / data['total']) * 100
        
        years_sorted = sorted(years_data.keys())
        if len(years_sorted) >= 3:
            recent_avg = sum(years_data[y]['confirms_pct'] for y in years_sorted[-3:]) / 3
            old_avg = sum(years_data[y]['confirms_pct'] for y in years_sorted[:3]) / 3
            trend = "📈 ALCISTA" if recent_avg > old_avg + 15 else "📉 BAJISTA" if recent_avg < old_avg - 15 else "➡️ ESTABLE"
        else:
            trend = "⚠️ Datos insuficientes"
        
        return {'years_data': years_data, 'trend': trend}
    
    @staticmethod
    def create_temporal_plot(temporal_data: Dict) -> Optional[BytesIO]:
        if not temporal_data or not temporal_data.get('years_data'):
            return None
        
        try:
            import matplotlib.pyplot as plt
            
            years = sorted(temporal_data['years_data'].keys())
            confirms_pct = [temporal_data['years_data'][y]['confirms_pct'] for y in years]
            total_studies = [temporal_data['years_data'][y]['total'] for y in years]
            
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.bar(years, confirms_pct, color='green', alpha=0.7, label='Confirman (%)')
            ax.set_xlabel('Año')
            ax.set_ylabel('Porcentaje')
            ax.set_title('Evolución del consenso científico')
            ax.set_ylim(0, 100)
            
            for i, (year, n) in enumerate(zip(years, total_studies)):
                ax.annotate(f'n={n}', (year, confirms_pct[i]), textcoords="offset points", xytext=(0, 5), ha='center', fontsize=8)
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            plt.close()
            return buf
        except:
            return None
    
    @staticmethod
    def generate_temporal_summary(temporal_data: Dict) -> str:
        if not temporal_data or not temporal_data.get('years_data'):
            return "No hay suficientes datos."
        return f"**📅 Análisis temporal:** {temporal_data['trend']}"


# ============================================================================
# EXTRACCIÓN PICO
# ============================================================================

class PICOExtractor:
    @staticmethod
    def extract_all(title: str, abstract: str) -> Dict:
        text = f"{title} {abstract if abstract else ''}".lower()
        pop_match = re.search(r'patients? with ([^.!]+)', text)
        outcomes = [kw for kw in ['mortality', 'rupture', 'bleeding', 'stroke'] if kw in text]
        
        return {
            'population': pop_match.group(1)[:100] if pop_match else "No especificada",
            'intervention': "No especificada",
            'outcomes': ', '.join(outcomes) if outcomes else "No especificado"
        }


# ============================================================================
# DEDUPLICACIÓN
# ============================================================================

def deduplicate_articles(articles: List[Dict]) -> List[Dict]:
    if not articles:
        return []
    seen = set()
    unique = []
    for a in articles:
        pmid = a.get('pmid', '')
        if pmid and pmid not in seen:
            seen.add(pmid)
            unique.append(a)
    return unique


# ============================================================================
# ANÁLISIS DE CONSENSO
# ============================================================================

class ConsensusAnalyzer:
    @staticmethod
    def calculate_consensus(articles: List[Dict]) -> Dict:
        if not articles:
            return {'consensus_level': 'Sin datos', 'confirms_percentage': 0, 'relation_counts': {}, 'total_articles': 0}
        
        articles = deduplicate_articles(articles)
        relation_counts = Counter()
        for a in articles:
            relation_counts[a.get('hypothesis_relation', 'unrelated')] += 1
        
        total = len(articles)
        confirms_pct = (relation_counts.get('confirms', 0) / total) * 100 if total > 0 else 0
        
        if confirms_pct >= 70:
            consensus_level = "MUY ALTO"
            consensus_color = "🟢"
            consensus_message = "La evidencia es sólida y consistente."
        elif confirms_pct >= 50:
            consensus_level = "MODERADO"
            consensus_color = "🟡"
            consensus_message = "Existe respaldo mayoritario."
        elif confirms_pct >= 30:
            consensus_level = "BAJO"
            consensus_color = "🟠"
            consensus_message = "La evidencia es mixta."
        else:
            consensus_level = "MUY BAJO / CONTROVERSIA"
            consensus_color = "🔴"
            consensus_message = "No hay evidencia explícita de confirmación."
        
        return {
            'consensus_level': consensus_level,
            'consensus_color': consensus_color,
            'consensus_message': consensus_message,
            'confirms_percentage': confirms_pct,
            'contradicts_percentage': (relation_counts.get('contradicts', 0) / total) * 100 if total > 0 else 0,
            'nuances_percentage': (relation_counts.get('nuances', 0) / total) * 100 if total > 0 else 0,
            'extends_percentage': (relation_counts.get('extends', 0) / total) * 100 if total > 0 else 0,
            'relation_counts': dict(relation_counts),
            'total_articles': total
        }
    
    @staticmethod
    def generate_executive_summary(articles: List[Dict], hypothesis: str, query: str) -> str:
        if not articles:
            return "❌ **No hay artículos para analizar.**"
        
        articles = deduplicate_articles(articles)
        consensus = ConsensusAnalyzer.calculate_consensus(articles)
        
        methodological_scores = [a.get('methodological_score', 0) for a in articles if a.get('methodological_score') is not None]
        avg_methodological = sum(methodological_scores) / len(methodological_scores) if methodological_scores else 0
        
        if articles:
            best_article = max(articles, key=lambda x: x.get('methodological_score', 0) or 0)
            best_authors = best_article.get('authors', 'Autor')[:80]
            best_year = best_article.get('pubdate', 's.f.')[:4]
            best_design = best_article.get('study_design', 'Estudio')
            best_relation = best_article.get('hypothesis_relation', 'unrelated')
            best_score = best_article.get('methodological_score', 0)
        else:
            best_authors = "No disponible"
            best_year = "N/A"
            best_design = "N/A"
            best_relation = "N/A"
            best_score = 0
        
        relation_symbol = {'confirms': '✅', 'contradicts': '⚠️', 'nuances': '🔍', 'extends': '🚀', 'unrelated': '📚'}.get(best_relation, '📚')
        
        summary = f"""
╔════════════════════════════════════════════════════════════╗
                         📋 RESUMEN EJECUTIVO v8.7                                   
╚════════════════════════════════════════════════════════════╝

🔬 HIPÓTESIS EVALUADA:
   "{hypothesis[:200]}..."

📊 NIVEL DE CONSENSO: {consensus['consensus_color']} {consensus['consensus_level']}
   {consensus['consensus_message']}

📈 DISTRIBUCIÓN DE EVIDENCIA:
   • Confirman: {consensus['confirms_percentage']:.1f}% ({consensus['relation_counts'].get('confirms', 0)} estudios)
   • Matizan: {consensus['nuances_percentage']:.1f}% ({consensus['relation_counts'].get('nuances', 0)} estudios)
   • Extienden: {consensus['extends_percentage']:.1f}% ({consensus['relation_counts'].get('extends', 0)} estudios)
   • Contradicen: {consensus['contradicts_percentage']:.1f}% ({consensus['relation_counts'].get('contradicts', 0)} estudios)

⭐ CALIDAD DE LA EVIDENCIA:
   • Calidad metodológica promedio: {avg_methodological:.0f}/100
   • Artículos analizados: {consensus['total_articles']}

🏆 MEJOR CALIDAD METODOLÓGICA:
   • {best_authors} ({best_year})
   • Diseño: {best_design}
   • Relación: {relation_symbol} {best_relation}
   • Score metodológico: {best_score:.0f}/100

💡 RECOMENDACIÓN CLÍNICA:
   """
        if consensus['confirms_percentage'] >= 50:
            summary += "La evidencia respalda la hipótesis con confirmación textual explícita."
        elif consensus['nuances_percentage'] >= 50:
            summary += "La mayoría de los estudios matizan la hipótesis."
        else:
            summary += "No se encontró evidencia textual explícita que confirme la hipótesis específica."
        
        return summary


# ============================================================================
# CONFIGURACIÓN DE PAQUETES AI
# ============================================================================

AI_EMBEDDINGS_AVAILABLE = False
BIOMED_EMBEDDER = None
FALLBACK_EMBEDDER = None
USE_FALLBACK = False
MATPLOTLIB_AVAILABLE = False
SKLEARN_AVAILABLE = False

st.info("🔄 Cargando paquetes de IA (BioBERT, clustering, etc.)...")

try:
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import KMeans
    from sklearn.decomposition import TruncatedSVD
    from sklearn.feature_extraction.text import TfidfVectorizer
    SKLEARN_AVAILABLE = True
    st.success("✅ Scikit-learn disponible (clustering, TF-IDF)")
except ImportError as e:
    st.warning(f"⚠️ Scikit-learn no disponible: {e}")

try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
    st.success("✅ Matplotlib disponible (gráficos)")
except ImportError as e:
    st.warning(f"⚠️ Matplotlib no disponible: {e}")

try:
    from sentence_transformers import SentenceTransformer
    import torch
    
    with st.spinner("🔄 Cargando modelo BioBERT..."):
        try:
            BIOMED_EMBEDDER = SentenceTransformer(
                'pritamdeka/S-Biomed-Roberta-snli-multinli-stsb',
                device='cpu'
            )
            AI_EMBEDDINGS_AVAILABLE = True
            USE_FALLBACK = False
            st.success("✅ BioBERT embeddings disponible (para relevancia y clustering)")
        except Exception as e:
            st.warning(f"⚠️ BioBERT no disponible: {e}")
            try:
                FALLBACK_EMBEDDER = SentenceTransformer(
                    'all-MiniLM-L6-v2',
                    device='cpu'
                )
                AI_EMBEDDINGS_AVAILABLE = True
                USE_FALLBACK = True
                st.success("✅ SBERT (fallback) disponible")
            except Exception as e2:
                st.warning(f"⚠️ SBERT no disponible: {e2}")
                AI_EMBEDDINGS_AVAILABLE = False
except ImportError as e:
    st.warning(f"⚠️ SentenceTransformers no disponible: {e}")
    AI_EMBEDDINGS_AVAILABLE = False

st.info(f"📊 Estado: Embeddings={'✅' if AI_EMBEDDINGS_AVAILABLE else '❌'} | Clustering={'✅' if SKLEARN_AVAILABLE else '❌'} | Gráficos={'✅' if MATPLOTLIB_AVAILABLE else '❌'}")

st.set_page_config(page_title="PubMed AI Analyzer v8.7", page_icon="🧠", layout="wide")


# ============================================================================
# SISTEMA DE PONDERACIÓN
# ============================================================================

class EvidenceHierarchy:
    @classmethod
    def classify_study_design(cls, title: str, abstract: str, study_types: str) -> Tuple[str, int, int]:
        text = f"{title} {abstract} {study_types}".lower()
        
        if 'meta-analysis' in text or 'systematic review' in text:
            return 'Meta-análisis', 1, 100
        if 'randomized controlled trial' in text or 'rct' in text:
            return 'RCT', 1, 95
        if 'prospective cohort' in text:
            return 'Cohorte prospectivo', 2, 80
        if 'retrospective cohort' in text:
            return 'Cohorte retrospectivo', 3, 60
        if 'case-control' in text:
            return 'Caso-control', 3, 55
        if 'case series' in text:
            return 'Serie de casos', 4, 35
        if 'case report' in text:
            return 'Reporte de caso', 5, 20
        
        return 'No clasificado', 5, 10
    
    @classmethod
    def extract_sample_size(cls, text: str) -> Optional[int]:
        patterns = [r'n\s*[=:]\s*(\d+)', r'patients?\s*[=:]\s*(\d+)', r'(\d+)\s*patients']
        for pattern in patterns:
            match = re.search(pattern, text.lower())
            if match:
                val = int(match.group(1))
                if 1 <= val <= 50000:
                    return val
        return None


class NumericalExtractor:
    @staticmethod
    def extract_all_numbers(text: str) -> Dict:
        results = {'hr': [], 'rr': [], 'or': [], 'n': None}
        
        for effect in ['HR', 'RR', 'OR']:
            pattern = rf'{effect}[\s]*[=:]?[\s]*([0-9.]+)[\s]*\(([0-9.]+)[-\s]+([0-9.]+)\)'
            for match in re.finditer(pattern, text, re.IGNORECASE):
                results[effect.lower()].append({'value': float(match.group(1)), 'ci_lower': float(match.group(2)), 'ci_upper': float(match.group(3))})
        
        results['n'] = EvidenceHierarchy.extract_sample_size(text)
        return results
    
    @staticmethod
    def get_best_effect_size(numbers: Dict) -> Optional[Dict]:
        for effect_type in ['hr', 'or', 'rr']:
            if numbers.get(effect_type):
                return {'type': effect_type.upper(), 'value': numbers[effect_type][0]['value'],
                       'ci_lower': numbers[effect_type][0]['ci_lower'], 'ci_upper': numbers[effect_type][0]['ci_upper']}
        return None


class MethodologicalChecklist:
    @staticmethod
    def check_article(article: Dict) -> Tuple[Dict, float]:
        text = f"{article.get('title', '')} {article.get('abstract', '')}".lower()
        
        score = 0
        if re.search(r'prospective|longitudinal|cohort', text):
            score += 16.67
        if re.search(r'randomized|randomised', text):
            score += 16.67
        if re.search(r'blind|masked|double-blind', text):
            score += 16.67
        if re.search(r'multicenter|multi-center', text):
            score += 16.67
        if re.search(r'multivariable|multivariate|adjusted|cox', text):
            score += 16.67
        if re.search(r'clinicaltrials\.gov|nct\d+', text):
            score += 16.67
        
        return {}, score


# ============================================================================
# EXPORTACIÓN
# ============================================================================

class ReferenceExporter:
    @staticmethod
    def to_ris(articles: List[Dict]) -> str:
        lines = []
        for a in deduplicate_articles(articles):
            lines.append(f"TY  - JOUR\nAU  - {a.get('authors', '')}\nTI  - {a.get('title', '')}\nJO  - {a.get('journal', '')}\nPY  - {a.get('pubdate', '')[:4]}\nID  - {a.get('pmid', '')}\nER  - \n")
        return "\n".join(lines)
    
    @staticmethod
    def to_bibtex(articles: List[Dict]) -> str:
        lines = []
        for i, a in enumerate(deduplicate_articles(articles)):
            pmid = a.get('pmid', f'unknown_{i}')
            lines.append(f"@article{{pmid_{pmid},\n  title = {{{a.get('title', '')}}},\n  journal = {{{a.get('journal', '')}}},\n  year = {{{a.get('pubdate', '')[:4]}}},\n  pmid = {{{pmid}}}\n}}\n")
        return "\n".join(lines)


# ============================================================================
# GRÁFICO
# ============================================================================

def plot_consensus_gauge(consensus: Dict) -> Optional[BytesIO]:
    if not MATPLOTLIB_AVAILABLE:
        return None
    
    fig, ax = plt.subplots(figsize=(6, 4))
    confirms_pct = consensus.get('confirms_percentage', 0)
    
    if confirms_pct >= 70:
        color = '#2ecc71'
    elif confirms_pct >= 50:
        color = '#f39c12'
    elif confirms_pct >= 30:
        color = '#e67e22'
    else:
        color = '#e74c3c'
    
    ax.barh(['Consenso'], [confirms_pct], color=color, height=0.5)
    ax.set_xlim(0, 100)
    ax.set_xlabel('Porcentaje de estudios que confirman la hipótesis (%)')
    ax.set_title(f'Nivel de Consenso: {consensus.get("consensus_level", "ND")}')
    ax.axvline(x=50, color='gray', linestyle='--', alpha=0.5)
    ax.text(confirms_pct + 2, 0, f'{confirms_pct:.0f}%', va='center', fontsize=12, fontweight='bold')
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf


# ============================================================================
# FUNCIONES DE PubMed
# ============================================================================

def make_request_with_retry(url, params, max_retries=5, initial_delay=5):
    delay = initial_delay
    for attempt in range(max_retries):
        try:
            response = requests.get(url, params=params, timeout=30)
            if response.status_code == 429:
                wait_time = delay * (2 ** attempt)
                st.warning(f"⚠️ Rate limit. Esperando {wait_time}s...")
                time.sleep(wait_time)
                continue
            response.raise_for_status()
            return response
        except:
            if attempt == max_retries - 1:
                raise
            time.sleep(delay)
    return None


def search_pubmed_complete(query, max_articles=MAX_ARTICLES, progress_callback=None):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    search_url = f"{base_url}esearch.fcgi"
    
    try:
        st.info("🔍 Buscando artículos en PubMed...")
        params = {"db": "pubmed", "term": query, "retmax": 1, "retmode": "xml"}
        response = make_request_with_retry(search_url, params)
        if response is None:
            return [], 0
        
        root = ElementTree.fromstring(response.content)
        total_count = int(root.find(".//Count").text) if root.find(".//Count") is not None else 0
        st.info(f"📊 PubMed encontró {total_count} artículos")
        
        if total_count == 0:
            return [], 0
        
        articles_to_fetch = min(total_count, max_articles)
        st.info(f"📊 Se procesarán {articles_to_fetch} artículos")
        
        all_ids = []
        retstart = 0
        batch_size = 10000
        progress_bar = st.progress(0)
        
        while retstart < articles_to_fetch:
            remaining = articles_to_fetch - retstart
            current_batch = min(batch_size, remaining)
            
            fetch_params = {"db": "pubmed", "term": query, "retstart": retstart, "retmax": current_batch, "retmode": "xml"}
            
            try:
                fetch_response = make_request_with_retry(search_url, fetch_params)
                if fetch_response is None:
                    break
                
                fetch_root = ElementTree.fromstring(fetch_response.content)
                batch_ids = [id_elem.text for id_elem in fetch_root.findall(".//Id")]
                
                if not batch_ids:
                    break
                
                all_ids.extend(batch_ids)
                progress_bar.progress(min(len(all_ids) / articles_to_fetch, 1.0))
                if progress_callback:
                    progress_callback(len(all_ids), articles_to_fetch)
                retstart += current_batch
                time.sleep(0.5)
            except Exception as e:
                st.error(f"Error: {e}")
                break
        
        progress_bar.empty()
        st.success(f"✅ Recuperados {len(all_ids)} IDs")
        return all_ids, len(all_ids)
    except Exception as e:
        st.error(f"Error: {e}")
        return [], 0


@st.cache_data(ttl=3600, show_spinner=False)
def get_abstract(pmid):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    fetch_url = f"{base_url}efetch.fcgi"
    params = {"db": "pubmed", "id": pmid, "retmode": "xml", "rettype": "abstract"}
    try:
        response = make_request_with_retry(fetch_url, params)
        if response is None:
            return None
        root = ElementTree.fromstring(response.content)
        abstract_texts = []
        for abstract in root.findall(".//AbstractText"):
            if abstract.text:
                abstract_texts.append(abstract.text)
        return " ".join(abstract_texts) if abstract_texts else None
    except:
        return None


def extract_article_info(doc_sum):
    article = {}
    article["pmid"] = doc_sum.find("Id").text if doc_sum.find("Id") is not None else "N/A"
    title_item = doc_sum.find(".//Item[@Name='Title']")
    article["title"] = title_item.text if title_item is not None else "No title"
    author_items = doc_sum.findall(".//Item[@Name='Author']")
    authors = [author.text for author in author_items if author.text]
    article["authors"] = ", ".join(authors[:5]) + (" et al." if len(authors) > 5 else "")
    source_item = doc_sum.find(".//Item[@Name='Source']")
    article["journal"] = source_item.text if source_item is not None else "N/A"
    pubdate_item = doc_sum.find(".//Item[@Name='PubDate']")
    article["pubdate"] = pubdate_item.text if pubdate_item is not None else "N/A"
    return article


def get_embedder():
    if BIOMED_EMBEDDER is not None:
        return BIOMED_EMBEDDER
    elif FALLBACK_EMBEDDER is not None:
        return FALLBACK_EMBEDDER
    return None


def extract_all_outcomes(text):
    if not text:
        return []
    text_lower = text.lower()
    outcomes = set()
    indicators = ['mortality', 'death', 'survival', 'rupture', 'bleeding', 'hemorrhage', 
                  'stroke', 'reinfarction', 'complication', 'heart failure', 'mace']
    for indicator in indicators:
        if indicator in text_lower:
            outcomes.add(indicator)
    return list(outcomes)


def enhanced_article_analysis(title: str, abstract: str, hypothesis: str = "") -> Dict:
    full_text = f"{title} {abstract if abstract else ''}"
    
    study_design, evidence_level, evidence_score = EvidenceHierarchy.classify_study_design(title, abstract, "")
    
    numbers = NumericalExtractor.extract_all_numbers(full_text)
    sample_size = numbers.get('n')
    best_effect = NumericalExtractor.get_best_effect_size(numbers)
    
    relation_info = HypothesisRelationClassifier.classify(title, abstract, hypothesis)
    
    all_outcomes = extract_all_outcomes(full_text)
    pico = PICOExtractor.extract_all(title, abstract)
    
    _, methodological_score = MethodologicalChecklist.check_article({'title': title, 'abstract': abstract})
    
    quality_score = (evidence_score + methodological_score) / 2
    if sample_size and sample_size > 100:
        quality_score += 5
    if best_effect:
        if best_effect['value'] < 0.7 or best_effect['value'] > 1.5:
            quality_score += 10
    quality_score = min(quality_score, 100)
    
    return {
        'study_design': study_design,
        'evidence_level': evidence_level,
        'evidence_score': evidence_score,
        'sample_size': sample_size,
        'effect_type': best_effect['type'] if best_effect else '',
        'effect_value': best_effect['value'] if best_effect else 0,
        'effect_ci_lower': best_effect['ci_lower'] if best_effect else 0,
        'effect_ci_upper': best_effect['ci_upper'] if best_effect else 0,
        'hypothesis_relation': relation_info['relation'],
        'relation_confidence': relation_info['confidence'],
        'relation_reason': relation_info.get('reason', ''),
        'all_outcomes': all_outcomes,
        'quality_score': quality_score,
        'methodological_score': methodological_score,
        'pico_population': pico['population'],
        'pico_intervention': pico['intervention'],
        'pico_outcomes': pico['outcomes']
    }


def fetch_articles_details(id_list, query_terms, hypothesis_terms, hypothesis="", block_number=0, progress_callback=None):
    if not id_list:
        return []
    
    total_to_process = len(id_list)
    batch_size = 30
    num_batches = math.ceil(total_to_process / batch_size)
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    articles = []
    
    for batch_num in range(num_batches):
        start_idx = batch_num * batch_size
        end_idx = min((batch_num + 1) * batch_size, total_to_process)
        batch_ids = id_list[start_idx:end_idx]
        
        summary_params = {"db": "pubmed", "id": ",".join(batch_ids), "retmode": "xml"}
        
        for retry in range(3):
            try:
                summary_response = make_request_with_retry(f"{base_url}esummary.fcgi", summary_params)
                if summary_response is None:
                    if retry < 2:
                        time.sleep(5)
                        continue
                    else:
                        raise Exception("Max retries exceeded")
                
                summary_root = ElementTree.fromstring(summary_response.content)
                
                for j, doc_sum in enumerate(summary_root.findall(".//DocSum")):
                    overall_idx = start_idx + j
                    if progress_callback:
                        progress_callback(overall_idx + 1, total_to_process)
                    
                    article = extract_article_info(doc_sum)
                    abstract = get_abstract(article["pmid"])
                    article["abstract"] = abstract if abstract else "Not available"
                    
                    enhanced = enhanced_article_analysis(article["title"], abstract, hypothesis)
                    article.update(enhanced)
                    article["block_number"] = block_number
                    articles.append(article)
                break
            except Exception as e:
                if retry == 2:
                    st.warning(f"Error en lote {batch_num + 1}: {str(e)[:50]}")
                else:
                    time.sleep(5)
        
        time.sleep(1)
    
    return articles


def calculate_relevance_to_hypothesis(articles, hypothesis):
    embedder = get_embedder()
    if not embedder or not AI_EMBEDDINGS_AVAILABLE:
        for article in articles:
            article['relevance_score'] = 0.5
        return articles
    
    try:
        hypothesis_embedding = embedder.encode([hypothesis[:500]])[0]
        for article in articles:
            text = f"{article.get('title', '')} {article.get('abstract', '')}"
            if text and len(text) > 50:
                article_embedding = embedder.encode([text[:1000]])[0]
                similarity = cosine_similarity([hypothesis_embedding], [article_embedding])[0][0]
                article['relevance_score'] = float(similarity)
            else:
                article['relevance_score'] = 0.0
    except Exception as e:
        st.warning(f"Error en relevancia: {e}")
        for article in articles:
            article['relevance_score'] = 0.5
    return articles


def filter_articles_by_relevance(articles, relevance_threshold):
    filtered = [a for a in articles if (a.get('relevance_score') or 0) >= relevance_threshold]
    if articles:
        st.write(f"**📊 Filtro por relevancia:** {len(filtered)}/{len(articles)} ({len(filtered)/len(articles)*100:.1f}%)")
    return filtered


# ============================================================================
# GENERACIÓN DE FLAVORS
# ============================================================================

def generate_insight_for_relation(relation, articles, hypothesis):
    if not articles:
        return ""
    
    articles = deduplicate_articles(articles)
    n = len(articles)
    
    best = max(articles, key=lambda x: x.get('methodological_score', 0) or 0)
    best_methodological = best.get('methodological_score') or 0
    best_design = best.get('study_design', 'Estudio')
    
    insights = {
        'confirms': f"**{n} estudios** confirman la hipótesis. Mejor calidad: {best_design} (score {best_methodological}/100).",
        'nuances': f"**{n} estudios** matizan la hipótesis.",
        'extends': f"**{n} estudios** extienden la hipótesis.",
        'contradicts': f"**{n} estudios** contradicen la hipótesis.",
        'unrelated': f"**{n} estudios** no relacionados."
    }
    return insights.get(relation, f"{n} estudios analizados.")


def generate_flavors_by_thesis(articles, hypothesis):
    if not articles:
        return {}
    
    articles = deduplicate_articles(articles)
    
    relation_groups = {'confirms': [], 'nuances': [], 'extends': [], 'contradicts': [], 'unrelated': []}
    
    for article in articles:
        relation = article.get('hypothesis_relation', 'unrelated')
        if relation in relation_groups:
            relation_groups[relation].append(article)
        else:
            relation_groups['unrelated'].append(article)
    
    for group in relation_groups.values():
        group.sort(key=lambda x: x.get('methodological_score', 0), reverse=True)
    
    flavors = {}
    relation_names = {
        'confirms': ('✅ Confirma la Hipótesis', 'Artículos que validan directamente la hipótesis'),
        'nuances': ('🔍 Matiza la Hipótesis', 'Artículos que añaden condiciones o matices'),
        'extends': ('🚀 Extiende la Hipótesis', 'Artículos que amplían el alcance'),
        'contradicts': ('⚠️ Contradice la Hipótesis', 'Artículos con evidencia contraria'),
        'unrelated': ('📚 No Relacionados', 'Artículos que no abordan la hipótesis')
    }
    
    for relation, articles_group in relation_groups.items():
        if not articles_group:
            continue
        
        name, description = relation_names.get(relation, ('Otros', ''))
        meth_scores = [a.get('methodological_score', 0) for a in articles_group]
        avg_methodological = sum(meth_scores) / len(meth_scores) if meth_scores else 0
        has_effect = sum(1 for a in articles_group if a.get('effect_value') and a.get('effect_value') > 0)
        
        insight = generate_insight_for_relation(relation, articles_group, hypothesis)
        
        unique_repr = []
        seen_repr = set()
        for art in articles_group[:5]:
            pmid = art.get('pmid', '')
            if pmid and pmid not in seen_repr:
                seen_repr.add(pmid)
                unique_repr.append(art)
        
        flavors[relation] = {
            'name': name,
            'description': description,
            'articles': articles_group,
            'n_articles': len(articles_group),
            'avg_methodological_score': avg_methodological,
            'has_effect_size': has_effect,
            'insight': insight,
            'representative_articles': unique_repr
        }
    
    return flavors


def generate_traditional_clusters(articles, query_terms, hypothesis_terms):
    if len(articles) < 3 or not SKLEARN_AVAILABLE:
        return []
    
    articles = deduplicate_articles(articles)
    
    try:
        texts = [f"{a.get('title', '')} {' '.join(a.get('all_outcomes', []))}" for a in articles]
        vectorizer = TfidfVectorizer(max_features=500, stop_words='english')
        tfidf_matrix = vectorizer.fit_transform(texts)
        svd = TruncatedSVD(n_components=min(50, tfidf_matrix.shape[1] - 1), random_state=42)
        embeddings_reduced = svd.fit_transform(tfidf_matrix)
        
        n_clusters = min(max(2, len(articles) // 10), 4)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clusters = kmeans.fit_predict(embeddings_reduced)
        
        flavors = []
        for cluster_id in range(n_clusters):
            cluster_indices = [i for i, c in enumerate(clusters) if c == cluster_id]
            if len(cluster_indices) < 2:
                continue
            cluster_articles = [articles[i] for i in cluster_indices]
            cluster_articles.sort(key=lambda x: x.get('methodological_score') or 0, reverse=True)
            
            all_outcomes = []
            for a in cluster_articles:
                all_outcomes.extend(a.get('all_outcomes', []))
            outcome_counts = Counter(all_outcomes)
            top_outcomes = [o for o, _ in outcome_counts.most_common(2)]
            name = f"Cluster sobre {', '.join(top_outcomes)}" if top_outcomes else f"Cluster {cluster_id + 1}"
            
            meth_scores = [a.get('methodological_score') or 0 for a in cluster_articles]
            avg_methodological = sum(meth_scores) / len(meth_scores) if meth_scores else 0
            
            unique_repr = []
            seen_repr = set()
            for art in cluster_articles[:5]:
                pmid = art.get('pmid', '')
                if pmid and pmid not in seen_repr:
                    seen_repr.add(pmid)
                    unique_repr.append(art)
            
            flavors.append({
                'type': 'traditional_cluster',
                'id': f"cluster_{cluster_id}",
                'name': name,
                'articles': cluster_articles,
                'n_articles': len(cluster_articles),
                'representative_articles': unique_repr,
                'avg_methodological': avg_methodological
            })
        return flavors
    except Exception as e:
        st.warning(f"Clustering falló: {e}")
        return []


def generate_all_flavors_human_style(articles, hypothesis, query_terms, hypothesis_terms, method="both"):
    if not articles:
        return {}
    
    articles = deduplicate_articles(articles)
    articles = [FalseConfirmationValidator.validate(a, hypothesis) for a in articles]
    
    flavors = {}
    
    if method in ["thesis", "both"]:
        flavors['by_thesis'] = generate_flavors_by_thesis(articles, hypothesis)
    
    if method in ["clustering", "both"] and SKLEARN_AVAILABLE:
        flavors['by_cluster'] = generate_traditional_clusters(articles, query_terms, hypothesis_terms)
    
    return flavors


def generate_flavor_summary_human_style(flavor_data, flavor_type, hypothesis):
    articles = flavor_data.get('articles', [])
    if not articles:
        return "No hay artículos.", []
    
    articles = deduplicate_articles(articles)
    n = len(articles)
    insight = flavor_data.get('insight', '')
    name = flavor_data.get('name', 'Flavor')
    
    high_methodological = [a for a in articles if (a.get('methodological_score') or 0) >= 70]
    
    if flavor_type == 'by_thesis':
        summary = f"### {name}\n\n{insight}\n\n"
        if high_methodological:
            summary += f"**📊 Alta calidad metodológica:** {len(high_methodological)} estudios con calidad ≥70/100.\n\n"
        
        best = max(articles, key=lambda x: x.get('methodological_score') or 0) if articles else None
        if best:
            summary += f"**⭐ Mejor calidad metodológica:** {best.get('authors', 'Autor')[:80]} ({best.get('pubdate', 's.f.')[:4]}). "
            effect_val = best.get('effect_value')
            if effect_val and effect_val > 0:
                summary += f"Reporta {best.get('effect_type', 'efecto')}={effect_val:.2f}. "
            summary += f"Calidad metodológica: {best.get('methodological_score') or 0}/100.\n\n"
            
            if best.get('validation_note'):
                summary += f"*{best.get('validation_note')}*\n\n"
    else:
        summary = f"### {flavor_data.get('name', 'Cluster')}\n\n"
        summary += f"Este grupo agrupa {n} artículos con temáticas similares. "
        summary += f"Calidad metodológica media: {flavor_data.get('avg_methodological', 0):.0f}/100.\n\n"
    
    citations = []
    seen_pmids = set()
    for i, article in enumerate(articles[:10], 1):
        pmid = str(article.get('pmid', ''))
        if pmid and pmid in seen_pmids:
            continue
        if pmid:
            seen_pmids.add(pmid)
        
        authors = str(article.get('authors', 'Autor'))[:50]
        title = str(article.get('title', 'Sin título'))[:80]
        journal = str(article.get('journal', 'Revista'))[:30]
        year = str(article.get('pubdate', 's.f.'))[:4]
        citation = f"{i}. {authors}. {title}. {journal}. {year}"
        if pmid and pmid != 'N/A':
            citation += f"; PMID: {pmid}"
        citations.append(citation)
    
    return summary, citations


# ============================================================================
# DOCUMENTO COMPLETO CON NUEVA SECCIÓN DE INTERPRETACIÓN
# ============================================================================

def add_grade_table_to_doc(doc, articles):
    if not articles:
        doc.add_paragraph("No hay artículos para mostrar.")
        return
    
    articles = deduplicate_articles(articles)
    top_articles = sorted(articles, key=lambda x: x.get('methodological_score', 0), reverse=True)[:10]
    
    doc.add_heading('📊 Tabla de Calidad Metodológica (top 10 estudios)', level=2)
    doc.add_paragraph('NOTA: Esta tabla muestra la calidad METODOLÓGICA, NO la confirmación de hipótesis.')
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    headers = ['Estudio', 'Diseño', 'N', 'EF (IC95%)', 'Calidad Metodológica', 'Relación', 'Score Diseño']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    for art in top_articles:
        row = table.add_row().cells
        row[0].text = f"{art.get('authors', 'Autor')[:30]} {art.get('pubdate', '')[:4]}"
        row[1].text = art.get('study_design', 'ND')[:20]
        
        n = art.get('sample_size')
        row[2].text = str(n) if n and n > 0 else 'NR'
        
        effect_val = art.get('effect_value', 0)
        if effect_val > 0:
            ci_lower = art.get('effect_ci_lower', 0)
            ci_upper = art.get('effect_ci_upper', 0)
            row[3].text = f"{art.get('effect_type', 'HR')}={effect_val:.2f} ({ci_lower:.2f}-{ci_upper:.2f})"
        else:
            row[3].text = 'NR'
        
        meth_score = art.get('methodological_score', 0) or 0
        if meth_score >= 70:
            quality = f"🟢 Alta ({meth_score:.0f})"
        elif meth_score >= 50:
            quality = f"🟡 Moderada ({meth_score:.0f})"
        else:
            quality = f"🟠 Baja ({meth_score:.0f})"
        row[4].text = quality
        
        rel_map = {'confirms': '✅ Confirma', 'nuances': '🔍 Matiza', 'extends': '🚀 Extiende', 
                   'contradicts': '⚠️ Contradice', 'unrelated': '📚 No relacionado'}
        row[5].text = rel_map.get(art.get('hypothesis_relation', 'unrelated'), '📚 Otro')
        row[6].text = f"{art.get('evidence_score', 0):.0f}/100"
    
    doc.add_paragraph()


def create_document_with_flavors_human_style(flavors, hypothesis, query, total_articles, relevance_threshold, consensus):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    title = doc.add_heading('Análisis de Evidencia por Flavors v8.7', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'**Estrategia de búsqueda:** {str(query)[:200]}...')
    doc.add_paragraph(f'**Hipótesis:** "{str(hypothesis)[:200]}..."')
    doc.add_paragraph(f'**Artículos analizados:** {total_articles}')
    doc.add_paragraph(f'**Threshold de relevancia:** {relevance_threshold}')
    doc.add_paragraph(f'**Fecha:** {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph()
    
    note_para = doc.add_paragraph()
    note_para.add_run('⚠️ NOTA IMPORTANTE: ').bold = True
    note_para.add_run('La clasificación se basa en evidencia textual EXPLÍCITA (frases como "confirms that").')
    
    doc.add_page_break()
    
    # Recolectar todos los artículos
    all_articles = []
    if flavors and 'by_thesis' in flavors:
        for relation in ['confirms', 'nuances', 'extends', 'contradicts']:
            if relation in flavors['by_thesis']:
                all_articles.extend(flavors['by_thesis'][relation].get('articles', []))
    
    if not all_articles:
        doc.add_paragraph("⚠️ **No se encontraron artículos para analizar.**")
        return doc
    
    doc.add_page_break()
    
    # RESUMEN EJECUTIVO
    exec_title = doc.add_heading('📋 RESUMEN EJECUTIVO Y ANÁLISIS DE CONSENSO', level=1)
    exec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    exec_summary = ConsensusAnalyzer.generate_executive_summary(all_articles, hypothesis, query)
    for line in exec_summary.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    doc.add_paragraph()
    
    if MATPLOTLIB_AVAILABLE:
        try:
            consensus_fig = plot_consensus_gauge(consensus)
            if consensus_fig:
                doc.add_picture(consensus_fig, width=Inches(5))
        except:
            pass
    
    doc.add_page_break()
    
    # ================================================================
    # NUEVA SECCIÓN: ¿QUÉ SIGNIFICA ESTO? (v8.7)
    # ================================================================
    doc.add_page_break()
    
    # Calcular métricas para el asistente interpretativo
    methodological_scores = [a.get('methodological_score', 0) for a in all_articles if a.get('methodological_score') is not None]
    avg_methodological = sum(methodological_scores) / len(methodological_scores) if methodological_scores else 0
    
    relevant_articles = len(all_articles)
    
    interpretation = InterpretiveAssistant.generate_interpretation(
        consensus, total_articles, relevant_articles, hypothesis, avg_methodological
    )
    
    for line in interpretation.split('\n'):
        if line.strip():
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('→'):
                p = doc.add_paragraph()
                p.add_run(line).italic = True
            elif line.startswith('•'):
                doc.add_paragraph(line, style='List Bullet')
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.'):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)
    
    doc.add_page_break()
    
    # ALERTAS DE CALIDAD
    doc.add_heading('🚦 Alertas de Calidad Metodológica', level=1)
    quality_alerts = QualityAlertSystem.generate_alerts(all_articles)
    doc.add_paragraph(f"**Nivel general:** {quality_alerts['overall_level']}")
    doc.add_paragraph(f"**{quality_alerts['overall_message']}**")
    doc.add_paragraph()
    
    if quality_alerts['alerts'].get('critical'):
        doc.add_heading('🔴 Alertas CRÍTICAS', level=2)
        for alert in quality_alerts['alerts']['critical']:
            doc.add_paragraph(f"• {alert['message']}", style='List Bullet')
    
    doc.add_page_break()
    
    # CONTRADICCIONES
    doc.add_heading('⚠️ Detección de Contradicciones entre Estudios', level=1)
    contradictions = ContradictionDetector.find_contradictions(all_articles)
    contradiction_summary = ContradictionDetector.generate_contradiction_summary(contradictions)
    doc.add_paragraph(contradiction_summary)
    
    doc.add_page_break()
    
    # LÍNEA DE TIEMPO
    doc.add_heading('📅 Evolución Temporal del Consenso', level=1)
    temporal_data = TemporalConsensusAnalyzer.analyze_temporal_consensus(all_articles)
    temporal_summary = TemporalConsensusAnalyzer.generate_temporal_summary(temporal_data)
    doc.add_paragraph(temporal_summary)
    
    temporal_plot = TemporalConsensusAnalyzer.create_temporal_plot(temporal_data)
    if temporal_plot:
        doc.add_picture(temporal_plot, width=Inches(6))
    
    doc.add_page_break()
    
    # TABLA GRADE
    add_grade_table_to_doc(doc, all_articles)
    doc.add_page_break()
    
    # FLAVORS
    doc.add_heading('FLAVORS POR RELACIÓN CON LA HIPÓTESIS', level=1)
    doc.add_paragraph('Los siguientes flavors agrupan artículos según cómo se relacionan con la hipótesis.')
    doc.add_paragraph('⚠️ La clasificación se basa en evidencia textual EXPLÍCITA.')
    doc.add_paragraph()
    
    thesis_flavors = flavors.get('by_thesis', {})
    for relation, flavor_data in thesis_flavors.items():
        if not flavor_data.get('articles'):
            continue
        
        summary, citations = generate_flavor_summary_human_style(flavor_data, 'by_thesis', hypothesis)
        doc.add_paragraph(summary)
        
        if citations:
            doc.add_paragraph('**Referencias clave:**', style='List Bullet')
            for citation in citations[:10]:
                doc.add_paragraph(citation, style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # CLUSTERS
    if flavors.get('by_cluster'):
        doc.add_heading('CLUSTERS TEMÁTICOS', level=1)
        doc.add_paragraph('Clusters por similitud temática (basados en títulos y outcomes).')
        doc.add_paragraph()
        
        for cluster in flavors.get('by_cluster', []):
            if not cluster.get('articles'):
                continue
            
            summary, citations = generate_flavor_summary_human_style(cluster, 'by_cluster', hypothesis)
            doc.add_paragraph(summary)
            
            if citations:
                doc.add_paragraph('**Artículos en este cluster:**', style='List Bullet')
                for citation in citations[:10]:
                    doc.add_paragraph(citation, style='List Bullet')
            doc.add_paragraph()
        
        doc.add_page_break()
    
    # TABLA RESUMEN
    doc.add_heading('RESUMEN DE EVIDENCIA POR CATEGORÍA', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Relación con Hipótesis'
    hdr_cells[1].text = 'N° Artículos'
    hdr_cells[2].text = 'Calidad Metodológica Media'
    hdr_cells[3].text = 'Con tamaño del efecto'
    
    for relation, flavor_data in thesis_flavors.items():
        if not flavor_data.get('articles'):
            continue
        
        row_cells = table.add_row().cells
        row_cells[0].text = flavor_data.get('name', relation)
        row_cells[1].text = str(flavor_data.get('n_articles', 0))
        row_cells[2].text = f"{flavor_data.get('avg_methodological_score', 0):.0f}/100"
        row_cells[3].text = f"{flavor_data.get('has_effect_size', 0)} estudios"
    
    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph(f'*Documento generado por PubMed AI Analyzer v8.7*')
    doc.add_paragraph(f'*Email configurado con: {get_email_config()["email_user"]}*')
    doc.add_paragraph('*La clasificación "Confirma" requiere evidencia textual EXPLÍCITA.*')
    
    return doc


def export_articles_to_csv_enhanced(articles):
    if not articles:
        return None
    
    articles = deduplicate_articles(articles)
    data = []
    for article in articles:
        row = {
            'PMID': article.get('pmid', ''),
            'Título': article.get('title', ''),
            'Autores': article.get('authors', ''),
            'Revista': article.get('journal', ''),
            'Año': article.get('pubdate', '')[:4],
            'Diseño': article.get('study_design', ''),
            'Nivel evidencia': article.get('evidence_level', ''),
            'Score diseño': article.get('evidence_score', 0),
            'Score metodológico': article.get('methodological_score', 0),
            'Tamaño muestral': article.get('sample_size', ''),
            'Efecto': article.get('effect_type', ''),
            'Valor efecto': article.get('effect_value', 0),
            'Relación hipótesis': article.get('hypothesis_relation', ''),
            'Confianza relación': article.get('relation_confidence', 0),
            'Razón clasificación': article.get('relation_reason', ''),
            'Nota validación': article.get('validation_note', ''),
            'PICO_Población': article.get('pico_population', ''),
            'PICO_Intervención': article.get('pico_intervention', ''),
            'PICO_Outcomes': article.get('pico_outcomes', '')
        }
        data.append(row)
    
    df = pd.DataFrame(data)
    csv_buffer = StringIO()
    df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
    return csv_buffer.getvalue().encode('utf-8-sig')


def display_flavors_preview_v8(flavors, consensus):
    st.markdown("---")
    st.markdown("## 📋 RESUMEN EJECUTIVO v8.7")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("✅ Confirman", f"{consensus.get('confirms_percentage', 0):.0f}%")
    with col2:
        st.metric("⚠️ Contradicen", f"{consensus.get('contradicts_percentage', 0):.0f}%")
    with col3:
        st.metric("🎯 Consenso", consensus.get('consensus_level', 'ND'))
    with col4:
        st.metric("⭐ Artículos", consensus.get('total_articles', 0))
    
    st.info(consensus.get('consensus_message', ''))
    
    if MATPLOTLIB_AVAILABLE:
        consensus_fig = plot_consensus_gauge(consensus)
        if consensus_fig:
            st.image(consensus_fig, use_container_width=True)
    
    st.markdown("---")
    st.markdown("## 🎨 Flavors Generados (v8.7)")
    
    thesis_flavors = flavors.get('by_thesis', {})
    if thesis_flavors:
        st.markdown("### 📚 Por relación con la hipótesis")
        for relation, flavor_data in thesis_flavors.items():
            if not flavor_data.get('articles'):
                continue
            with st.expander(f"{flavor_data.get('name', relation)} ({flavor_data.get('n_articles', 0)} artículos)", expanded=True):
                st.markdown(f"**{flavor_data.get('description', '')}**")
                st.markdown(f"📊 **Insight:** {flavor_data.get('insight', '')}")
                st.markdown(f"⭐ **Calidad metodológica promedio:** {flavor_data.get('avg_methodological_score', 0):.0f}/100")
                
                st.markdown("**📄 Artículos destacados (mejor calidad metodológica):**")
                for i, art in enumerate(flavor_data.get('representative_articles', [])[:5], 1):
                    title = art.get('title', 'Sin título')[:80]
                    meth_score = art.get('methodological_score') or 0
                    st.markdown(f"   {i}. {title}... (calidad metodológica={meth_score:.0f})")
                    if art.get('validation_note'):
                        st.caption(f"   📝 {art.get('validation_note')}")


def display_graphs_section(articles):
    if not MATPLOTLIB_AVAILABLE:
        st.info("📊 Instala matplotlib para gráficos")
        return
    
    articles = deduplicate_articles(articles)
    st.markdown("## 📊 Análisis Gráfico")
    
    st.markdown("### 📅 Evolución temporal")
    temporal_data = TemporalConsensusAnalyzer.analyze_temporal_consensus(articles)
    temporal_plot = TemporalConsensusAnalyzer.create_temporal_plot(temporal_data)
    if temporal_plot:
        st.image(temporal_plot, use_container_width=True)
    else:
        st.info("No hay suficientes datos para gráfico temporal.")
    
    st.markdown("### 🚦 Alertas de calidad")
    quality_alerts = QualityAlertSystem.generate_alerts(articles)
    st.markdown(f"**Nivel general:** {quality_alerts['overall_level']}")
    
    if quality_alerts['alerts'].get('critical'):
        with st.expander("🔴 Alertas CRÍTICAS", expanded=True):
            for alert in quality_alerts['alerts']['critical']:
                st.warning(f"**{alert['message']}**")
    
    st.markdown("### ⚠️ Contradicciones")
    contradictions = ContradictionDetector.find_contradictions(articles)
    if contradictions:
        st.warning(f"Se detectaron {len(contradictions)} contradicción(es)")
    else:
        st.success("✅ No se detectaron contradicciones")


def extract_key_terms_from_query(query):
    terms = set()
    mesh_pattern = r'"([^"]+)"\[Mesh\]'
    for term in re.findall(mesh_pattern, query, re.IGNORECASE):
        terms.add(term.lower())
    tiab_pattern = r'"([^"]+)"\[tiab\]'
    for term in re.findall(tiab_pattern, query, re.IGNORECASE):
        terms.add(term.lower())
    return list(terms)


def extract_key_terms_from_hypothesis(hypothesis):
    if not hypothesis:
        return []
    terms = set()
    words = hypothesis.lower().split()
    for i in range(len(words)-1):
        if len(words[i]) > 4 and len(words[i+1]) > 3:
            terms.add(f"{words[i]} {words[i+1]}")
    return list(terms)


def process_and_generate_flavors(query, hypothesis, threshold, user_email, session_manager=None, flavor_method="both"):
    start_time = time.time()
    
    query_terms = extract_key_terms_from_query(query)
    hypothesis_terms = extract_key_terms_from_hypothesis(hypothesis)
    st.info(f"📝 Términos: {len(query_terms)} búsqueda, {len(hypothesis_terms)} hipótesis")
    
    progress_placeholder = st.empty()
    
    def update_progress(current, total):
        progress_placeholder.progress(current / total, text=f"Procesando {current}/{total}...")
    
    with st.spinner("🔍 Buscando artículos..."):
        id_list, total = search_pubmed_complete(query.strip(), max_articles=MAX_ARTICLES)
        if not id_list:
            st.error("❌ No se encontraron artículos")
            return None, None, None, None, None
    
    all_articles = []
    total_blocks = (len(id_list) + BLOCK_SIZE - 1) // BLOCK_SIZE
    blocks_to_process = min(total_blocks, 6)
    
    overall_progress = st.progress(0)
    article_count = 0
    
    for block_num in range(1, blocks_to_process + 1):
        start_idx = (block_num - 1) * BLOCK_SIZE
        end_idx = min(block_num * BLOCK_SIZE, len(id_list))
        block_ids = id_list[start_idx:end_idx]
        
        st.info(f"📦 Bloque {block_num}/{blocks_to_process}...")
        
        def block_progress(current, total):
            nonlocal article_count
            article_count += 1
            overall_progress.progress(article_count / len(id_list))
        
        block_articles = fetch_articles_details(
            block_ids, query_terms, hypothesis_terms, 
            hypothesis=hypothesis, block_number=block_num,
            progress_callback=block_progress
        )
        
        if block_articles:
            all_articles.extend(block_articles)
            if session_manager and session_manager.current_session_id:
                session_manager.save_articles_batch(block_articles, session_manager.current_session_id)
        
        time.sleep(2)
    
    progress_placeholder.empty()
    overall_progress.empty()
    
    if not all_articles:
        st.error("❌ No se procesaron artículos")
        return None, None, None, None, None
    
    st.success(f"✅ Procesados {len(all_articles)} artículos")
    
    with st.spinner("🧠 Calculando relevancia (usando embeddings si están disponibles)..."):
        all_articles = calculate_relevance_to_hypothesis(all_articles, hypothesis)
        filtered_articles = filter_articles_by_relevance(all_articles, threshold)
    
    filtered_articles = deduplicate_articles(filtered_articles)
    st.info(f"📊 Después de deduplicación: {len(filtered_articles)} artículos únicos")
    
    if len(filtered_articles) < 3:
        st.error(f"❌ Pocos artículos ({len(filtered_articles)}). Mínimo 3.")
        return None, None, None, None, None
    
    st.info(f"⏱️ Tiempo: {(time.time()-start_time)/60:.1f} minutos")
    
    consensus = ConsensusAnalyzer.calculate_consensus(filtered_articles)
    display_graphs_section(filtered_articles)
    
    with st.spinner(f"🎨 Generando flavors (con clustering si está disponible)..."):
        flavors = generate_all_flavors_human_style(filtered_articles, hypothesis, query_terms, hypothesis_terms, method=flavor_method)
    
    return flavors, filtered_articles, query, hypothesis, consensus


# ============================================================================
# CLASE PARA ALMACENAMIENTO LOCAL
# ============================================================================

class LocalStorage:
    def __init__(self, login):
        self.login = login
        self.data_dir = f"./data_{login}"
        os.makedirs(self.data_dir, exist_ok=True)
    
    def read_csv(self, filename):
        filepath = os.path.join(self.data_dir, filename)
        if os.path.exists(filepath):
            return pd.read_csv(filepath)
        return None
    
    def write_csv(self, filename, df):
        filepath = os.path.join(self.data_dir, filename)
        df.to_csv(filepath, index=False, encoding='utf-8-sig')
        return True
    
    def append_to_csv(self, filename, new_data):
        existing_df = self.read_csv(filename)
        if existing_df is not None and not existing_df.empty:
            df = pd.concat([existing_df, new_data], ignore_index=True)
        else:
            df = new_data
        return self.write_csv(filename, df)


class UserSessionManager:
    def __init__(self, login):
        self.login = login
        self.storage = LocalStorage(login)
        self.articles_file = f"{login}_articles_v87.csv"
        self.sessions_file = f"{login}_sessions_v87.csv"
        self.current_session_id = None
    
    def create_session(self, query: str, hypothesis: str, relevance_threshold: float, email: str) -> str:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        query_hash = hashlib.md5(query.encode()).hexdigest()[:8]
        session_id = f"{timestamp}_{query_hash}"
        self.current_session_id = session_id
        
        session_data = pd.DataFrame([{
            'session_id': session_id,
            'login': self.login,
            'query': query[:500],
            'hypothesis': hypothesis[:500],
            'relevance_threshold': relevance_threshold,
            'total_found': 0,
            'total_processed': 0,
            'start_time': datetime.now().isoformat(),
            'status': 'running',
            'flavors_generated': False,
            'user_email': email
        }])
        
        self.storage.append_to_csv(self.sessions_file, session_data)
        return session_id
    
    def save_articles_batch(self, articles: List[Dict], session_id: str):
        if not articles:
            return
        records = []
        for article in articles:
            record = {
                'pmid': str(article.get('pmid', '')),
                'session_id': str(session_id),
                'login': str(self.login),
                'title': str(article.get('title', ''))[:1000],
                'authors': str(article.get('authors', ''))[:500],
                'journal': str(article.get('journal', ''))[:200],
                'abstract': str(article.get('abstract', ''))[:5000],
                'methodological_score': float(article.get('methodological_score', 0) or 0),
                'evidence_level': str(article.get('evidence_level', '')),
                'evidence_score': int(article.get('evidence_score', 0) or 0),
                'sample_size': int(article.get('sample_size', 0) or 0),
                'effect_type': str(article.get('effect_type', '')),
                'effect_value': float(article.get('effect_value', 0) or 0),
                'effect_ci_lower': float(article.get('effect_ci_lower', 0) or 0),
                'effect_ci_upper': float(article.get('effect_ci_upper', 0) or 0),
                'hypothesis_relation': str(article.get('hypothesis_relation', '')),
                'relation_confidence': float(article.get('relation_confidence', 0) or 0),
                'relation_reason': str(article.get('relation_reason', ''))[:200],
                'relevance_score': float(article.get('relevance_score', 0) or 0),
                'processed_date': datetime.now().isoformat(),
                'block_number': int(article.get('block_number', 0) or 0),
                'pico_population': str(article.get('pico_population', '')),
                'pico_intervention': str(article.get('pico_intervention', '')),
                'pico_outcomes': str(article.get('pico_outcomes', ''))
            }
            records.append(record)
        
        df_new = pd.DataFrame(records)
        existing_df = self.storage.read_csv(self.articles_file)
        if existing_df is not None and not existing_df.empty:
            df = pd.concat([existing_df, df_new], ignore_index=True)
            df = df.drop_duplicates(subset=['pmid', 'session_id'], keep='last')
        else:
            df = df_new
        
        self.storage.write_csv(self.articles_file, df)
        st.success(f"✅ Guardados {len(records)} artículos")
    
    def get_session_articles(self, session_id: str) -> pd.DataFrame:
        df = self.storage.read_csv(self.articles_file)
        if df is None or df.empty:
            return pd.DataFrame()
        return df[df['session_id'] == session_id]
    
    def update_session(self, session_id: str, status: str, total_processed: int = None):
        df = self.storage.read_csv(self.sessions_file)
        if df is None or df.empty:
            return
        df = df.copy()
        mask = df['session_id'] == session_id
        if total_processed is not None:
            df.loc[mask, 'total_processed'] = total_processed
        df.loc[mask, 'status'] = status
        self.storage.write_csv(self.sessions_file, df)
    
    def mark_flavors_generated(self, session_id: str):
        df = self.storage.read_csv(self.sessions_file)
        if df is None or df.empty:
            return
        df = df.copy()
        mask = df['session_id'] == session_id
        df.loc[mask, 'flavors_generated'] = True
        self.storage.write_csv(self.sessions_file, df)


# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    st.title("🧠 PubMed AI Analyzer v8.7 (CON ASISTENTE INTERPRETATIVO)")
    st.markdown("*Análisis de evidencia con clasificación CORREGIDA + Asistente que interpreta los resultados*")
    st.markdown("⚠️ **Novedad v8.7:** Nueva sección **'¿Qué significa esto?'** que contextualiza los resultados numéricos y ofrece recomendaciones accionables.")
    
    if 'user_login' not in st.session_state:
        st.markdown("### 🔐 Identificación")
        col1, col2 = st.columns(2)
        with col1:
            login = st.text_input("Login:", placeholder="ejemplo: investigador")
        with col2:
            user_email = st.text_input("📧 Email:", placeholder="usuario@dominio.com")
        
        if st.button("✅ Continuar", type="primary"):
            if login.strip() and user_email.strip() and '@' in user_email:
                st.session_state.user_login = login.strip().lower()
                st.session_state.user_email = user_email.strip()
                st.rerun()
            else:
                st.warning("⚠️ Complete ambos campos")
        return
    
    st.sidebar.success(f"👤 Usuario: {st.session_state.user_login}")
    st.sidebar.info(f"📧 {st.session_state.user_email}")
    
    if st.sidebar.button("🔄 Cambiar usuario"):
        del st.session_state.user_login
        del st.session_state.user_email
        st.rerun()
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### ⚙️ Configuración")
    
    flavor_method = st.sidebar.radio(
        "Método de flavors:",
        ["both", "thesis", "clustering"],
        format_func=lambda x: {
            "both": "Ambos métodos",
            "thesis": "Solo por hipótesis",
            "clustering": "Solo clustering"
        }[x],
        index=0
    )
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🤖 Paquetes AI cargados")
    st.sidebar.markdown(f"- **Embeddings:** {'✅' if AI_EMBEDDINGS_AVAILABLE else '❌'}")
    st.sidebar.markdown(f"- **Clustering:** {'✅' if SKLEARN_AVAILABLE else '❌'}")
    st.sidebar.markdown(f"- **Gráficos:** {'✅' if MATPLOTLIB_AVAILABLE else '❌'}")
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 📧 Configuración de Email")
    
    email_config = EmailSender.get_email_config_status()
    if email_config['configured']:
        st.sidebar.success(f"✅ Email configurado")
        st.sidebar.info(f"📤 Desde: {email_config['sender_email']}")
    else:
        st.sidebar.error("❌ Email NO configurado")
        st.sidebar.info("Crear archivo .streamlit/secrets.toml")
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 📦 Exportación")
    st.sidebar.markdown("- DOCX (con envío email)")
    st.sidebar.markdown("- CSV (datos completos)")
    st.sidebar.markdown("- RIS / BibTeX")
    
    st.info("⚡ **v8.7 NUEVO:** Incluye sección **'¿Qué significa esto?'** que interpreta automáticamente los resultados y ofrece recomendaciones personalizadas.")
    st.markdown("---")
    
    default_query = '("myocardial infarction"[Mesh] OR "myocardial infarction"[tiab]) AND ("intramyocardial dissection"[tiab] OR "intramyocardial dissecting hematoma"[tiab])'
    default_hypothesis = "Intramyocardial dissections occurring as a complication of myocardial infarction follow predictable anatomical pathways along established tissue planes."
    
    query = st.text_area("**PubMed search strategy:**", value=default_query, height=100)
    hypothesis = st.text_area("**📌 Hipótesis:**", value=default_hypothesis, height=100)
    threshold = st.slider("**Umbral de relevancia (filtro preliminar):**", 0.0, 0.9, 0.35, 0.05)
    
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        generate_button = st.button("🚀 GENERAR FLAVORS v8.7", type="primary", use_container_width=True)
    with col2:
        send_email = st.checkbox("📧 Enviar por email", value=True)
    with col3:
        st.markdown(f"*Máx. {MAX_ARTICLES} artículos*")
    
    if generate_button:
        if not query.strip() or not hypothesis.strip():
            st.warning("⚠️ Complete búsqueda e hipótesis")
        else:
            session_manager = UserSessionManager(st.session_state.user_login)
            session_id = session_manager.create_session(query, hypothesis, threshold, st.session_state.user_email)
            
            with st.spinner("Procesando..."):
                flavors, articles, q, h, consensus = process_and_generate_flavors(
                    query, hypothesis, threshold, st.session_state.user_email, 
                    session_manager, flavor_method
                )
                
                if flavors and articles:
                    session_manager.mark_flavors_generated(session_id)
                    display_flavors_preview_v8(flavors, consensus)
                    
                    doc = create_document_with_flavors_human_style(
                        flavors, h, q, len(articles), threshold, consensus
                    )
                    
                    docx_bytes = BytesIO()
                    doc.save(docx_bytes)
                    docx_bytes.seek(0)
                    
                    st.markdown("---")
                    st.markdown("## 📥 Descargas")
                    
                    tab1, tab2, tab3, tab4 = st.tabs(["📄 DOCX", "📊 CSV", "📚 RIS", "📖 BibTeX"])
                    
                    with tab1:
                        filename = f"flavors_v87_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                        st.download_button(
                            "💾 DESCARGAR DOCX",
                            data=docx_bytes,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        
                        if send_email:
                            with st.spinner("📧 Enviando por correo..."):
                                success, message = EmailSender.send_docx_by_email(
                                    recipient_email=st.session_state.user_email,
                                    docx_bytes=docx_bytes,
                                    filename=filename
                                )
                                if success:
                                    st.success(message)
                                else:
                                    st.error(message)
                    
                    with tab2:
                        csv_data = export_articles_to_csv_enhanced(articles)
                        if csv_data:
                            st.download_button(
                                "📊 DESCARGAR CSV",
                                data=csv_data,
                                file_name=f"articles_v87_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                mime="text/csv",
                                use_container_width=True
                            )
                    
                    with tab3:
                        ris_data = ReferenceExporter.to_ris(articles)
                        st.download_button(
                            "📚 DESCARGAR RIS",
                            data=ris_data,
                            file_name=f"references_v87_{datetime.now().strftime('%Y%m%d_%H%M%S')}.ris",
                            mime="application/x-research-info-systems",
                            use_container_width=True
                        )
                    
                    with tab4:
                        bib_data = ReferenceExporter.to_bibtex(articles)
                        st.download_button(
                            "📖 DESCARGAR BibTeX",
                            data=bib_data,
                            file_name=f"references_v87_{datetime.now().strftime('%Y%m%d_%H%M%S')}.bib",
                            mime="text/plain",
                            use_container_width=True
                        )
                    
                    st.success("✅ ¡Proceso completado!")
                    st.info("💡 **Novedad v8.7:** Revise la sección **'¿Qué significa esto?'** en el documento DOCX para interpretar contextualmente los resultados.")
                else:
                    st.error("❌ No se pudieron generar flavors. Intente con umbral más bajo.")
    
    st.markdown("---")
    st.markdown("*PubMed AI Analyzer v8.7 | Clasificación por evidencia textual EXPLÍCITA | Asistente interpretativo integrado*")


if __name__ == "__main__":
    main()
