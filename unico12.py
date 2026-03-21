import streamlit as st
import pandas as pd
from datetime import datetime
import time
import requests
from xml.etree import ElementTree
import re
from collections import Counter
import numpy as np
import math
import uuid
import string
import json
import hashlib
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# CONFIGURACIÓN PARA STREAMLIT CLOUD
# ============================================================================

# Detectar si estamos en Streamlit Cloud
IN_STREAMLIT_CLOUD = os.getenv('STREAMLIT_RUNTIME_ENV') == 'cloud' or os.getenv('STREAMLIT_SHARING_MODE') is not None

# Configurar BioBERT (funciona perfectamente en la nube)
AI_EMBEDDINGS_AVAILABLE = False
BIOMED_EMBEDDER = None

try:
    from sentence_transformers import SentenceTransformer
    import torch
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import KMeans, HDBSCAN
    from sklearn.feature_extraction.text import TfidfVectorizer
    
    # Cargar modelo (se descargará automáticamente la primera vez)
    with st.spinner("🔄 Cargando BioBERT (primera vez puede tomar ~2 minutos)..."):
        BIOMED_EMBEDDER = SentenceTransformer('pritamdeka/S-Biomed-Roberta-snli-multinli-stsb')
    AI_EMBEDDINGS_AVAILABLE = True
    print("✅ BioBERT embeddings disponible")
except ImportError:
    print("⚠️ sentence-transformers no instalado")
except Exception as e:
    print(f"⚠️ Error cargando BioBERT: {e}")

# Configurar spaCy (deshabilitado en Streamlit Cloud)
if IN_STREAMLIT_CLOUD:
    DEPENDENCY_PARSING_AVAILABLE = False
    NLP_MODEL = None
    print("ℹ️ spaCy deshabilitado en Streamlit Cloud")
else:
    try:
        import spacy
        try:
            NLP_MODEL = spacy.load("en_core_web_sm")
            DEPENDENCY_PARSING_AVAILABLE = True
            print("✅ spaCy cargado")
        except:
            DEPENDENCY_PARSING_AVAILABLE = False
            NLP_MODEL = None
    except ImportError:
        DEPENDENCY_PARSING_AVAILABLE = False
        NLP_MODEL = None

# Configuración de la página
st.set_page_config(
    page_title="PubMed AI Analyzer - Generador de Flavors",
    page_icon="🧠",
    layout="centered"
)

# ============================================================================
# FUNCIONES DE UTILIDAD PARA PubMed
# ============================================================================

@st.cache_data(ttl=3600, show_spinner=False)
def get_abstract(pmid):
    """Obtiene el abstract de un artículo de PubMed con cache"""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    fetch_url = f"{base_url}efetch.fcgi"
    
    params = {
        "db": "pubmed",
        "id": pmid,
        "retmode": "xml",
        "rettype": "abstract"
    }
    
    try:
        response = requests.get(fetch_url, params=params, timeout=10)
        response.raise_for_status()
        
        root = ElementTree.fromstring(response.content)
        
        abstract_texts = []
        for abstract in root.findall(".//AbstractText"):
            if abstract.text:
                abstract_texts.append(abstract.text)
        
        return " ".join(abstract_texts) if abstract_texts else None
        
    except Exception as e:
        return None

def preprocess_text(text):
    """Preprocesa el texto para análisis NLP"""
    if not text:
        return ""
    
    text = text.lower()
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\d+', '', text)
    text = ' '.join(text.split())
    
    return text

def extract_medical_entities(text):
    """Extrae entidades médicas importantes del texto"""
    if not text:
        return []
    
    text_lower = text.lower()
    entities = []
    
    medical_patterns = {
        'medications': r'\b(ticagrelor|clopidogrel|aspirin|prasugrel|heparin|enoxaparin|bivalirudin|warfarin|rivaroxaban|apixaban|dabigatran|metformin|insulin|statins|atorvastatin|rosuvastatin|lisinopril|enalapril|losartan|valsartan|amlodipine|nifedipine|furosemide|hydrochlorothiazide|omeprazole|pantoprazole|acetaminophen|ibuprofen|naproxen|prednisone|hydrocortisone|levothyroxine|sertraline|fluoxetine)\b',
        'conditions': r'\b(myocardial infarction|heart attack|acute coronary syndrome|unstable angina|nstemi|stemi|ischemia|angina|coronary disease|kounis syndrome|reperfusion injury|hypertension|high blood pressure|diabetes|diabetes mellitus|hyperlipidemia|high cholesterol|heart failure|cardiomyopathy|atrial fibrillation|stroke|cva|tia|asthma|copd|pneumonia|covid-19|arthritis|osteoarthritis|rheumatoid arthritis|depression|anxiety|schizophrenia|bipolar|epilepsy|seizure|migraine|alzheimer|parkinson|dementia|cancer|malignancy|tumor|carcinoma)\b',
        'procedures': r'\b(pci|percutaneous coronary intervention|cabg|bypass|stent|angioplasty|thrombectomy|fibrinolysis|thrombolysis|angiography|echocardiogram|ekg|ecg|mri|ct scan|x-ray|ultrasound|endoscopy|colonoscopy|biopsy|surgery|transplant|dialysis|ventilation|intubation|resuscitation)\b',
        'study_types': r'\b(randomized|rct|randomised|trial|cohort|prospective|retrospective|observational|clinical trial|meta-analysis|systematic review|case-control|cross-sectional|longitudinal|registry|real-world|real world)\b',
        'outcomes': r'\b(mortality|death|survival|bleeding|hemorrhage|stroke|reinfarction|mace|major adverse cardiac events|efficacy|safety|tolerability|effectiveness|complication|adverse event|side effect|toxicity|risk|rate|incidence|prevalence|occurrence|outcome|endpoint|result|finding|response|remission|recovery|improvement|worsening|pain|symptom|function|quality of life|qol)\b'
    }
    
    for entity_type, pattern in medical_patterns.items():
        matches = re.findall(pattern, text_lower)
        entities.extend([(match, entity_type) for match in matches])
    
    return entities

def extract_pico_elements(text):
    """Extrae elementos PICO del texto"""
    if not text:
        return {}
    
    text_lower = text.lower()
    pico = {}
    
    population_patterns = [
        r'patients? with ([^,.]+)',
        r'subjects? with ([^,.]+)',
        r'adults? (?:aged|with) ([^,.]+)',
        r'population:? ([^,.]+)'
    ]
    
    for pattern in population_patterns:
        match = re.search(pattern, text_lower)
        if match:
            pico['population'] = match.group(1).strip()
            break
    
    intervention_patterns = [
        r'(?:treated with|received|administered) ([^,.]+)',
        r'intervention:? ([^,.]+)',
        r'study drug:? ([^,.]+)'
    ]
    
    for pattern in intervention_patterns:
        match = re.search(pattern, text_lower)
        if match:
            pico['intervention'] = match.group(1).strip()
            break
    
    comparison_patterns = [
        r'compared with ([^,.]+)',
        r'versus ([^,.]+)',
        r'vs\.? ([^,.]+)'
    ]
    
    for pattern in comparison_patterns:
        match = re.search(pattern, text_lower)
        if match:
            pico['comparison'] = match.group(1).strip()
            break
    
    return pico

def extract_numeric_results(text):
    """Extrae resultados numéricos importantes"""
    if not text:
        return []
    
    results = []
    
    hr_pattern = r'(?:HR|hazard ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(hr_pattern, text, re.IGNORECASE):
        results.append({'type': 'HR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    
    rr_pattern = r'(?:RR|risk ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(rr_pattern, text, re.IGNORECASE):
        results.append({'type': 'RR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    
    or_pattern = r'(?:OR|odds ratio)[\s:=]+([0-9]+\.[0-9]+)\s*\(([0-9]+\.[0-9]+)[-\s]+([0-9]+\.[0-9]+)\)'
    for match in re.findall(or_pattern, text, re.IGNORECASE):
        results.append({'type': 'OR', 'value': float(match[0]), 'ci_lower': float(match[1]), 'ci_upper': float(match[2])})
    
    p_pattern = r'[Pp]\s*[=<]\s*([0-9]+\.[0-9]+)'
    for match in re.findall(p_pattern, text):
        results.append({'type': 'p-value', 'value': float(match)})
    
    return results

def extract_all_outcomes(text):
    """Extrae todos los outcomes mencionados"""
    if not text:
        return []
    
    text_lower = text.lower()
    outcomes = []
    
    outcome_patterns = [
        r'\b(mortality|death|survival|fatality)\b',
        r'\b(bleeding|hemorrhage|haemorrhage|blood loss)\b',
        r'\b(stroke|cva|tia|ischaemic stroke|hemorrhagic stroke)\b',
        r'\b(reinfarction|mi|heart attack|myocardial infarction)\b',
        r'\b(mace|major adverse cardiac events)\b',
        r'\b(efficacy|effectiveness|safety|tolerability)\b',
        r'\b(complication|adverse event|side effect|toxicity|ae)\b',
        r'\b(risk|rate|incidence|prevalence|occurrence|frequency)\b',
        r'\b(outcome|endpoint|result|finding|conclusion)\b',
        r'\b(response|remission|recovery|improvement|worsening|deterioration)\b'
    ]
    
    for pattern in outcome_patterns:
        matches = re.findall(pattern, text_lower)
        outcomes.extend(matches)
    
    return list(set(outcomes))

def analyze_sentiment_by_outcome(text):
    """Analiza el sentimiento para cada outcome"""
    if not text:
        return {}
    
    text_lower = text.lower()
    outcomes_sentiment = {}
    all_outcomes = extract_all_outcomes(text)
    
    negation_words = ['no', 'not', 'without', 'absence', 'lack', 'failed']
    
    for outcome in all_outcomes:
        outcome_index = text_lower.find(outcome)
        if outcome_index == -1:
            continue
            
        start = max(0, outcome_index - 200)
        end = min(len(text_lower), outcome_index + 200)
        context = text_lower[start:end]
        
        positive_terms = ['reduced', 'decreased', 'lower', 'improved', 'better', 'benefit', 'protective', 'effective']
        negative_terms = ['increased', 'higher', 'worse', 'elevated', 'risk', 'adverse', 'harm', 'detrimental']
        
        sentiment_score = 0
        
        for term in positive_terms:
            if term in context:
                term_pos = context.find(term)
                prev_words = context[max(0, term_pos-30):term_pos]
                if any(neg in prev_words for neg in negation_words):
                    sentiment_score -= 1
                else:
                    sentiment_score += 1
        
        for term in negative_terms:
            if term in context:
                term_pos = context.find(term)
                prev_words = context[max(0, term_pos-30):term_pos]
                if any(neg in prev_words for neg in negation_words):
                    sentiment_score += 1
                else:
                    sentiment_score -= 1
        
        if sentiment_score > 0:
            outcomes_sentiment[outcome] = 'BENEFICIO'
        elif sentiment_score < 0:
            outcomes_sentiment[outcome] = 'RIESGO'
        else:
            outcomes_sentiment[outcome] = 'SIN EFECTO SIGNIFICATIVO'
    
    return outcomes_sentiment

def semantic_similarity_analysis(text, concepts):
    """Análisis de similitud semántica usando embeddings biomédicos"""
    if not AI_EMBEDDINGS_AVAILABLE or not text or not BIOMED_EMBEDDER:
        return {}
    
    try:
        sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 20][:20]
        if not sentences:
            return {}
        
        sentence_embeddings = BIOMED_EMBEDDER.encode(sentences)
        concept_embeddings = BIOMED_EMBEDDER.encode(concepts)
        similarity_matrix = cosine_similarity(sentence_embeddings, concept_embeddings)
        
        concept_analysis = {}
        for i, concept in enumerate(concepts):
            max_sim_idx = np.argmax(similarity_matrix[:, i])
            concept_analysis[concept] = {
                'similarity': float(similarity_matrix[max_sim_idx, i]),
                'relevant_sentence': sentences[max_sim_idx][:200]
            }
        
        return concept_analysis
    except Exception as e:
        return {'error': str(e)}

def assess_risk_of_bias(text, study_type):
    """Evalúa riesgo de sesgo según dominios Cochrane"""
    if not text:
        return {}
    
    text_lower = text.lower()
    rob_assessment = {}
    
    domains = {
        'selection_bias': {'low': ['random sequence generation', 'random allocation'], 'high': ['quasi-random', 'alternate allocation']},
        'performance_bias': {'low': ['double-blind', 'double blind', 'masked'], 'high': ['open-label', 'unblinded']},
        'detection_bias': {'low': ['blinded outcome assessment', 'independent adjudication'], 'high': ['unblinded assessment']},
        'attrition_bias': {'low': ['low dropout', 'complete follow-up'], 'high': ['high dropout', 'loss to follow-up']},
        'reporting_bias': {'low': ['registered', 'protocol published'], 'high': ['selective reporting']}
    }
    
    for domain, criteria in domains.items():
        if any(term in text_lower for term in criteria.get('low', [])):
            rob_assessment[domain] = 'LOW'
        elif any(term in text_lower for term in criteria.get('high', [])):
            rob_assessment[domain] = 'HIGH'
        else:
            rob_assessment[domain] = 'UNCLEAR'
    
    if 'RCT' not in study_type:
        rob_assessment['overall'] = 'HIGH (non-RCT)'
    else:
        risk_counts = Counter(rob_assessment.values())
        if risk_counts.get('HIGH', 0) >= 3:
            rob_assessment['overall'] = 'HIGH'
        elif risk_counts.get('HIGH', 0) >= 1 or risk_counts.get('UNCLEAR', 0) >= 3:
            rob_assessment['overall'] = 'SOME CONCERNS'
        else:
            rob_assessment['overall'] = 'LOW'
    
    return rob_assessment

def enhanced_quality_scoring(ai_analysis, full_text):
    """Sistema de puntuación de calidad mejorado"""
    score = 0
    factors = []
    text_lower = full_text.lower() if full_text else ""
    
    study_type_weights = {'Meta-análisis': 2.5, 'RCT': 2.0, 'Cohorte': 1.6, 'Caso-control': 1.4, 'Observacional': 1.2}
    
    study_types = ai_analysis.get('study_types', [])
    for study_type in study_types:
        if study_type in study_type_weights:
            score += study_type_weights[study_type] * 20
            factors.append(f"Tipo estudio: {study_type}")
    
    sample_patterns = [r'n\s*[=:]\s*(\d+)', r'sample size\s*[=:]\s*(\d+)', r'patients?\s*[=:]\s*(\d+)']
    max_sample = 0
    for pattern in sample_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                max_sample = max(max_sample, int(match))
            except:
                pass
    
    if max_sample > 0:
        if max_sample > 1000:
            score += 20
            factors.append(f"Muestra grande: n={max_sample}")
        elif max_sample > 500:
            score += 15
            factors.append(f"Muestra moderada: n={max_sample}")
        elif max_sample > 200:
            score += 10
            factors.append(f"Muestra adecuada: n={max_sample}")
        elif max_sample > 100:
            score += 5
            factors.append(f"Muestra pequeña: n={max_sample}")
    
    if any(term in text_lower for term in ['multicenter', 'multi-center']):
        score += 10
        factors.append("Estudio multicéntrico")
    
    if any(term in text_lower for term in ['double-blind', 'double blind', 'masked']):
        score += 15
        factors.append("Con enmascaramiento/cego")
    
    if any(re.search(pattern, text_lower) for pattern in [r'nct\d+', r'clinicaltrials\.gov']):
        score += 5
        factors.append("Ensayo registrado")
    
    if any(term in text_lower for term in ['intention-to-treat', 'intention to treat']):
        score += 10
        factors.append("Análisis por intención de tratar")
    
    return min(score, 100), factors

def get_effect_direction(result_value, ci_lower=None, ci_upper=None):
    """Determina la direccionalidad del efecto"""
    try:
        value = float(result_value)
        if value < 1:
            if ci_upper and float(ci_upper) < 1:
                return "PROTECTOR (IC<1)"
            else:
                return "PROTECTOR (tendencia)"
        elif value > 1:
            if ci_lower and float(ci_lower) > 1:
                return "DAÑINO (IC>1)"
            else:
                return "DAÑINO (tendencia)"
        else:
            return "SIN EFECTO"
    except:
        return "NO DETERMINADA"

def calculate_evidence_strength(statistical_results, quality_score, sample_size=None):
    """Calcula la fuerza de la evidencia"""
    if not statistical_results:
        return "Sin datos estadísticos", 0, {}
    
    strength_score = 0
    directions = {}
    
    for i, result in enumerate(statistical_results):
        if result['type'] in ['HR', 'RR', 'OR']:
            effect = result['value']
            if effect < 0.5 or effect > 2.0:
                strength_score += 30
            elif effect < 0.7 or effect > 1.5:
                strength_score += 20
            elif effect != 1.0:
                strength_score += 10
            
            directions[f"result_{i}"] = get_effect_direction(effect, result.get('ci_lower'), result.get('ci_upper'))
        
        if result['type'] == 'p-value':
            if result['value'] < 0.001:
                strength_score += 25
            elif result['value'] < 0.01:
                strength_score += 20
            elif result['value'] < 0.05:
                strength_score += 15
    
    strength_score = strength_score * (quality_score / 50)
    
    if strength_score >= 50:
        strength = "EVIDENCIA FUERTE"
    elif strength_score >= 30:
        strength = "EVIDENCIA MODERADA"
    elif strength_score >= 15:
        strength = "EVIDENCIA DÉBIL"
    else:
        strength = "EVIDENCIA INSUFICIENTE"
    
    return strength, min(strength_score, 100), directions

def analyze_article_with_ai(title, abstract):
    """Analiza un artículo con IA y extrae toda la información"""
    if not title and not abstract:
        return {}
    
    full_text = f"{title} {abstract if abstract else ''}"
    processed_text = preprocess_text(full_text)
    
    entities = extract_medical_entities(full_text)
    
    study_types = []
    study_keywords = {
        'RCT': ['randomized', 'rct', 'randomised', 'trial', 'double-blind'],
        'Cohorte': ['cohort', 'follow-up', 'longitudinal', 'prospective', 'retrospective'],
        'Meta-análisis': ['meta-analysis', 'meta analysis', 'systematic review'],
        'Observacional': ['observational', 'registry', 'real-world'],
        'Caso-control': ['case-control', 'matched']
    }
    
    for study_type, keywords in study_keywords.items():
        if any(keyword in processed_text for keyword in keywords):
            study_types.append(study_type)
    
    words = processed_text.split()
    word_freq = Counter(words)
    top_keywords = [word for word, _ in word_freq.most_common(20) if len(word) > 3][:10]
    
    population_patterns = {
        'adultos': r'\badults?\b|\bpatients?\b',
        'ancianos': r'\belderly\b|\baged\b|\bolder\b',
        'mujeres': r'\bwomen\b|\bfemale\b',
        'hombres': r'\bmen\b|\bmale\b'
    }
    
    population = [pop for pop, pattern in population_patterns.items() if re.search(pattern, processed_text)]
    
    numeric_results = extract_numeric_results(full_text)
    all_outcomes = extract_all_outcomes(full_text)
    
    quality_score, quality_factors = enhanced_quality_scoring({'study_types': study_types}, full_text)
    
    evidence_strength, evidence_score, directions = calculate_evidence_strength(numeric_results, quality_score)
    
    sentiment = analyze_sentiment_by_outcome(full_text)
    pico = extract_pico_elements(full_text)
    
    semantic_analysis = {}
    if AI_EMBEDDINGS_AVAILABLE and BIOMED_EMBEDDER:
        key_concepts = ['efficacy', 'safety', 'mortality', 'bleeding']
        semantic_analysis = semantic_similarity_analysis(full_text, key_concepts)
    
    rob_assessment = assess_risk_of_bias(full_text, ', '.join(study_types))
    
    # Formatear resultados numéricos como string
    numeric_results_str = ' | '.join([f"{r['type']}={r['value']}" for r in numeric_results[:3]])
    
    return {
        'entities': entities,
        'study_types': ', '.join(study_types) if study_types else 'No especificado',
        'top_keywords': ', '.join(top_keywords),
        'population': ', '.join(population) if population else 'No especificada',
        'quality_score': quality_score,
        'quality_factors': ' | '.join(quality_factors) if quality_factors else 'Sin factores',
        'pico': str(pico),
        'numeric_results': numeric_results,
        'numeric_results_str': numeric_results_str,
        'outcomes_analysis': sentiment,
        'all_outcomes': all_outcomes,
        'evidence_strength': evidence_strength,
        'evidence_score': evidence_score,
        'effect_directions': directions,
        'semantic_analysis': semantic_analysis,
        'risk_of_bias': rob_assessment
    }

def extract_article_info(doc_sum):
    """Extrae la información básica de un artículo desde DocSum"""
    article = {}
    article["pmid"] = doc_sum.find("Id").text if doc_sum.find("Id") is not None else "N/A"
    
    title_item = doc_sum.find(".//Item[@Name='Title']")
    article["title"] = title_item.text if title_item is not None else "Sin título"
    
    author_items = doc_sum.findall(".//Item[@Name='Author']")
    authors = [author.text for author in author_items if author.text]
    article["authors"] = ", ".join(authors[:5]) + (" et al." if len(authors) > 5 else "")
    
    source_item = doc_sum.find(".//Item[@Name='Source']")
    article["journal"] = source_item.text if source_item is not None else "N/A"
    
    pubdate_item = doc_sum.find(".//Item[@Name='PubDate']")
    article["pubdate"] = pubdate_item.text if pubdate_item is not None else "N/A"
    
    doi_item = doc_sum.find(".//Item[@Name='DOI']")
    article["doi"] = doi_item.text if doi_item is not None else "N/A"
    
    return article

def search_pubmed(query, retmax=500):
    """Busca artículos en PubMed"""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    
    search_url = f"{base_url}esearch.fcgi"
    search_params = {
        "db": "pubmed",
        "term": query,
        "retmax": retmax,
        "retmode": "xml",
        "sort": "relevance"
    }
    
    try:
        response = requests.get(search_url, params=search_params, timeout=30)
        response.raise_for_status()
        
        root = ElementTree.fromstring(response.content)
        id_list = [id_elem.text for id_elem in root.findall(".//Id")]
        count = root.find(".//Count").text if root.find(".//Count") is not None else "0"
        
        return id_list, int(count)
    except Exception as e:
        st.error(f"Error en búsqueda: {e}")
        return [], 0

def fetch_articles_details(id_list, max_articles=500):
    """Obtiene los detalles de los artículos y los analiza"""
    if not id_list:
        return []
    
    total_to_process = min(len(id_list), max_articles)
    batch_size = 50
    num_batches = math.ceil(total_to_process / batch_size)
    
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    articles = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for batch_num in range(num_batches):
        start_idx = batch_num * batch_size
        end_idx = min((batch_num + 1) * batch_size, total_to_process)
        batch_ids = id_list[start_idx:end_idx]
        
        status_text.text(f"📦 Procesando lote {batch_num + 1} de {num_batches}...")
        
        summary_params = {
            "db": "pubmed",
            "id": ",".join(batch_ids),
            "retmode": "xml"
        }
        
        try:
            summary_response = requests.get(f"{base_url}esummary.fcgi", params=summary_params, timeout=30)
            summary_response.raise_for_status()
            summary_root = ElementTree.fromstring(summary_response.content)
            
            for j, doc_sum in enumerate(summary_root.findall(".//DocSum")):
                overall_idx = start_idx + j
                progress_bar.progress((overall_idx + 1) / total_to_process)
                
                article = extract_article_info(doc_sum)
                abstract = get_abstract(article["pmid"])
                article["abstract"] = abstract if abstract else "No disponible"
                
                ai_analysis = analyze_article_with_ai(article["title"], abstract)
                article.update(ai_analysis)
                
                articles.append(article)
                time.sleep(0.05)
                
        except Exception as e:
            st.warning(f"Error en lote {batch_num + 1}: {str(e)[:100]}")
            continue
    
    progress_bar.empty()
    status_text.empty()
    
    return articles

def calculate_relevance_to_hypothesis(articles, hypothesis):
    """Calcula la relevancia semántica de cada artículo a la hipótesis"""
    if not AI_EMBEDDINGS_AVAILABLE or not BIOMED_EMBEDDER or not hypothesis:
        return articles
    
    try:
        hypothesis_embedding = BIOMED_EMBEDDER.encode([hypothesis])[0]
        
        for article in articles:
            text = f"{article.get('title', '')} {article.get('abstract', '')}"
            if text and len(text) > 50:
                article_embedding = BIOMED_EMBEDDER.encode([text[:1000]])[0]
                similarity = cosine_similarity([hypothesis_embedding], [article_embedding])[0][0]
                article['relevance_score'] = float(similarity)
            else:
                article['relevance_score'] = 0.0
        
        articles.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
    except Exception as e:
        st.warning(f"Error calculando relevancia: {e}")
        for article in articles:
            article['relevance_score'] = 0.5
    
    return articles

def discover_flavors_by_embeddings(articles):
    """Descubre flavors agrupando artículos por similitud semántica"""
    if len(articles) < 3 or not AI_EMBEDDINGS_AVAILABLE or not BIOMED_EMBEDDER:
        return []
    
    try:
        texts = [f"{a.get('title', '')} {a.get('outcomes_analysis', '')}" for a in articles]
        embeddings = BIOMED_EMBEDDER.encode(texts)
        
        n_clusters = min(max(2, len(articles) // 5), 8)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        cluster_labels = kmeans.fit_predict(embeddings)
        
        flavors = []
        for cluster_id in range(n_clusters):
            cluster_indices = [i for i, label in enumerate(cluster_labels) if label == cluster_id]
            if len(cluster_indices) < 2:
                continue
            
            cluster_articles = [articles[i] for i in cluster_indices]
            flavor = characterize_cluster(cluster_articles, cluster_id)
            if flavor:
                flavors.append(flavor)
        
        return flavors
    except Exception as e:
        st.warning(f"Error en clustering semántico: {e}")
        return []

def discover_flavors_by_outcomes(articles):
    """Agrupa artículos por outcomes compartidos"""
    if len(articles) < 2:
        return []
    
    try:
        all_outcomes = set()
        for article in articles:
            outcomes = article.get('all_outcomes', [])
            all_outcomes.update(outcomes)
        
        if len(all_outcomes) < 2:
            return []
        
        all_outcomes_list = list(all_outcomes)
        outcome_matrix = []
        for article in articles:
            outcomes = set(article.get('all_outcomes', []))
            row = [1 if o in outcomes else 0 for o in all_outcomes_list]
            outcome_matrix.append(row)
        
        n_clusters = min(max(2, len(articles) // 4), 6)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clusters = kmeans.fit_predict(outcome_matrix)
        
        flavors = []
        for cluster_id in range(n_clusters):
            cluster_indices = [i for i, c in enumerate(clusters) if c == cluster_id]
            if len(cluster_indices) < 2:
                continue
            
            cluster_articles = [articles[i] for i in cluster_indices]
            
            outcome_counts = Counter()
            for article in cluster_articles:
                outcomes = article.get('all_outcomes', [])
                outcome_counts.update(outcomes)
            
            top_outcomes = [o for o, _ in outcome_counts.most_common(3)]
            
            flavor = {
                'type': 'outcome_based',
                'id': f"outcome_cluster_{cluster_id}",
                'name': f"Outcomes: {', '.join(top_outcomes)}",
                'top_outcomes': top_outcomes,
                'articles': cluster_articles,
                'n_articles': len(cluster_articles),
                'avg_quality': np.mean([a.get('quality_score', 0) for a in cluster_articles])
            }
            flavors.append(flavor)
        
        return flavors
    except Exception as e:
        st.warning(f"Error en clustering por outcomes: {e}")
        return []

def discover_flavors_by_effect_direction(articles):
    """Agrupa artículos por dirección del efecto"""
    flavors = []
    
    direction_groups = {
        'beneficio': [],
        'riesgo': [],
        'sin_efecto': [],
        'contradictorio': []
    }
    
    for article in articles:
        directions = article.get('effect_directions', {})
        direction_values = list(directions.values()) if directions else []
        
        direction_str = ' '.join(direction_values).lower()
        
        if 'protector' in direction_str and 'dañino' not in direction_str:
            direction_groups['beneficio'].append(article)
        elif 'dañino' in direction_str and 'protector' not in direction_str:
            direction_groups['riesgo'].append(article)
        elif 'sin efecto' in direction_str:
            direction_groups['sin_efecto'].append(article)
        else:
            direction_groups['contradictorio'].append(article)
    
    direction_names = {
        'beneficio': 'Evidencia de beneficio',
        'riesgo': 'Evidencia de riesgo',
        'sin_efecto': 'Sin efecto significativo',
        'contradictorio': 'Resultados contradictorios'
    }
    
    for key, group_articles in direction_groups.items():
        if group_articles:
            flavors.append({
                'type': 'effect_direction',
                'id': f"effect_{key}",
                'name': direction_names[key],
                'articles': group_articles,
                'n_articles': len(group_articles),
                'avg_quality': np.mean([a.get('quality_score', 0) for a in group_articles])
            })
    
    return flavors

def discover_flavors_by_population(articles):
    """Agrupa artículos por población estudiada"""
    population_patterns = {
        'ancianos': ['elderly', 'aged', 'older', 'geriatric', '>75', '≥75', '> 75', '≥ 75'],
        'diabetes': ['diabetes', 'diabetic', 'dm', 'type 2 diabetes', 'type 1 diabetes'],
        'insuficiencia_renal': ['renal', 'kidney', 'ckd', 'creatinine', 'chronic kidney disease'],
        'mujeres': ['women', 'female', 'gender'],
        'hombres': ['men', 'male']
    }
    
    flavors = []
    
    for pop, patterns in population_patterns.items():
        pop_articles = []
        for article in articles:
            text = f"{article.get('title', '')} {article.get('abstract', '')}".lower()
            if any(pattern in text for pattern in patterns):
                pop_articles.append(article)
        
        if pop_articles:
            pop_names = {
                'ancianos': 'Población anciana', 
                'diabetes': 'Pacientes con diabetes', 
                'insuficiencia_renal': 'Insuficiencia renal', 
                'mujeres': 'Población femenina',
                'hombres': 'Población masculina'
            }
            
            flavors.append({
                'type': 'population',
                'id': f"pop_{pop}",
                'name': pop_names.get(pop, pop),
                'articles': pop_articles,
                'n_articles': len(pop_articles),
                'avg_quality': np.mean([a.get('quality_score', 0) for a in pop_articles])
            })
    
    return flavors

def characterize_cluster(articles, cluster_id):
    """Caracteriza un cluster de artículos"""
    if not articles:
        return None
    
    all_keywords = []
    all_outcomes = []
    directions = []
    
    for a in articles:
        keywords = a.get('top_keywords', '').split(', ')
        all_keywords.extend(keywords)
        outcomes = a.get('all_outcomes', [])
        all_outcomes.extend(outcomes)
        
        dirs = a.get('effect_directions', {})
        for d in dirs.values():
            if 'PROTECTOR' in d:
                directions.append('beneficio')
            elif 'DAÑINO' in d:
                directions.append('riesgo')
            elif 'SIN EFECTO' in d:
                directions.append('sin_efecto')
    
    top_keywords = Counter(all_keywords).most_common(3)
    top_outcomes = Counter(all_outcomes).most_common(3)
    dominant_direction = Counter(directions).most_common(1)[0][0] if directions else 'mixto'
    avg_quality = np.mean([a.get('quality_score', 0) for a in articles])
    
    representative = sorted(articles, key=lambda x: (x.get('quality_score', 0), x.get('relevance_score', 0)), reverse=True)[:3]
    
    direction_names = {'beneficio': 'beneficio', 'riesgo': 'riesgo', 'sin_efecto': 'sin efecto', 'mixto': 'mixto'}
    
    return {
        'type': 'semantic',
        'id': f"semantic_cluster_{cluster_id}",
        'name': f"{', '.join([k for k, _ in top_keywords[:2]])} | {', '.join([o for o, _ in top_outcomes[:2]])}",
        'top_keywords': [k for k, _ in top_keywords],
        'top_outcomes': [o for o, _ in top_outcomes],
        'dominant_direction': direction_names.get(dominant_direction, dominant_direction),
        'avg_quality': avg_quality,
        'representative_articles': representative,
        'articles': articles,
        'n_articles': len(articles)
    }

def generate_citation(article, index):
    """Genera una cita bibliográfica en formato Vancouver"""
    authors = article.get('authors', 'Autor')
    year = article.get('pubdate', 's.f.')[:4] if article.get('pubdate') else 's.f.'
    title = article.get('title', 'Sin título')
    journal = article.get('journal', 'Revista')
    pmid = article.get('pmid', '')
    
    citation = f"{index}. {authors}. {title}. {journal}. {year}"
    if pmid and pmid != 'N/A':
        citation += f"; PMID: {pmid}"
    
    return citation

def generate_introduction_paragraph(flavor, hypothesis):
    """Genera un párrafo para la introducción"""
    n = flavor['n_articles']
    outcomes = flavor.get('top_outcomes', ['resultados'])
    keywords = flavor.get('top_keywords', ['estudios'])
    representative = flavor.get('representative_articles', flavor['articles'][:3] if flavor['articles'] else [])
    
    direction_text = {
        'beneficio': 'ha demostrado consistentemente beneficios',
        'riesgo': 'ha identificado riesgos significativos',
        'sin_efecto': 'no ha encontrado diferencias significativas',
        'mixto': 'ha mostrado resultados heterogéneos'
    }.get(flavor.get('dominant_direction', 'mixto'), 'ha sido investigada')
    
    paragraph = f"""<b>{flavor['name']}</b><br/><br/>
La evidencia disponible sobre {', '.join(keywords[:3])} {direction_text}. 
Se han identificado {n} estudios que abordan específicamente los outcomes de 
{', '.join(outcomes[:2])}.<br/><br/>"""
    
    if representative:
        paragraph += f"""Entre los hallazgos más relevantes, {representative[0].get('authors', 'un estudio')} 
reportó {representative[0].get('numeric_results_str', 'resultados significativos')} 
(calidad: {representative[0].get('quality_score', 'N/A')}/100). 
Estos resultados fueron consistentes con los hallazgos de """
        
        if len(representative) > 1:
            paragraph += f"""{representative[1].get('authors', 'estudios posteriores')} """
        if len(representative) > 2:
            paragraph += f"""y {representative[2].get('authors', 'otros investigadores')}."""
        else:
            paragraph += """estudios previos."""
    
    paragraph += f"""<br/><br/>En el contexto de la hipótesis planteada: "{hypothesis}", 
los hallazgos de este conjunto de estudios proporcionan evidencia {'favorable' if flavor.get('dominant_direction') == 'beneficio' else 'desfavorable' if flavor.get('dominant_direction') == 'riesgo' else 'limitada'} 
que respalda la necesidad de continuar investigando en esta área.<br/><br/>"""
    
    citations = []
    for i, article in enumerate(representative, 1):
        citations.append(generate_citation(article, i))
    
    return paragraph, citations

def generate_discussion_paragraph(flavor, hypothesis):
    """Genera un párrafo para la discusión"""
    n = flavor['n_articles']
    outcomes = flavor.get('top_outcomes', ['resultados'])
    representative = flavor.get('representative_articles', flavor['articles'][:3] if flavor['articles'] else [])
    
    direction_compare = {
        'beneficio': 'coinciden con',
        'riesgo': 'son consistentes con',
        'sin_efecto': 'son similares a',
        'mixto': 'presentan similitudes y diferencias con'
    }.get(flavor.get('dominant_direction', 'mixto'), 'se comparan con')
    
    paragraph = f"""<b>{flavor['name']}</b><br/><br/>
Los resultados de nuestro análisis {direction_compare} los hallazgos previamente reportados en la literatura.
En un conjunto de {n} estudios que evaluaron {', '.join(outcomes[:2])}, 
la dirección del efecto fue predominantemente {'favorable' if flavor.get('dominant_direction') == 'beneficio' else 'desfavorable' if flavor.get('dominant_direction') == 'riesgo' else 'neutral'}.<br/><br/>"""
    
    if representative:
        paragraph += f"""Específicamente, {representative[0].get('authors', 'estudios previos')} 
observaron {representative[0].get('numeric_results_str', 'efectos significativos')}, 
lo cual es consistente con lo reportado por """
        
        if len(representative) > 1:
            paragraph += f"""{representative[1].get('authors', 'investigaciones posteriores')}. """
        if len(representative) > 2:
            paragraph += f"""Por otro lado, {representative[2].get('authors', 'algunos estudios')} 
reportaron resultados divergentes que merecen atención.<br/><br/>"""
    
    paragraph += f"""En relación con la hipótesis planteada: "{hypothesis}", 
la evidencia agregada sugiere que {'existe un respaldo consistente' if flavor.get('dominant_direction') == 'beneficio' else 'existen señales de alerta' if flavor.get('dominant_direction') == 'riesgo' else 'la evidencia es insuficiente para establecer conclusiones firmes'}.
Estos hallazgos deben interpretarse considerando las limitaciones metodológicas de los estudios incluidos.<br/><br/>"""
    
    citations = []
    for i, article in enumerate(representative, 1):
        citations.append(generate_citation(article, i))
    
    return paragraph, citations

def generate_general_paragraph(flavor, hypothesis):
    """Genera un párrafo general"""
    n = flavor['n_articles']
    outcomes = flavor.get('top_outcomes', ['resultados'])
    representative = flavor.get('representative_articles', flavor['articles'][:3] if flavor['articles'] else [])
    
    paragraph = f"""<b>{flavor['name']}</b><br/><br/>
Este flavor agrupa {n} artículos que abordan {', '.join(outcomes[:2])}. 
La calidad metodológica promedio es de {flavor.get('avg_quality', 0):.0f}/100.<br/><br/>"""
    
    if representative:
        paragraph += f"""Artículos representativos:<br/>"""
        for i, article in enumerate(representative, 1):
            paragraph += f"""{i}. {article.get('authors', 'Autor')} - {article.get('title', 'Título')[:100]}...<br/>"""
    
    citations = []
    for i, article in enumerate(representative, 1):
        citations.append(generate_citation(article, i))
    
    return paragraph, citations

def generate_all_flavors(articles, hypothesis):
    """Genera todos los flavors desde múltiples perspectivas"""
    if not articles:
        return {}
    
    all_flavors = {
        'semantic_clusters': [],
        'outcome_clusters': [],
        'effect_clusters': [],
        'population_clusters': []
    }
    
    if len(articles) >= 3 and AI_EMBEDDINGS_AVAILABLE:
        semantic_flavors = discover_flavors_by_embeddings(articles)
        for flavor in semantic_flavors:
            if flavor:
                all_flavors['semantic_clusters'].append(flavor)
    
    if len(articles) >= 2:
        all_flavors['outcome_clusters'] = discover_flavors_by_outcomes(articles)
        all_flavors['effect_clusters'] = discover_flavors_by_effect_direction(articles)
        all_flavors['population_clusters'] = discover_flavors_by_population(articles)
    
    # Generar párrafos para cada flavor
    for category in all_flavors.values():
        for flavor in category:
            if flavor:
                flavor['intro_paragraph'], flavor['intro_citations'] = generate_introduction_paragraph(flavor, hypothesis)
                flavor['discussion_paragraph'], flavor['discussion_citations'] = generate_discussion_paragraph(flavor, hypothesis)
                flavor['general_paragraph'], flavor['general_citations'] = generate_general_paragraph(flavor, hypothesis)
    
    return all_flavors

def create_document_with_flavors(flavors, hypothesis, query, total_articles):
    """Crea un documento DOCX con los flavors generados"""
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título
    title = doc.add_heading('Análisis de Literatura: Flavors Temáticos', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información de la búsqueda
    doc.add_heading('Información de la Búsqueda', level=1)
    doc.add_paragraph(f'Estrategia de búsqueda en PubMed:')
    p = doc.add_paragraph(query, style='Intense Quote')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    
    doc.add_paragraph(f'Hipótesis/Conjetura: "{hypothesis}"')
    doc.add_paragraph(f'Total de artículos analizados: {total_articles}')
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    
    doc.add_page_break()
    
    # Introducción
    doc.add_heading('Flavors para la Sección de Introducción', level=1)
    doc.add_paragraph('Los siguientes párrafos están organizados por flavors temáticos y pueden ser utilizados en la sección de Introducción de un artículo científico.')
    
    for category_name, flavor_list in flavors.items():
        if flavor_list:
            category_title = category_name.replace('_', ' ').title()
            doc.add_heading(category_title, level=2)
            
            for flavor in flavor_list:
                if not flavor:
                    continue
                    
                doc.add_heading(flavor['name'], level=3)
                
                # Párrafo de introducción (quitar etiquetas HTML)
                intro_text = flavor['intro_paragraph'].replace('<br/>', '\n').replace('<b>', '').replace('</b>', '')
                p = doc.add_paragraph(intro_text)
                p.paragraph_format.space_after = Pt(6)
                
                # Referencias
                if flavor.get('intro_citations'):
                    doc.add_paragraph('Referencias:', style='List Bullet')
                    for citation in flavor['intro_citations']:
                        doc.add_paragraph(citation, style='List Bullet')
                
                doc.add_paragraph()  # Espacio
    
    doc.add_page_break()
    
    # Discusión
    doc.add_heading('Flavors para la Sección de Discusión', level=1)
    doc.add_paragraph('Los siguientes párrafos están organizados por flavors temáticos y pueden ser utilizados en la sección de Discusión de un artículo científico.')
    
    for category_name, flavor_list in flavors.items():
        if flavor_list:
            category_title = category_name.replace('_', ' ').title()
            doc.add_heading(category_title, level=2)
            
            for flavor in flavor_list:
                if not flavor:
                    continue
                    
                doc.add_heading(flavor['name'], level=3)
                
                # Párrafo de discusión
                disc_text = flavor['discussion_paragraph'].replace('<br/>', '\n').replace('<b>', '').replace('</b>', '')
                p = doc.add_paragraph(disc_text)
                p.paragraph_format.space_after = Pt(6)
                
                # Referencias
                if flavor.get('discussion_citations'):
                    doc.add_paragraph('Referencias:', style='List Bullet')
                    for citation in flavor['discussion_citations']:
                        doc.add_paragraph(citation, style='List Bullet')
                
                doc.add_paragraph()  # Espacio
    
    doc.add_page_break()
    
    # Resumen general
    doc.add_heading('Resumen General de Flavors', level=1)
    doc.add_paragraph('A continuación se presenta un resumen general de cada flavor con sus artículos representativos.')
    
    for category_name, flavor_list in flavors.items():
        if flavor_list:
            category_title = category_name.replace('_', ' ').title()
            doc.add_heading(category_title, level=2)
            
            for flavor in flavor_list:
                if not flavor:
                    continue
                    
                doc.add_heading(flavor['name'], level=3)
                
                # Información general
                doc.add_paragraph(f'Número de artículos: {flavor["n_articles"]}')
                doc.add_paragraph(f'Calidad promedio: {flavor.get("avg_quality", 0):.0f}/100')
                
                if flavor.get('top_outcomes'):
                    doc.add_paragraph(f'Outcomes principales: {", ".join(flavor["top_outcomes"])}')
                if flavor.get('dominant_direction'):
                    doc.add_paragraph(f'Dirección dominante: {flavor["dominant_direction"]}')
                
                # Artículos representativos
                if flavor.get('representative_articles'):
                    doc.add_paragraph('Artículos representativos:')
                    for i, article in enumerate(flavor['representative_articles'][:5], 1):
                        doc.add_paragraph(
                            f"{i}. {article.get('authors', 'Autor')} - {article.get('title', 'Título')[:150]}... "
                            f"(PMID: {article.get('pmid', 'N/A')}, Calidad: {article.get('quality_score', 0)})",
                            style='List Bullet'
                        )
                
                doc.add_paragraph()  # Espacio
    
    return doc

# ============================================================================
# FUNCIÓN PRINCIPAL
# ============================================================================

def main():
    st.title("🧠 PubMed AI Analyzer - Generador de Flavors")
    
    if IN_STREAMLIT_CLOUD:
        st.info("☁️ Ejecutándose en Streamlit Cloud - Modo optimizado")
    
    st.markdown("""
    ### Genera párrafos temáticos (flavors) para tu artículo científico
    
    **El programa:**
    - 🔍 Busca artículos en PubMed según tu estrategia
    - 🧠 Analiza cada artículo con IA (BioBERT)
    - 📊 Agrupa los artículos en flavors temáticos
    - ✍️ Genera párrafos listos para Introducción y Discusión
    - 📚 Incluye referencias bibliográficas en formato Vancouver
    - 📄 Descarga un archivo DOCX con todo el contenido
    """)
    
    # Mostrar estado de módulos
    col_a, col_b = st.columns(2)
    with col_a:
        if AI_EMBEDDINGS_AVAILABLE:
            st.success("✅ BioBERT disponible (análisis semántico)")
        else:
            st.error("❌ BioBERT no disponible - algunas funciones limitadas")
    with col_b:
        st.info("📄 Formato de salida: DOCX con referencias Vancouver")
    
    # Entradas del usuario
    st.markdown("---")
    st.markdown("### 📝 Configuración de la búsqueda")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        query = st.text_area(
            "**Estrategia de búsqueda en PubMed:**",
            value="((\"Ticagrelor\"[Mesh]) OR (ticagrelor)) AND ((\"Acute Coronary Syndrome\"[Mesh]) OR (\"Myocardial Infarction\"[Mesh])) AND ((randomized controlled trial[pt]) OR (cohort studies[mesh]))",
            height=100,
            help="Pega aquí tu estrategia de búsqueda de PubMed con sintaxis MeSH"
        )
    
    with col2:
        max_articles = st.slider(
            "Máx. artículos a analizar:",
            min_value=50,
            max_value=500,
            value=200,
            step=50,
            help="Mayor número = más tiempo de procesamiento"
        )
    
    hypothesis = st.text_area(
        "**📌 Tu conjetura/hipótesis (lenguaje natural):**",
        value="Ticagrelor es superior a clopidogrel en la reducción de eventos cardiovasculares mayores en pacientes con síndrome coronario agudo sometidos a intervención coronaria percutánea, con un perfil de seguridad aceptable sin aumento significativo de sangrado mayor.",
        height=100,
        help="Escribe tu hipótesis en lenguaje natural. El programa la usará para priorizar artículos relevantes."
    )
    
    generate_button = st.button("🚀 GENERAR FLAVORS Y DOCUMENTO", type="primary", use_container_width=True)
    
    # Session state
    if 'docx_generated' not in st.session_state:
        st.session_state.docx_generated = False
        st.session_state.docx_data = None
        st.session_state.total_articles = 0
        st.session_state.n_flavors = 0
    
    if generate_button:
        if not query.strip():
            st.warning("⚠️ Por favor, introduce una estrategia de búsqueda")
        elif not hypothesis.strip():
            st.warning("⚠️ Por favor, introduce tu conjetura/hipótesis")
        else:
            start_time = time.time()
            
            with st.spinner("🔍 Buscando artículos en PubMed..."):
                id_list, total_count = search_pubmed(query.strip(), retmax=max_articles)
                
                if not id_list:
                    st.error("❌ No se encontraron artículos para esta búsqueda")
                    st.stop()
                
                st.info(f"📊 Se encontraron {total_count} artículos. Procesando {len(id_list)}...")
                
                articles = fetch_articles_details(id_list, max_articles=max_articles)
            
            if not articles:
                st.error("❌ No se pudo procesar ningún artículo")
                st.stop()
            
            st.success(f"✅ Procesados {len(articles)} artículos")
            
            with st.spinner("🧠 Calculando relevancia con la hipótesis..."):
                articles = calculate_relevance_to_hypothesis(articles, hypothesis)
                
                # Filtrar artículos con relevancia > 0.5 si hay suficientes
                relevant_articles = [a for a in articles if a.get('relevance_score', 0) > 0.5]
                if len(relevant_articles) >= 5:
                    articles = relevant_articles
                    st.info(f"📌 Filtrados a {len(articles)} artículos con alta relevancia a la hipótesis (>0.5)")
                else:
                    st.info(f"📌 Se mantienen {len(articles)} artículos (relevancia >0.3)")
                    articles = [a for a in articles if a.get('relevance_score', 0) > 0.3]
            
            with st.spinner("🔍 Descubriendo flavors temáticos..."):
                flavors = generate_all_flavors(articles, hypothesis)
            
            # Contar flavors totales
            total_flavors = sum(len(flavor_list) for flavor_list in flavors.values())
            
            st.success(f"✅ Generados {total_flavors} flavors desde {len(articles)} artículos")
            
            with st.spinner("📄 Creando documento DOCX..."):
                doc = create_document_with_flavors(flavors, hypothesis, query, len(articles))
                
                docx_bytes = BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)
                
                st.session_state.docx_data = docx_bytes
                st.session_state.docx_generated = True
                st.session_state.total_articles = len(articles)
                st.session_state.n_flavors = total_flavors
            
            elapsed_time = time.time() - start_time
            st.info(f"⏱️ Tiempo total de procesamiento: {elapsed_time/60:.1f} minutos")
    
    # Mostrar resultados y botón de descarga
    if st.session_state.docx_generated and st.session_state.docx_data is not None:
        st.markdown("---")
        st.markdown("### 📥 Documento Generado - Listo para Descargar")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Artículos procesados", st.session_state.total_articles)
        with col2:
            st.metric("Flavors generados", st.session_state.n_flavors)
        with col3:
            st.metric("Formato", "DOCX con referencias")
        
        st.download_button(
            label="📥 DESCARGAR DOCUMENTO CON FLAVORS",
            data=st.session_state.docx_data,
            file_name=f"flavors_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_docx"
        )
        
        with st.expander("📖 Estructura del documento"):
            st.markdown("""
            **Estructura del documento generado:**
            
            1. **Información de la búsqueda** - Estrategia, hipótesis y metadatos
            2. **Flavors para Introducción** - Párrafos listos con referencias
            3. **Flavors para Discusión** - Párrafos listos con referencias  
            4. **Resumen general** - Características de cada flavor
            
            **Tipos de flavors generados:**
            - **Semánticos**: Agrupados por similitud de contenido
            - **Por outcomes**: Agrupados por resultados medidos
            - **Por dirección del efecto**: Beneficio, riesgo, sin efecto
            - **Por población**: Subgrupos poblacionales específicos
            
            Los párrafos incluyen citas bibliográficas en formato Vancouver.
            """)
    
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: gray; font-size: 0.8em;'>
            🧠 PubMed AI Analyzer - Generador de Flavors v5.0<br>
            Análisis semántico con BioBERT | Agrupamiento temático automático | Referencias Vancouver
        </div>
        """,
        unsafe_allow_html=True
    )

# ============================================================================
# PUNTO DE ENTRADA
# ============================================================================

if __name__ == "__main__":
    main()
